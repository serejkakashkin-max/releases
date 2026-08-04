from __future__ import annotations

import io
import json
import os
import re
import tempfile
import time
import unittest
from unittest import mock
import zipfile
from pathlib import Path

from docx import Document
from flask import Flask

from tests._support import PROJECT_ROOT, prepare_config_import

prepare_config_import()

from routes.document_template_routes import document_template_bp
from services.document_template_candidate_service import validate_staged_candidate
from services.document_template_csrf_service import (
    CSRF_COOKIE_NAME, DOCUMENT_TEMPLATE_ACTOR, csrf_cookie_path,
)
from services.document_template_read_service import build_document_whitelist
from services.document_template_storage_service import (
    get_candidate, list_candidates, list_history, update_candidate,
    write_uploaded_candidate,
)
from services.document_template_publish_service import publish_candidate, recover_stale_operations, rollback_version
from services.document_template_read_service import resolve_document
from services.oplot_ui_service import register_oplot_ui
from services.document_template_validation_service import inspect_docx
from services.release_template_catalog_service import clear_template_catalog_cache


SECRET = "stage2-strong-secret-0123456789abcdef0123456789"


def make_template(path: Path, *, heading: str = "Шаблон") -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph("Версия RELEASE_VERSION")
    document.add_paragraph("Дата DATE")
    table = document.add_table(rows=1, cols=4)
    headers = table.rows[0].cells
    headers[0].text = "№"
    headers[1].text = "ЗНИ/JIRA ID"
    headers[2].text = "Issue"
    headers[3].text = "Issue Type"
    document.save(path)
    return path.read_bytes()


def build_app(root: Path, runtime: Path, *, secret=SECRET, enabled=None) -> Flask:
    app = Flask(__name__, template_folder=str(PROJECT_ROOT / "templates"), static_folder=str(PROJECT_ROOT / "static"))
    app.config.update(
        TESTING=True, SECRET_KEY=secret,
        DOCUMENT_TEMPLATE_CENTER_ROOT=root,
        DOCUMENT_TEMPLATE_CENTER_RUNTIME_ROOT=runtime,
    )
    if enabled is not None:
        app.config["DOCUMENT_TEMPLATE_CENTER_ENABLED"] = enabled
    app.add_url_rule("/", endpoint="main.index", view_func=lambda: "home")
    app.add_url_rule("/help", endpoint="main.help_page", view_func=lambda: "help")
    app.add_url_rule("/dashboard", endpoint="dashboard.dashboard", view_func=lambda: "dashboard")
    app.add_url_rule("/release-monitor", endpoint="dashboard.release_monitor_page", view_func=lambda: "monitor")
    app.add_url_rule("/mpr", endpoint="mpr.mpr_page", view_func=lambda: "mpr")
    app.register_blueprint(document_template_bp)
    register_oplot_ui(app)
    return app


class Stage2AccessBoundaryTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        base = Path(self.temp.name)
        self.root = base / "doc_templates"
        make_template(self.root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx")
        clear_template_catalog_cache()
        self.app = build_app(self.root, base / "runtime")
        self.client = self.app.test_client()

    def tearDown(self):
        clear_template_catalog_cache(); self.temp.cleanup()

    def test_full_and_htmx_catalog_are_directly_available(self):
        full = self.client.get("/dashboard/release-monitor/document-templates")
        htmx = self.client.get("/dashboard/release-monitor/document-templates", headers={"HX-Request": "true"})
        self.assertEqual(200, full.status_code)
        self.assertEqual(200, htmx.status_code)
        self.assertNotIn("HX-Redirect", htmx.headers)

    def test_old_auth_routes_are_absent_and_flag_is_ignored(self):
        for path, method in (("/dashboard/release-monitor/document-templates/login", "get"), ("/dashboard/release-monitor/document-templates/session/login", "post"), ("/dashboard/release-monitor/document-templates/session/logout", "post")):
            self.assertEqual(404, getattr(self.client, method)(path).status_code)
        disabled = build_app(self.root, Path(self.temp.name) / "disabled-runtime", secret=None, enabled=False)
        self.assertEqual(200, disabled.test_client().get("/dashboard/release-monitor/document-templates").status_code)

    def test_cookie_csrf_matches_hidden_field_and_is_not_a_flask_session(self):
        response = self.client.get("/dashboard/release-monitor/document-templates")
        html = response.get_data(as_text=True)
        token = re.search(r'name="_csrf_token" value="([A-Za-z0-9_-]+)"', html).group(1)
        self.assertEqual(token, self.client.get_cookie(CSRF_COOKIE_NAME, path="/dashboard/release-monitor/document-templates/").value)
        cookie = response.headers.get("Set-Cookie", "")
        self.assertIn("HttpOnly", cookie); self.assertIn("SameSite=Lax", cookie)
        self.assertIn("Path=/dashboard/release-monitor/document-templates/", cookie)
        self.assertNotIn("Max-Age", cookie); self.assertNotIn("Expires=", cookie)
        with self.client.session_transaction() as session:
            self.assertFalse([key for key in session if str(key).startswith("document_template_editor_")])

    def test_csrf_cookie_path_is_prefix_safe(self):
        cases = (
            ({}, {}, "/dashboard/release-monitor/document-templates/"),
            ({"SCRIPT_NAME": "/script"}, {}, "/script/dashboard/release-monitor/document-templates/"),
            ({}, {"X-Forwarded-Prefix": "/proxy"}, "/proxy/dashboard/release-monitor/document-templates/"),
        )
        for environ, headers, expected in cases:
            with self.subTest(expected=expected), self.app.test_request_context("/dashboard/release-monitor/document-templates/", environ_overrides=environ, headers=headers):
                self.assertEqual(expected, csrf_cookie_path())
        with mock.patch.dict(os.environ, {"BASE_PATH": "/base"}, clear=False):
            with self.app.test_request_context("/dashboard/release-monitor/document-templates/"):
                self.assertEqual("/base/dashboard/release-monitor/document-templates/", csrf_cookie_path())

    def test_blueprint_csp_and_no_store_are_applied(self):
        response = self.client.get("/dashboard/release-monitor/document-templates")
        self.assertEqual("no-store", response.headers["Cache-Control"])
        self.assertEqual("nosniff", response.headers["X-Content-Type-Options"])
        csp = response.headers["Content-Security-Policy"]
        self.assertIn("connect-src 'self'", csp); self.assertIn("object-src 'none'", csp)
        html = response.get_data(as_text=True)
        self.assertIn("htmx.min.js", html)
        self.assertIn("jszip.min.js", html)
        self.assertIn("docx-preview.min.js", html)
        self.assertNotRegex(html, r'(?:src|href|action|hx-get)=["\'](?:https?:)?//')

class Stage2WorkflowTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        base = Path(self.temp.name)
        self.root = base / "doc_templates"
        self.active_path = self.root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx"
        self.original = make_template(self.active_path, heading="Активный")
        self.candidate = make_template(base / "source.docx", heading="Кандидат")
        clear_template_catalog_cache()
        self.app = build_app(self.root, base / "runtime")
        self.client = self.app.test_client()
        self.client.get("/dashboard/release-monitor/document-templates")
        self.csrf = self.client.get_cookie(CSRF_COOKIE_NAME, path="/dashboard/release-monitor/document-templates/").value
        self.document_id = next(iter(build_document_whitelist(self.root)))

    def tearDown(self):
        clear_template_catalog_cache(); self.temp.cleanup()

    def _upload(self, payload=None):
        return self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates", data={
            "_csrf_token": self.csrf, "comment": "Обновлено оформление",
            "file": (io.BytesIO(payload or self.candidate), "Новая версия.docx"),
        }, content_type="multipart/form-data")

    def _upload_and_validate(self, payload=None):
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            candidate = write_uploaded_candidate(
                io.BytesIO(payload or self.candidate),
                document_id=self.document_id,
                source_filename="Новая версия.docx",
                active_filename=document.filename,
                active_sha=document.sha256,
                uploaded_by=DOCUMENT_TEMPLATE_ACTOR,
                comment="Обновлено оформление",
            )
            validate_staged_candidate(
                self.document_id, candidate["candidate_uuid"], document.path,
            )
        return candidate["candidate_uuid"]

    def test_upload_automatically_validates_publishes_and_preserves_history(self):
        upload = self._upload()
        self.assertEqual(303, upload.status_code)
        self.assertIn("replacement=success", upload.headers["Location"])
        with self.app.app_context():
            candidate_uuid = list_candidates(self.document_id)[0]["candidate_uuid"]
            metadata = get_candidate(candidate_uuid, document_id=self.document_id)
        self.assertEqual("published", metadata["state"])
        preview = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/preview")
        test_preview = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/test-document/preview")
        self.assertEqual(200, preview.status_code)
        self.assertEqual(200, test_preview.status_code)
        self.assertEqual(self.candidate, self.active_path.read_bytes())
        audit_text = (Path(self.temp.name) / "runtime/data/document_template_center/audit/events.jsonl").read_text(encoding="utf-8")
        self.assertIn(DOCUMENT_TEMPLATE_ACTOR, audit_text); self.assertNotIn(self.csrf, audit_text); self.assertNotIn(str(self.root), audit_text)
        self.assertEqual("publish", json.loads(audit_text.splitlines()[-1])["action"])
        with self.app.app_context():
            versions = list_history(self.document_id)
        self.assertEqual(1, len(versions))
        self.assertEqual("Обновлено оформление", versions[0]["comment"])
        self.assertEqual(self.original, (Path(self.temp.name) / f"runtime/data/document_template_center/history/{self.document_id}/{versions[0]['version_uuid']}/document.docx").read_bytes())
        rolled = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/history/{versions[0]['version_uuid']}/rollback", data={"_csrf_token": self.csrf, "reason": "Возвращаем проверенный вариант", "expected_active_sha": metadata["candidate_sha"]})
        self.assertEqual(303, rolled.status_code, rolled.get_data(as_text=True))
        self.assertEqual(self.original, self.active_path.read_bytes())

    def test_multiple_candidates_become_conflicted_after_publish(self):
        first = self._upload_and_validate()
        second_bytes = make_template(Path(self.temp.name) / "second.docx", heading="Второй кандидат")
        second = self._upload_and_validate(second_bytes)
        response = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{first}/publish", data={"_csrf_token": self.csrf})
        self.assertEqual(303, response.status_code)
        with self.app.app_context():
            self.assertEqual("published", get_candidate(first)["state"])
            self.assertEqual("conflict", get_candidate(second)["state"])

    def test_automatic_replacement_preserves_external_change_on_sha_conflict(self):
        external = make_template(
            Path(self.temp.name) / "external.docx", heading="External change",
        )

        def validate_then_change(*args, **kwargs):
            result = validate_staged_candidate(*args, **kwargs)
            self.active_path.write_bytes(external)
            return result

        with mock.patch(
            "routes.document_template_routes.validate_staged_candidate",
            side_effect=validate_then_change,
        ):
            response = self._upload()
        self.assertEqual(303, response.status_code)
        self.assertNotIn("replacement=success", response.headers["Location"])
        self.assertEqual(external, self.active_path.read_bytes())
        with self.app.app_context():
            candidate = list_candidates(self.document_id)[0]
            self.assertEqual("conflict", candidate["state"])
            self.assertEqual([], list_history(self.document_id))

    def test_recovery_never_repeats_replace(self):
        candidate_uuid = self._upload_and_validate()
        class SimulatedCrash(BaseException):
            pass
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            def injector(point):
                if point == "after_replace":
                    raise SimulatedCrash()
            with self.assertRaises(SimulatedCrash):
                publish_candidate(document, candidate_uuid, "Редактор", failure_injector=injector)
            changed = self.active_path.read_bytes()
            outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 180)
            self.assertEqual(1, outcome["published"])
            self.assertEqual(changed, self.active_path.read_bytes())
            self.assertEqual("published", get_candidate(candidate_uuid)["state"])

    def test_prepared_history_crash_recovers_without_replace(self):
        candidate_uuid = self._upload_and_validate()
        class SimulatedCrash(BaseException):
            pass
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            def injector(point):
                if point == "after_prepared_history":
                    raise SimulatedCrash()
            with self.assertRaises(SimulatedCrash):
                publish_candidate(document, candidate_uuid, "Редактор", failure_injector=injector)
            self.assertEqual(self.original, self.active_path.read_bytes())
            outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 180)
            self.assertEqual(1, outcome["recovered"])
            self.assertEqual(self.original, self.active_path.read_bytes())

    def test_stale_validation_is_interrupted_without_automatic_retry(self):
        candidate_uuid = self._upload_and_validate()
        with self.app.app_context():
            update_candidate(candidate_uuid, {"state": "validating"})
            outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 16 * 60)
            self.assertEqual(1, outcome["validation_interrupted"])
            self.assertEqual("validation_interrupted", get_candidate(candidate_uuid)["state"])

    def test_publish_failure_injection_points_have_deterministic_recovery(self):
        points = {
            "before_history": None,
            "after_prepared_history": "recovered",
            "after_publishing_metadata": "recovered",
            "before_replace": "recovered",
            "after_replace": "published",
            "after_active_verification": "published",
            "during_sibling_invalidation": "terminal_pending",
            "before_audit": "terminal_pending",
            "after_audit": "terminal_done",
        }

        class SimulatedCrash(BaseException):
            pass

        for point, expected in points.items():
            with self.subTest(point=point), tempfile.TemporaryDirectory() as temporary:
                base = Path(temporary); root = base / "doc_templates"; runtime = base / "runtime"
                active_path = root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx"
                old_bytes = make_template(active_path, heading="Active")
                new_bytes = make_template(base / "new.docx", heading="Candidate")
                app = build_app(root, runtime)
                with app.app_context():
                    document = next(iter(build_document_whitelist(root).values()))
                    candidate = write_uploaded_candidate(io.BytesIO(new_bytes), document_id=document.document_id, source_filename="new.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Failure Tester", comment="Failure injection coverage")
                    validate_staged_candidate(document.document_id, candidate["candidate_uuid"], document.path)
                    sibling_uuid = None
                    if point == "during_sibling_invalidation":
                        sibling = write_uploaded_candidate(io.BytesIO(new_bytes), document_id=document.document_id, source_filename="sibling.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Failure Tester", comment="Sibling candidate")
                        validate_staged_candidate(document.document_id, sibling["candidate_uuid"], document.path)
                        sibling_uuid = sibling["candidate_uuid"]

                    def injector(current):
                        if current == point:
                            raise SimulatedCrash()

                    with self.assertRaises(SimulatedCrash):
                        publish_candidate(document, candidate["candidate_uuid"], "Failure Tester", failure_injector=injector)
                    before_recovery = active_path.read_bytes()
                    outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, root), now=time.time() + 180)
                    self.assertEqual(before_recovery, active_path.read_bytes(), "Recovery must never repeat replace")
                    if expected is None:
                        self.assertEqual(old_bytes, active_path.read_bytes())
                        self.assertEqual("valid", get_candidate(candidate["candidate_uuid"])["state"])
                    elif expected in {"terminal_pending", "terminal_done"}:
                        self.assertEqual("published", get_candidate(candidate["candidate_uuid"])["state"])
                        self.assertEqual(1 if expected == "terminal_pending" else 0, outcome["audit_completed"])
                    else:
                        self.assertEqual(1, outcome[expected])
                        expected_state = "valid" if point == "after_prepared_history" else expected
                        self.assertEqual(expected_state, get_candidate(candidate["candidate_uuid"])["state"])
                    if sibling_uuid and expected in {"published", "terminal_pending", "terminal_done"}:
                        self.assertEqual("conflict", get_candidate(sibling_uuid)["state"])

    def test_rollback_failure_injection_points_have_deterministic_recovery(self):
        points = {
            "before_history": None,
            "after_prepared_history": "recovered",
            "after_publishing_metadata": "recovered",
            "before_replace": "recovered",
            "after_replace": "published",
            "after_active_verification": "published",
            "before_audit": "terminal_pending",
            "after_audit": "terminal_done",
        }

        class SimulatedCrash(BaseException):
            pass

        for point, expected in points.items():
            with self.subTest(point=point), tempfile.TemporaryDirectory() as temporary:
                base = Path(temporary); root = base / "doc_templates"; runtime = base / "runtime"
                active_path = root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx"
                original = make_template(active_path, heading="Original")
                replacement = make_template(base / "replacement.docx", heading="Replacement")
                app = build_app(root, runtime)
                with app.app_context():
                    document = next(iter(build_document_whitelist(root).values()))
                    candidate = write_uploaded_candidate(io.BytesIO(replacement), document_id=document.document_id, source_filename="replacement.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Rollback Tester", comment="Prepare rollback history")
                    validate_staged_candidate(document.document_id, candidate["candidate_uuid"], document.path)
                    publish_candidate(document, candidate["candidate_uuid"], "Rollback Tester")
                    target = list_history(document.document_id)[0]
                    active = resolve_document(document.document_id, root)

                    def injector(current):
                        if current == point:
                            raise SimulatedCrash()

                    with self.assertRaises(SimulatedCrash):
                        rollback_version(active, target["version_uuid"], actor="Rollback Tester", reason="Failure injection rollback", expected_active_sha=active.sha256, failure_injector=injector)
                    before_recovery = active_path.read_bytes()
                    outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, root), now=time.time() + 180)
                    self.assertEqual(before_recovery, active_path.read_bytes(), "Rollback recovery must never repeat replace")
                    if expected is None:
                        self.assertEqual(replacement, active_path.read_bytes())
                    elif expected in {"terminal_pending", "terminal_done"}:
                        self.assertEqual(1 if expected == "terminal_pending" else 0, outcome["audit_completed"])
                        self.assertEqual(original, active_path.read_bytes())
                    else:
                        self.assertEqual(1, outcome[expected])
                        self.assertEqual(original if expected == "published" else replacement, active_path.read_bytes())

    def test_upload_actual_byte_limit_removes_partial_candidate(self):
        response = self._upload(b"x" * (10 * 1024 * 1024 + 1))
        try:
            self.assertEqual(413, response.status_code)
        finally:
            response.request.environ["wsgi.input"].close()
            response.close()
        candidates = Path(self.temp.name) / "runtime/cache/document_template_center/candidates"
        self.assertFalse(candidates.exists() and any(candidates.iterdir()))

    def test_upload_request_content_length_is_rejected_early(self):
        response = self.client.open(
            f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates",
            method="POST",
            headers={"X-CSRF-Token": self.csrf, "Content-Type": "multipart/form-data; boundary=x"},
            environ_overrides={"CONTENT_LENGTH": str(10 * 1024 * 1024 + 256 * 1024 + 1)},
            data=b"",
        )
        self.assertEqual(413, response.status_code)

    def test_csrf_and_unknown_candidate_are_rejected(self):
        response = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates", data={"comment": "Комментарий", "file": (io.BytesIO(self.candidate), "x.docx")}, content_type="multipart/form-data")
        self.assertEqual(403, response.status_code)
        missing = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/00000000-0000-4000-8000-000000000000")
        self.assertEqual(404, missing.status_code)
        malformed = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/not-a-uuid")
        self.assertEqual(400, malformed.status_code)

    def test_synthetic_generation_never_resolves_real_jira(self):
        with mock.patch("services.docx_service.get_jira_domain_and_token", side_effect=AssertionError("network configuration must not be used")):
            response = self._upload()
        self.assertEqual(303, response.status_code)
        self.assertIn("replacement=success", response.headers["Location"])

    def test_security_rejects_traversal_zip_and_dtd(self):
        unsafe = Path(self.temp.name) / "unsafe.docx"
        with zipfile.ZipFile(unsafe, "w") as archive:
            archive.writestr("../escape.xml", "x")
            archive.writestr("[Content_Types].xml", "<!DOCTYPE x [<!ENTITY y SYSTEM 'file:///x'>]><x>&y;</x>")
            archive.writestr("_rels/.rels", "<Relationships/>")
            archive.writestr("word/document.xml", "<document/>")
        errors, _ = inspect_docx(unsafe)
        codes = {item.code for item in errors}
        self.assertIn("zip_unsafe_path", codes); self.assertNotIn("xml_unsafe", codes)


if __name__ == "__main__":
    unittest.main()
