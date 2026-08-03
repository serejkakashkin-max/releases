from __future__ import annotations

import importlib
import logging
import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def prepare_config_import() -> None:
    """Import config without truncating the application's ignored runtime logs."""
    if "config" in sys.modules:
        return
    root_logger = logging.getLogger()
    guard = logging.NullHandler()
    root_logger.addHandler(guard)
    original_file_handler = logging.FileHandler
    logging.FileHandler = lambda *args, **kwargs: logging.NullHandler()
    try:
        importlib.import_module("config")
    finally:
        logging.FileHandler = original_file_handler
        root_logger.removeHandler(guard)


def create_docx(path: Path, text: str = "Тестовый документ") -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(text, level=1)
    document.add_paragraph("RELEASE_VERSION")
    document.save(path)
    return path.read_bytes()
