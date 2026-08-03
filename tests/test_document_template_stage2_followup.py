from __future__ import annotations

import io
import json
import os
import stat
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from tests._support import prepare_config_import

prepare_config_import()

from services.docx_service import add_hyperlink
from services.document_template_audit_service import append_audit_event as real_append_audit_event
from services.document_template_candidate_service import candidate_preview_payload, validate_staged_candidate
from services.document_template_generation_service import (
    SYNTHETIC_ISSUE,
    SYNTHETIC_JIRA_BASE,
    generate_synthetic_document,
)
from services.document_template_publish_service import (
    DocumentMutationBlocked,
    publish_candidate,
    recover_stale_operations,
    rollback_version,
)
from services.document_template_read_service import build_document_whitelist, resolve_document
from services.document_template_storage_service import (
    HistoryPreviewDenied,
    committed_history_payload,
    create_history_version,
    get_candidate,
    history_file,
    list_history,
    sha256_bytes,
    sha256_file,
    update_candidate,
    utc_now,
    write_uploaded_candidate,
)
from services.document_template_validation_service import build_contract, validate_candidate
from services.release_template_catalog_service import clear_template_catalog_cache
from tests.test_document_template_stage2 import TOKEN, build_app, make_template


def make_no_jira_template(path: Path, *, heading: str = "Шаблон без Jira", unexpected_link: bool = False) -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph("Версия RELEASE_VERSION")
    document.add_paragraph("Дата DATE")
    if unexpected_link:
        paragraph = document.add_paragraph("Ссылка: ")
        add_hyperlink(paragraph, f"{SYNTHETIC_JIRA_BASE}/browse/{SYNTHETIC_ISSUE['key']}", SYNTHETIC_ISSUE["key"])
    document.save(path)
    return path.read_bytes()


def make_jira_structure(path: Path, headers: list[str], *, heading: str = "Jira contract") -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph("Версия RELEASE_VERSION")
    document.add_paragraph("Дата DATE")
    table = document.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    document.save(path)
    return path.read_bytes()


class JiraContractActiveBasedTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)

    def tearDown(self):
        self.temp.cleanup()

    def _validate(self, active_headers: list[str] | None, candidate_headers: list[str] | None):
        active = self.root / "active.docx"
        candidate = self.root / "candidate.docx"
        if active_headers is None:
            make_no_jira_template(active)
        else:
            make_jira_structure(active, active_headers)
        if candidate_headers is None:
            make_no_jira_template(candidate)
        else:
            make_jira_structure(candidate, candidate_headers)
        return validate_candidate(active, candidate), candidate

    def test_four_known_legacy_patterns_accept_identical_candidate_with_warning(self):
        patterns = [
            ["№", "ЗНИ/JIRA ID", "Суть доработки", "Issue \nType"],
            ["№", "ЗНИ/JIRA ID", "Суть доработки", "Issue \nType"],
            ["№", "ЗНИ/JIRA ID", "Суть доработки", "Issue \nType"],
            ["№", "ЗНИ/JIRA ID", "Суть доработки", "Issue \nType"],
        ]
        for index, headers in enumerate(patterns):
            with self.subTest(pattern=index):
                result, _ = self._validate(headers, headers)
                self.assertTrue(result["ok"], result)
                self.assertNotIn("jira_table_columns", {item["code"] for item in result["errors"]})
                self.assertEqual(["legacy_jira_table"], [item["code"] for item in result["warnings"]])

    def test_legacy_candidate_can_add_safe_type_column(self):
        legacy = ["№", "ЗНИ/JIRA ID", "Issue", "Команда"]
        improved = ["№", "ЗНИ/JIRA ID", "Issue", "Issue Type", "Команда"]
        result, candidate = self._validate(legacy, improved)
        self.assertTrue(result["ok"], result)
        self.assertNotIn("legacy_jira_table", {item["code"] for item in result["warnings"]})
        generated, errors = generate_synthetic_document(candidate, self.root / "improved-test.docx")
        self.assertIsNotNone(generated, errors)

    def test_legacy_candidate_cannot_remove_or_break_jira_id(self):
        legacy = ["№", "ЗНИ/JIRA ID", "Суть доработки", "Issue \nType"]
        removed, _ = self._validate(legacy, ["№", "Система", "Суть доработки", "Issue Type"])
        unsupported, _ = self._validate(legacy, ["№", "Issue", "ЗНИ/JIRA ID", "Issue Type"])
        self.assertFalse(removed["ok"])
        self.assertIn("jira_table_count", {item["code"] for item in removed["errors"]})
        self.assertFalse(unsupported["ok"])
        self.assertIn("jira_table_columns", {item["code"] for item in unsupported["errors"]})

    def test_modern_candidate_cannot_remove_type_column(self):
        modern = ["№", "ЗНИ/JIRA ID", "Issue", "Issue Type"]
        result, _ = self._validate(modern, ["№", "ЗНИ/JIRA ID", "Issue", "Команда"])
        self.assertFalse(result["ok"])
        self.assertIn("jira_table_columns", {item["code"] for item in result["errors"]})

    def test_no_jira_contract_and_new_table_modern_rule(self):
        no_jira, _ = self._validate(None, None)
        invalid_new, _ = self._validate(None, ["№", "ЗНИ/JIRA ID", "Issue"])
        valid_new, candidate = self._validate(None, ["№", "ЗНИ/JIRA ID", "Issue", "Issue Type"])
        self.assertTrue(no_jira["ok"], no_jira)
        self.assertFalse(invalid_new["ok"])
        self.assertIn("jira_table_columns", {item["code"] for item in invalid_new["errors"]})
        self.assertTrue(valid_new["ok"], valid_new)
        generated, errors = generate_synthetic_document(candidate, self.root / "new-jira-test.docx")
        self.assertIsNotNone(generated, errors)


class Stage2FollowupTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.base = Path(self.temp.name)
        self.root = self.base / "doc_templates"
        self.active_path = self.root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx"
        self.active_bytes = make_no_jira_template(self.active_path, heading="Активный")
        self.candidate_bytes = make_no_jira_template(self.base / "candidate.docx", heading="Кандидат")
        clear_template_catalog_cache()
        self.app = build_app(self.root, self.base / "runtime")
        self.client = self.app.test_client()
        self.client.post("/admin/document-templates/session/login", data={"display_name": "Редактор", "token": TOKEN})
        with self.client.session_transaction() as session:
            self.csrf = session["document_template_editor_csrf_nonce"]
        with self.app.app_context():
            self.document_id = next(iter(build_document_whitelist(self.root)))
            self.document = resolve_document(self.document_id, self.root)

    def tearDown(self):
        clear_template_catalog_cache()
        self.temp.cleanup()

    def _validated_candidate(self, payload: bytes | None = None):
        candidate = write_uploaded_candidate(
            io.BytesIO(payload or self.candidate_bytes),
            document_id=self.document_id,
            source_filename="Кандидат.docx",
            active_filename=self.document.filename,
            active_sha=self.document.sha256,
            uploaded_by="Редактор",
            comment="Проверка повторного аудита",
        )
        validated = validate_staged_candidate(self.document_id, candidate["candidate_uuid"], self.document.path)
        self.assertEqual("valid", validated["state"], validated.get("validation"))
        return validated

    def _history_version(self, payload: bytes, *, state: str = "committed", sha: str | None = None):
        return create_history_version(self.document_id, payload, {
            "created_at": utc_now(),
            "updated_at": utc_now(),
            "state": state,
            "source_filename": self.document.filename,
            "sha256": sha or sha256_bytes(payload),
            "actor": "Редактор",
            "action": "seed",
            "comment": "Synthetic history",
            "contract": build_contract(self.active_path),
        })

    def _audit_events(self):
        path = self.base / "runtime/data/document_template_center/audit/events.jsonl"
        return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines()] if path.is_file() else []

    def test_no_jira_template_validates_generates_and_publishes(self):
        with self.app.app_context():
            candidate = self._validated_candidate()
            self.assertEqual(0, candidate["validation"]["contract"]["jira_table_count"])
            payload, metadata = candidate_preview_payload(self.document_id, candidate["candidate_uuid"], test_document=True)
            self.assertEqual(metadata["test_sha"], sha256_bytes(payload))
            published = publish_candidate(self.document, candidate["candidate_uuid"], "Редактор")
            self.assertEqual("published", published["state"])
            self.assertEqual(candidate["candidate_sha"], sha256_file(self.active_path))

    def test_jira_contract_requires_exact_row_and_hyperlink(self):
        source = self.base / "jira.docx"
        make_template(source)
        output, errors = generate_synthetic_document(source, self.base / "jira-test.docx")
        self.assertIsNotNone(output, errors)
        generated = Document(output)
        matching_rows = []
        for table in generated.tables:
            for row in table.rows[1:]:
                values = [cell.text.strip() for cell in row.cells]
                if len(values) >= 4 and values[1] == SYNTHETIC_ISSUE["key"]:
                    matching_rows.append(values)
        self.assertEqual([["1", SYNTHETIC_ISSUE["key"], SYNTHETIC_ISSUE["summary"], SYNTHETIC_ISSUE["type"]]], matching_rows)
        links = [rel.target_ref for rel in generated.part.rels.values() if rel.reltype == RT.HYPERLINK and rel.is_external]
        self.assertIn(f"{SYNTHETIC_JIRA_BASE}/browse/{SYNTHETIC_ISSUE['key']}", links)

    def test_no_jira_template_rejects_unexpected_synthetic_hyperlink(self):
        source = self.base / "unexpected-link.docx"
        make_no_jira_template(source, unexpected_link=True)
        output, errors = generate_synthetic_document(source, self.base / "unexpected-test.docx")
        self.assertIsNone(output)
        self.assertEqual(["generation_jira_unexpected"], [item["code"] for item in errors])

    def test_publish_regeneration_refreshes_test_metadata_before_replace(self):
        with self.app.app_context():
            candidate = self._validated_candidate()
            candidate_uuid = candidate["candidate_uuid"]
            update_candidate(candidate_uuid, {"test_sha": "0" * 64, "test_size": 1}, document_id=self.document_id)

            def stop_after_generation(point):
                if point == "before_history":
                    raise RuntimeError("stop after regeneration")

            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace") as replace:
                with self.assertRaisesRegex(RuntimeError, "stop after regeneration"):
                    publish_candidate(self.document, candidate_uuid, "Редактор", failure_injector=stop_after_generation)
                replace.assert_not_called()
            refreshed = get_candidate(candidate_uuid, document_id=self.document_id)
            test_payload, _ = candidate_preview_payload(self.document_id, candidate_uuid, test_document=True)
            self.assertEqual("valid", refreshed["state"])
            self.assertTrue(refreshed["validation"]["checks"]["generation"])
            self.assertEqual(sha256_bytes(test_payload), refreshed["test_sha"])
            self.assertEqual(len(test_payload), refreshed["test_size"])

    def test_publish_recovered_and_restore_failed_have_terminal_audit(self):
        with self.app.app_context():
            recovered = self._validated_candidate()
            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace", side_effect=[OSError("replace"), None]):
                with self.assertRaises(OSError):
                    publish_candidate(self.document, recovered["candidate_uuid"], "Редактор")
            publish_events = [item for item in self._audit_events() if item.get("action", "").startswith("publish")]
            self.assertEqual(["attempt", "recovered"], [item["result"] for item in publish_events])

        # Use a fresh fixture because a recovered candidate is terminal.
        with self.app.app_context():
            failed = self._validated_candidate()

            def audit_with_terminal_failure(**values):
                if values.get("result") == "publish_failed":
                    raise OSError("audit unavailable")
                return real_append_audit_event(**values)

            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace", side_effect=[OSError("replace"), OSError("restore")]), mock.patch("services.document_template_publish_service.append_audit_event", side_effect=audit_with_terminal_failure):
                with self.assertRaises(DocumentMutationBlocked):
                    publish_candidate(self.document, failed["candidate_uuid"], "Редактор")
            metadata = get_candidate(failed["candidate_uuid"], document_id=self.document_id)
            self.assertEqual("publish_failed", metadata["state"])
            self.assertEqual("publish_failed", metadata["audit_pending"]["result"])
            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace") as replace:
                outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root))
                replace.assert_not_called()
            self.assertEqual(1, outcome["audit_completed"])
            self.assertIsNone(get_candidate(failed["candidate_uuid"], document_id=self.document_id).get("audit_pending"))
            failed_events = [item for item in self._audit_events() if item.get("candidate_uuid") == failed["candidate_uuid"] and item.get("action", "").startswith("publish")]
            self.assertEqual(["attempt", "publish_failed"], [item["result"] for item in failed_events])

    def test_rollback_recovered_and_restore_failed_have_terminal_audit(self):
        target_path = self.base / "target.docx"
        target_bytes = make_no_jira_template(target_path, heading="Исторический")
        with self.app.app_context():
            target = create_history_version(self.document_id, target_bytes, {
                "created_at": utc_now(), "updated_at": utc_now(), "state": "committed",
                "source_filename": self.document.filename, "sha256": sha256_bytes(target_bytes),
                "actor": "Редактор", "action": "seed", "comment": "Target",
                "contract": build_contract(target_path),
            })
            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace", side_effect=[OSError("replace"), None]):
                with self.assertRaises(OSError):
                    rollback_version(self.document, target["version_uuid"], actor="Редактор", reason="Проверка recovered", expected_active_sha=self.document.sha256)
            rollback_events = [item for item in self._audit_events() if item.get("action", "").startswith("rollback")]
            self.assertEqual(["attempt", "recovered"], [item["result"] for item in rollback_events])

        # A fresh app/runtime keeps the second rollback independent from recovery evidence.
        other_runtime = self.base / "runtime-rollback-failed"
        other_app = build_app(self.root, other_runtime)
        with other_app.app_context():
            document = resolve_document(self.document_id, self.root)
            target = create_history_version(self.document_id, target_bytes, {
                "created_at": utc_now(), "updated_at": utc_now(), "state": "committed",
                "source_filename": document.filename, "sha256": sha256_bytes(target_bytes),
                "actor": "Редактор", "action": "seed", "comment": "Target",
                "contract": build_contract(target_path),
            })

            def audit_with_terminal_failure(**values):
                if values.get("result") == "publish_failed":
                    raise OSError("audit unavailable")
                return real_append_audit_event(**values)

            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace", side_effect=[OSError("replace"), OSError("restore")]), mock.patch("services.document_template_publish_service.append_audit_event", side_effect=audit_with_terminal_failure):
                with self.assertRaises(DocumentMutationBlocked):
                    rollback_version(document, target["version_uuid"], actor="Редактор", reason="Проверка failure", expected_active_sha=document.sha256)
            recovery_version = next(item for item in list_history(self.document_id) if item.get("action") == "rollback_previous")
            self.assertEqual("publish_failed", recovery_version["state"])
            self.assertEqual("publish_failed", recovery_version["audit_pending"]["result"])
            with mock.patch("services.document_template_publish_service._write_adjacent_and_replace") as replace:
                outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root))
                replace.assert_not_called()
            self.assertEqual(1, outcome["audit_completed"])
            rollback_audit = other_runtime / "data/document_template_center/audit/events.jsonl"
            failed_events = [json.loads(line) for line in rollback_audit.read_text(encoding="utf-8").splitlines() if json.loads(line).get("action", "").startswith("rollback")]
            self.assertEqual(["attempt", "publish_failed"], [item["result"] for item in failed_events])

    def test_history_preview_download_require_committed_safe_consistent_file(self):
        with self.app.app_context():
            committed = self._history_version(self.active_bytes)
            prepared = self._history_version(self.active_bytes, state="prepared")
            publishing = self._history_version(self.active_bytes, state="publishing")
            publish_failed = self._history_version(self.active_bytes, state="publish_failed")
            wrong_sha = self._history_version(self.active_bytes, sha="0" * 64)
            unsafe = self._history_version(b"not a docx")

        base = f"/admin/document-templates/documents/{self.document_id}/history"
        preview = self.client.get(f"{base}/{committed['version_uuid']}/preview")
        download = self.client.get(f"{base}/{committed['version_uuid']}/download")
        self.assertEqual(200, preview.status_code)
        self.assertEqual(self.active_bytes, preview.data)
        self.assertEqual(200, download.status_code)
        self.assertIn("attachment", download.headers["Content-Disposition"])
        for blocked in (prepared, publishing, publish_failed, wrong_sha, unsafe):
            self.assertEqual(409, self.client.get(f"{base}/{blocked['version_uuid']}/preview").status_code)
            self.assertEqual(409, self.client.get(f"{base}/{blocked['version_uuid']}/download").status_code)

        with self.app.app_context():
            target = history_file(self.document_id, committed["version_uuid"])
        real_lstat = os.lstat

        def symlink_lstat(path):
            value = real_lstat(path)
            if Path(path) == target:
                fields = list(value)
                fields[stat.ST_MODE] = stat.S_IFLNK | 0o777
                return os.stat_result(fields)
            return value

        with mock.patch("services.document_template_storage_service.os.lstat", side_effect=symlink_lstat):
            self.assertEqual(409, self.client.get(f"{base}/{committed['version_uuid']}/preview").status_code)
            with self.app.app_context():
                with self.assertRaises(HistoryPreviewDenied):
                    committed_history_payload(self.document_id, committed["version_uuid"])


if __name__ == "__main__":
    unittest.main()
