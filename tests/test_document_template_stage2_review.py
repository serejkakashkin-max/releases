from __future__ import annotations

import io
import json
import tempfile
import threading
import time
import unittest
import zipfile
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta, timezone
from pathlib import Path
from unittest import mock

from docx import Document

from tests._support import PROJECT_ROOT, prepare_config_import

prepare_config_import()

from services.document_template_candidate_service import validate_staged_candidate
from services.document_template_csrf_service import CSRF_COOKIE_NAME
from services.document_template_generation_service import generate_synthetic_document
from services.document_template_publish_service import (
    DocumentMutationBlocked,
    _blocked_by_recovery,
    publish_candidate,
    recover_stale_operations,
    rollback_version,
)
import services.document_template_publish_service as publish_service
import services.document_template_read_service as read_service
import services.document_template_validation_service as validation_service
import services.document_template_vendor_service as vendor_service
from services.document_template_read_service import (
    build_document_whitelist,
    clear_document_template_read_cache,
    resolve_document,
)
from services.document_template_runtime_service import atomic_write_json
from services.document_template_storage_service import (
    candidate_directory,
    candidate_file,
    cleanup_candidates,
    create_history_version,
    get_candidate,
    get_history_version,
    history_directory,
    list_history,
    prune_history,
    sha256_bytes,
    update_candidate,
    utc_now,
    write_uploaded_candidate,
)
from services.document_template_validation_service import (
    build_contract,
    compare_contract,
    inspect_docx,
    validate_candidate,
)
from services.document_template_vendor_service import (
    clear_document_template_vendor_cache,
    verify_vendor_assets,
)
from services.release_template_catalog_service import clear_template_catalog_cache
from tests.test_document_template_stage2 import build_app, make_template


def _set_candidate_metadata(candidate_uuid: str, **updates):
    path = candidate_directory(candidate_uuid) / "metadata.json"
    value = json.loads(path.read_text(encoding="utf-8"))
    value.update(updates)
    atomic_write_json(path, value)
    return value


def _make_invalid_contract(path: Path) -> bytes:
    make_template(path, heading="Invalid contract")
    document = Document(path)
    for paragraph in document.paragraphs:
        if "DATE" in paragraph.text:
            paragraph.text = paragraph.text.replace("DATE", "15.01.2030")
    document.save(path)
    return path.read_bytes()


def _rewrite_zip(path: Path, transform):
    temporary = path.with_suffix(".rewritten")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for info in source.infolist():
            payload = source.read(info.filename)
            target.writestr(info, transform(info.filename, payload))
    temporary.replace(path)


class ZipValidationReviewTests(unittest.TestCase):
    def test_zip_bomb_short_circuits_before_any_entry_read_or_python_docx(self):
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "bomb.docx"
            with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("_rels/.rels", "<Relationships/>")
                archive.writestr("word/document.xml", b"A" * (4 * 1024 * 1024))
            with mock.patch.object(zipfile.ZipFile, "open", side_effect=AssertionError("entry decompression is forbidden")) as opened, mock.patch.object(validation_service, "Document") as document:
                errors, _ = inspect_docx(path)
            self.assertIn("zip_compression_ratio", {item.code for item in errors})
            opened.assert_not_called()
            document.assert_not_called()


class Stage2ReviewWorkflowTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        base = Path(self.temp.name)
        self.root = base / "doc_templates"
        self.active_path = self.root / "CAT" / "Комплект PL (12345)" / "Шаблон.docx"
        self.original = make_template(self.active_path, heading="Активный")
        self.candidate = make_template(base / "candidate.docx", heading="Кандидат")
        clear_template_catalog_cache(); clear_document_template_read_cache(); clear_document_template_vendor_cache()
        self.app = build_app(self.root, base / "runtime")
        self.client = self.app.test_client()
        self.client.get("/dashboard/release-monitor/document-templates")
        self.csrf = self.client.get_cookie(CSRF_COOKIE_NAME, path="/dashboard/release-monitor/document-templates/").value
        self.document_id = next(iter(build_document_whitelist(self.root)))

    def tearDown(self):
        clear_template_catalog_cache(); clear_document_template_read_cache(); clear_document_template_vendor_cache(); self.temp.cleanup()

    def _upload(self, payload=None):
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            item = write_uploaded_candidate(
                io.BytesIO(payload or self.candidate),
                document_id=self.document_id,
                source_filename="candidate.docx",
                active_filename=document.filename,
                active_sha=document.sha256,
                uploaded_by="Review Tester",
                comment="Regression candidate",
            )
        return item["candidate_uuid"]

    def _validate(self, candidate_uuid):
        return self.client.post(
            f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/validate",
            data={"_csrf_token": self.csrf},
        )

    def _service_candidate(self, payload=None):
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            item = write_uploaded_candidate(
                io.BytesIO(payload or self.candidate), document_id=self.document_id,
                source_filename="candidate.docx", active_filename=document.filename,
                active_sha=document.sha256, uploaded_by="Review Tester", comment="Review regression",
            )
            validate_staged_candidate(self.document_id, item["candidate_uuid"], document.path)
            return document, item["candidate_uuid"]

    def test_candidate_and_test_preview_server_gates(self):
        candidate_uuid = self._upload()
        preview_url = f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/preview"
        test_url = f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/test-document/preview"
        self.assertEqual(409, self.client.get(preview_url).status_code)
        self.assertEqual(409, self.client.get(test_url).status_code)
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            validate_staged_candidate(self.document_id, candidate_uuid, document.path)
        self.assertEqual(200, self.client.get(preview_url).status_code)
        self.assertEqual(200, self.client.get(test_url).status_code)
        with self.app.app_context():
            candidate_file(candidate_uuid, "test.docx").write_bytes(b"stale")
        self.assertEqual(409, self.client.get(test_url).status_code)
        self.assertEqual(409, self.client.get(test_url.replace("/preview", "/download")).status_code)
        with self.app.app_context():
            candidate_file(candidate_uuid).write_bytes(b"stale-candidate")
        self.assertEqual(409, self.client.get(preview_url).status_code)

    def test_uploaded_zip_bomb_and_invalid_security_cannot_be_previewed(self):
        payload = io.BytesIO()
        with zipfile.ZipFile(payload, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "<Types/>")
            archive.writestr("_rels/.rels", "<Relationships/>")
            archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))
        candidate_uuid = self._upload(payload.getvalue())
        url = f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/preview"
        self.assertEqual(409, self.client.get(url).status_code)
        self.assertEqual(422, self._validate(candidate_uuid).status_code)
        with self.app.app_context():
            self.assertEqual("invalid_security", get_candidate(candidate_uuid)["state"])
        self.assertEqual(409, self.client.get(url).status_code)

    def test_invalid_contract_allows_only_basic_candidate_preview_and_removes_old_test(self):
        invalid = _make_invalid_contract(Path(self.temp.name) / "invalid.docx")
        candidate_uuid = self._upload(invalid)
        with self.app.app_context():
            old_test = candidate_directory(candidate_uuid) / "test.docx"
            old_test.write_bytes(self.candidate)
        self.assertEqual(422, self._validate(candidate_uuid).status_code)
        preview = f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/preview"
        test_preview = f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/test-document/preview"
        self.assertEqual(200, self.client.get(preview).status_code)
        self.assertEqual(409, self.client.get(test_preview).status_code)
        with self.app.app_context():
            metadata = get_candidate(candidate_uuid)
            self.assertEqual("invalid_contract", metadata["state"])
            self.assertFalse((candidate_directory(candidate_uuid) / "test.docx").exists())

    def test_cleanup_preserves_recovery_evidence_and_removes_terminal_bytes(self):
        with self.app.app_context():
            document, failed_uuid = self._service_candidate()
            old = (datetime.now(timezone.utc) - timedelta(hours=25)).isoformat()
            update_candidate(failed_uuid, {"state": "publish_failed", "recovery_blocking": True})
            _set_candidate_metadata(failed_uuid, created_at=old, updated_at=old)
            cleanup_candidates(now=time.time())
            self.assertTrue(candidate_file(failed_uuid).exists())
            self.assertTrue(_blocked_by_recovery(self.document_id))

            publishing = write_uploaded_candidate(io.BytesIO(self.candidate), document_id=self.document_id, source_filename="publishing.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Tester", comment="publishing evidence")
            update_candidate(publishing["candidate_uuid"], {"state": "publishing", "operation_uuid": "operation"})
            _set_candidate_metadata(publishing["candidate_uuid"], created_at=old, updated_at=old)
            cleanup_candidates(now=time.time())
            self.assertTrue(candidate_file(publishing["candidate_uuid"]).exists())

            published = write_uploaded_candidate(io.BytesIO(self.candidate), document_id=self.document_id, source_filename="published.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Tester", comment="published bytes")
            (candidate_directory(published["candidate_uuid"]) / "test.docx").write_bytes(self.candidate)
            update_candidate(published["candidate_uuid"], {"state": "published"})
            cleanup_candidates(now=time.time())
            self.assertFalse((candidate_directory(published["candidate_uuid"]) / "candidate.docx").exists())
            self.assertFalse((candidate_directory(published["candidate_uuid"]) / "test.docx").exists())

            expired = write_uploaded_candidate(io.BytesIO(self.candidate), document_id=self.document_id, source_filename="expired.docx", active_filename=document.filename, active_sha=document.sha256, uploaded_by="Tester", comment="expired bytes")
            _set_candidate_metadata(expired["candidate_uuid"], state="expired", created_at=old, updated_at=old)
            cleanup_candidates(now=time.time())
            self.assertFalse(candidate_directory(expired["candidate_uuid"]).exists())

    def test_history_pruning_keeps_recovery_versions_and_thirty_committed(self):
        with self.app.app_context():
            committed = []
            for index in range(32):
                payload = f"committed-{index}".encode()
                committed.append(create_history_version(self.document_id, payload, {"created_at": f"2030-01-{index % 28 + 1:02d}T00:00:00+00:00", "updated_at": utc_now(), "state": "committed", "sha256": sha256_bytes(payload)}))
            prepared = create_history_version(self.document_id, b"prepared", {"created_at": utc_now(), "updated_at": utc_now(), "state": "prepared", "sha256": sha256_bytes(b"prepared")})
            failed = create_history_version(self.document_id, b"failed", {"created_at": utc_now(), "updated_at": utc_now(), "state": "publish_failed", "sha256": sha256_bytes(b"failed")})
            prune_history(self.document_id, keep=30)
            versions = list_history(self.document_id)
            self.assertEqual(30, sum(item["state"] == "committed" for item in versions))
            self.assertTrue(history_directory(self.document_id, prepared["version_uuid"]).exists())
            self.assertTrue(history_directory(self.document_id, failed["version_uuid"]).exists())

    def test_attempt_audit_failure_prevents_publish_and_rollback_replace(self):
        with self.app.app_context():
            document, candidate_uuid = self._service_candidate()
            with mock.patch.object(publish_service, "append_audit_event", side_effect=OSError("audit unavailable")), mock.patch.object(publish_service, "_write_adjacent_and_replace") as replace:
                with self.assertRaises(DocumentMutationBlocked):
                    publish_candidate(document, candidate_uuid, "Review Tester")
                replace.assert_not_called()
            self.assertEqual(self.original, self.active_path.read_bytes())

            publish_candidate(document, candidate_uuid, "Review Tester")
            active = resolve_document(self.document_id, self.root)
            version = list_history(self.document_id)[0]
            with mock.patch.object(publish_service, "append_audit_event", side_effect=OSError("audit unavailable")), mock.patch.object(publish_service, "_write_adjacent_and_replace") as replace:
                with self.assertRaises(DocumentMutationBlocked):
                    rollback_version(active, version["version_uuid"], actor="Review Tester", reason="Audit ordering", expected_active_sha=active.sha256)
                replace.assert_not_called()
            self.assertEqual(self.candidate, self.active_path.read_bytes())

    def test_terminal_audit_failure_is_pending_and_recovery_never_replaces(self):
        with self.app.app_context():
            document, candidate_uuid = self._service_candidate()
            real_append = publish_service.append_audit_event
            calls = {"count": 0}

            def fail_terminal(**values):
                calls["count"] += 1
                if values.get("action") == "publish_attempt":
                    return real_append(**values)
                raise OSError("terminal audit unavailable")

            with mock.patch.object(publish_service, "append_audit_event", side_effect=fail_terminal):
                result = publish_candidate(document, candidate_uuid, "Publisher")
            self.assertEqual("published", result["state"])
            self.assertEqual("Publisher", result["published_by"])
            self.assertIsInstance(get_candidate(candidate_uuid)["audit_pending"], dict)
            with mock.patch.object(publish_service, "_write_adjacent_and_replace") as replace:
                outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 180)
                replace.assert_not_called()
            self.assertEqual(1, outcome["audit_completed"])
            self.assertIsNone(get_candidate(candidate_uuid).get("audit_pending"))

    def test_corrupt_history_forces_publish_failed_without_replace(self):
        class Crash(BaseException):
            pass

        with self.app.app_context():
            document, candidate_uuid = self._service_candidate()

            def injector(point):
                if point == "after_publishing_metadata":
                    raise Crash()

            with self.assertRaises(Crash):
                publish_candidate(document, candidate_uuid, "Review Tester", failure_injector=injector)
            metadata = get_candidate(candidate_uuid)
            (history_directory(self.document_id, metadata["history_version_uuid"]) / "document.docx").write_bytes(b"corrupt")
            with mock.patch.object(publish_service, "_write_adjacent_and_replace") as replace:
                outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 180)
                replace.assert_not_called()
            self.assertEqual(1, outcome["publish_failed"])
            self.assertEqual("publish_failed", get_candidate(candidate_uuid)["state"])
            self.assertEqual(self.original, self.active_path.read_bytes())

    def test_corrupt_rollback_target_forces_blocked_recovery_without_replace(self):
        class Crash(BaseException):
            pass

        with self.app.app_context():
            document, candidate_uuid = self._service_candidate()
            publish_candidate(document, candidate_uuid, "Review Tester")
            active = resolve_document(self.document_id, self.root)
            target = next(item for item in list_history(self.document_id) if item.get("action") == "publish_previous")

            def injector(point):
                if point == "after_publishing_metadata":
                    raise Crash()

            with self.assertRaises(Crash):
                rollback_version(active, target["version_uuid"], actor="Review Tester", reason="Corrupt rollback recovery", expected_active_sha=active.sha256, failure_injector=injector)
            (history_directory(self.document_id, target["version_uuid"]) / "document.docx").write_bytes(b"corrupt-target")
            before = self.active_path.read_bytes()
            with mock.patch.object(publish_service, "_write_adjacent_and_replace") as replace:
                outcome = recover_stale_operations(document_resolver=lambda value: resolve_document(value, self.root), now=time.time() + 180)
                replace.assert_not_called()
            self.assertEqual(1, outcome["publish_failed"])
            self.assertEqual(before, self.active_path.read_bytes())
            recovery = next(item for item in list_history(self.document_id) if item.get("action") == "rollback_previous")
            self.assertEqual("publish_failed", recovery["state"])

    def test_expected_sha_requires_lowercase_hex(self):
        candidate_uuid = self._upload()
        with self.app.app_context():
            document = resolve_document(self.document_id, self.root)
            validate_staged_candidate(self.document_id, candidate_uuid, document.path)
        published = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/publish", data={"_csrf_token": self.csrf})
        self.assertEqual(303, published.status_code)
        with self.app.app_context():
            version = list_history(self.document_id)[0]
        invalid = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/history/{version['version_uuid']}/rollback", data={"_csrf_token": self.csrf, "reason": "Invalid uppercase SHA", "expected_active_sha": "A" * 64})
        self.assertEqual(400, invalid.status_code)

    def test_upload_validation_and_cancel_are_audited_without_sensitive_values(self):
        invalid_path = Path(self.temp.name) / "invalid-contract.docx"
        invalid_document = Document()
        invalid_document.add_paragraph("No required placeholders")
        invalid_document.save(invalid_path)
        response = self.client.post(
            f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates",
            data={
                "_csrf_token": self.csrf,
                "comment": "Regression candidate",
                "file": (io.BytesIO(invalid_path.read_bytes()), "candidate.docx"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(303, response.status_code)
        candidate_uuid = response.headers["Location"].rstrip("/").split("/")[-1]
        with self.app.app_context():
            self.assertTrue(get_candidate(candidate_uuid)["state"].startswith("invalid_"))
        cancelled = self.client.post(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/candidates/{candidate_uuid}/cancel", data={"_csrf_token": self.csrf})
        self.assertEqual(303, cancelled.status_code)
        audit_path = Path(self.temp.name) / "runtime/data/document_template_center/audit/events.jsonl"
        events = [json.loads(line) for line in audit_path.read_text(encoding="utf-8").splitlines()]
        actions = {event["action"] for event in events}
        self.assertTrue({"upload", "validation_result", "cancel"}.issubset(actions))
        audit_text = audit_path.read_text(encoding="utf-8")
        self.assertNotIn(self.csrf, audit_text)
        self.assertNotIn(str(self.root), audit_text)


class ContractAndGenerationReviewTests(unittest.TestCase):
    def test_split_runs_table_placeholders_instruction_and_unsupported_parts(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            active = root / "active.docx"
            candidate = root / "candidate.docx"
            for path in (active, candidate):
                document = Document()
                paragraph = document.add_paragraph()
                paragraph.add_run("RELEASE_")
                paragraph.add_run("VERSION")
                table = document.add_table(rows=2, cols=4)
                table.rows[0].cells[0].text = "№"; table.rows[0].cells[1].text = "ЗНИ/JIRA ID"; table.rows[0].cells[2].text = "Issue"; table.rows[0].cells[3].text = "Issue Type"
                cell_paragraph = table.rows[1].cells[0].paragraphs[0]
                cell_paragraph.add_run("DA"); cell_paragraph.add_run("TE")
                document.add_paragraph("ИНСТРУКЦИЯ")
                document.save(path)
            errors, _ = compare_contract(active, candidate)
            self.assertEqual([], errors)
            output, generation_errors = generate_synthetic_document(candidate, root / "test.docx")
            self.assertFalse(generation_errors); self.assertIsNotNone(output)
            reopened = Document(output)
            self.assertFalse(any(validation_service.PLACEHOLDER_RE.search(p.text) for table in reopened.tables for row in table.rows for cell in row.cells for p in cell.paragraphs))

            def add_unsupported(name, payload):
                return payload

            temporary_zip = candidate.with_suffix(".tmp")
            with zipfile.ZipFile(candidate) as source, zipfile.ZipFile(temporary_zip, "w", zipfile.ZIP_DEFLATED) as target:
                for info in source.infolist():
                    target.writestr(info, source.read(info.filename))
                target.writestr("word/footnotes.xml", '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:footnote w:id="1"><w:p><w:r><w:t>CHECKER</w:t></w:r></w:p></w:footnote></w:footnotes>')
                target.writestr("word/endnotes.xml", '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:endnote w:id="1"><w:p><w:r><w:t>OPLOT</w:t></w:r></w:p></w:endnote></w:endnotes>')
            temporary_zip.replace(candidate)
            contract = build_contract(candidate)
            self.assertIn("CHECKER", contract["areas"]["unsupported"])
            self.assertIn("OPLOT", contract["areas"]["unsupported"])

    def test_whitespace_damage_and_grandfathered_external_image_warning(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            active = root / "active.docx"; candidate = root / "candidate.docx"
            make_template(active); make_template(candidate)
            document = Document(candidate)
            document.add_paragraph("RELEASE_ VERSION")
            document.save(candidate)
            errors, contract = compare_contract(active, candidate)
            self.assertIn("placeholder_whitespace_damaged", {item.code for item in errors})

            candidate.write_bytes(active.read_bytes())
            relation = '<Relationship Id="rIdExternalImage" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="https://example.invalid/image.png" TargetMode="External"/>'
            for path in (active, candidate):
                _rewrite_zip(path, lambda name, payload: payload.replace(b"</Relationships>", relation.encode() + b"</Relationships>") if name == "word/_rels/document.xml.rels" else payload)
            result = validate_candidate(active, candidate)
            self.assertTrue(result["ok"])
            self.assertIn("grandfathered_external_image", {item["code"] for item in result["warnings"]})


class TargetReadAndCacheReviewTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        base = Path(self.temp.name)
        self.root = base / "doc_templates"
        for index in range(4):
            make_template(self.root / "CAT" / f"Комплект {index} (1234{index})" / f"Шаблон {index}.docx", heading=f"Template {index}")
        clear_document_template_read_cache(); clear_document_template_vendor_cache()
        self.app = build_app(self.root, base / "runtime")
        self.client = self.app.test_client()
        self.document_id = next(iter(build_document_whitelist(self.root)))
        clear_document_template_read_cache()

    def tearDown(self):
        clear_document_template_read_cache(); clear_document_template_vendor_cache(); self.temp.cleanup()

    def test_preview_does_not_hash_other_docx_and_detects_target_replacement_and_deletion(self):
        with mock.patch.object(read_service, "_sha256", wraps=read_service._sha256) as catalog_hash:
            response = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/preview")
        self.assertEqual(200, response.status_code)
        catalog_hash.assert_not_called()
        selected = resolve_document(self.document_id, self.root)
        replacement = make_template(Path(self.temp.name) / "replacement.docx", heading="Replacement")
        selected.path.write_bytes(replacement)
        replaced = self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/preview")
        self.assertEqual(200, replaced.status_code); self.assertEqual(replacement, replaced.data)
        selected.path.unlink()
        self.assertEqual(404, self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/preview").status_code)

    def test_direct_read_rejects_a_symlink_component_before_open(self):
        original = read_service._contains_symlink

        def injected(root, candidate, **kwargs):
            if candidate.suffix.casefold() == ".docx":
                return True
            return original(root, candidate, **kwargs)

        clear_document_template_read_cache()
        with mock.patch.object(read_service, "_contains_symlink", side_effect=injected):
            self.assertEqual(404, self.client.get(f"/dashboard/release-monitor/document-templates/documents/{self.document_id}/preview").status_code)

    def test_concurrent_read_and_vendor_caches_lru_ttl_and_invalidation(self):
        failures = []

        def read_worker():
            try:
                resolve_document(self.document_id, self.root)
                verify_vendor_assets(PROJECT_ROOT / "static")
            except Exception as exc:
                failures.append(type(exc).__name__)

        with ThreadPoolExecutor(max_workers=8) as pool:
            list(pool.map(lambda _: read_worker(), range(40)))
        self.assertEqual([], failures)
        self.assertLessEqual(len(read_service._PATH_CACHE), read_service._READ_CACHE_MAX_ROOTS)
        self.assertLessEqual(len(vendor_service._CACHE), vendor_service._CACHE_MAX_ROOTS)
        path_key = str(self.root.absolute())
        timestamp, mapping = read_service._PATH_CACHE[path_key]
        read_service._PATH_CACHE[path_key] = (time.monotonic() - read_service._READ_CACHE_TTL_SECONDS - 1, mapping)
        resolve_document(self.document_id, self.root)
        self.assertGreater(read_service._PATH_CACHE[path_key][0], timestamp)
        vendor_key = str((PROJECT_ROOT / "static").resolve())
        vendor_timestamp, signature, result = vendor_service._CACHE[vendor_key]
        vendor_service._CACHE[vendor_key] = (time.monotonic() - vendor_service._CACHE_TTL_SECONDS - 1, signature, result)
        verify_vendor_assets(PROJECT_ROOT / "static")
        self.assertGreater(vendor_service._CACHE[vendor_key][0], vendor_timestamp)
        read_roots = []
        for index in range(read_service._READ_CACHE_MAX_ROOTS + 1):
            root = Path(self.temp.name) / f"lru-root-{index}"
            path = root / "CAT" / "Kit (12345)" / "template.docx"
            make_template(path)
            read_roots.append(root)
            resolve_document(read_service.document_id_for_relative_path("CAT/Kit (12345)/template.docx"), root)
        self.assertNotIn(str(read_roots[0].absolute()), read_service._PATH_CACHE)
        vendor_roots = []
        for index in range(vendor_service._CACHE_MAX_ROOTS + 1):
            static_root = Path(self.temp.name) / f"static-{index}"
            manifest = static_root / "vendor" / "manifest.json"
            manifest.parent.mkdir(parents=True)
            manifest.write_text('{"assets": []}', encoding="utf-8")
            vendor_roots.append(static_root)
            self.assertTrue(verify_vendor_assets(static_root)["ok"])
        self.assertNotIn(str(vendor_roots[0].resolve()), vendor_service._CACHE)
        clear_document_template_read_cache(self.root)
        clear_document_template_vendor_cache(PROJECT_ROOT / "static")
        self.assertNotIn(str(self.root.absolute()), read_service._PATH_CACHE)
        self.assertNotIn(str((PROJECT_ROOT / "static").resolve()), vendor_service._CACHE)


if __name__ == "__main__":
    unittest.main()
