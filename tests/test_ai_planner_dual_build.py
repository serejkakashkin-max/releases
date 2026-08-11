import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest import mock

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from routes.release_routes import (
    detect_release_template_from_values,
    get_previous_build_versions_from_monitor_items,
)
from routes.dashboard_routes import _build_release_monitor_template_hints
from services.ai_planner_document_service import (
    BUILDER_DPM_URL,
    PRIMARY_DPM_URL,
    adapt_ai_planner_document,
)
from services.document_template_validation_service import validate_candidate
from services.jira_service import _build_release_snapshot
from services.release_build_service import normalize_release_builds, resolve_ai_planner_builds


PLANNER = {
    "version": "D-01.001.45.e2e-planner",
    "url": "registry.ca.sbrf.ru/team/ci14061745/e2e-planner:45",
    "PARENT_CI": "CI14061745",
}
BUILDER = {
    "version": "D-01.001.46.js-business-plan-builder",
    "url": "registry.ca.sbrf.ru/team/ci14061745/js-business-plan-builder:46",
    "PARENT_CI": "CI14061745",
}


def _add_hyperlink(paragraph, label, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_plan_row(table, number, token, *, dpm=True, time_window=""):
    cells = table.add_row().cells
    cells[0].text = f"{number}."
    cells[1].text = f"Средствами DevOps установить сборку {token} {time_window}".strip()
    if dpm:
        _add_hyperlink(cells[2].paragraphs[0], "Ссылка на дистрибутив", PRIMARY_DPM_URL)
    return cells


def _plan_document():
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "№"
    _add_plan_row(table, 2, "RELEASE_VERSION", time_window="21:15–23:05")
    _add_plan_row(table, 3, "Следующий шаг", dpm=False)
    _add_plan_row(table, 4, "PREV_VERSION")
    return document


def _checklist_document():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "№"
    cells = table.add_row().cells
    cells[0].text = "1."
    cells[1].text = "Проверить сборку RELEASE_VERSION"
    return document


def _roundtrip(document):
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return Document(buffer)


def _hyperlink_targets(document):
    return {
        str(relationship.target_ref)
        for relationship in document.part.rels.values()
        if relationship.reltype == RT.HYPERLINK
    }


class ReleaseBuildNormalizationTests(unittest.TestCase):
    def test_template_detection_normalizes_distributive_ke_forms(self):
        entry = {
            "category": "AI_AGENTS",
            "release_clean": "AEF Containers AI-Planner E2E",
            "release_full": "AEF Containers AI-Planner E2E(14061745)",
            "variant": "",
            "requires_playbooks": False,
        }
        context = {"entries": [entry], "by_ke": {"14061745": [entry]}}
        for raw_ke in ("14061745", "CI14061745", "140-617-45"):
            with self.subTest(raw_ke=raw_ke):
                detection = detect_release_template_from_values(
                    raw_ke,
                    "AI Planner",
                    catalog_context=context,
                )
                self.assertTrue(detection["found"])
                self.assertEqual("14061745", detection["template_sm_id"])

    def test_builder_then_planner_is_normalized_to_planner_then_builder(self):
        result = resolve_ai_planner_builds([BUILDER, PLANNER], ke_id="14061745")
        self.assertTrue(result["applies"])
        self.assertEqual(
            ["e2e_planner", "business_plan_builder"],
            [item["component"] for item in result["builds"]],
        )
        self.assertEqual(PRIMARY_DPM_URL, result["builds"][0]["dpm_url"])
        self.assertEqual(BUILDER_DPM_URL, result["builds"][1]["dpm_url"])

    def test_rule_does_not_apply_to_another_ke(self):
        result = resolve_ai_planner_builds([PLANNER, BUILDER], ke_id="14061746")
        self.assertFalse(result["applies"])
        self.assertEqual([], result["builds"])

    def test_single_build_and_ambiguous_component_contracts(self):
        planner = resolve_ai_planner_builds([PLANNER], ke_id="CI14061745")
        builder = resolve_ai_planner_builds([BUILDER], ke_id="CI14061745")
        duplicate = resolve_ai_planner_builds(
            [PLANNER, {**PLANNER, "url": "registry.ca.sbrf.ru/mirror/e2e-planner:45"}],
            ke_id="14061745",
        )
        ambiguous = resolve_ai_planner_builds(
            [PLANNER, {**PLANNER, "version": "D-01.001.46.e2e-planner"}],
            ke_id="14061745",
        )
        self.assertEqual(["e2e_planner"], [item["component"] for item in planner["builds"]])
        self.assertEqual(["business_plan_builder"], [item["component"] for item in builder["builds"]])
        self.assertEqual(1, len(duplicate["builds"]))
        self.assertEqual(["e2e_planner"], ambiguous["ambiguous_components"])

    def test_jira_snapshot_preserves_primary_aliases(self):
        issue = {"fields": {
            "summary": "Релиз AI-Планировщик E2E(14061745)",
            "customfield_21711": {"id": "14061745", "value": "AEF Containers AI-Планировщик E2E(14061745)"},
            "customfield_27011": [BUILDER, PLANNER],
        }}
        snapshot = _build_release_snapshot("AIGAS-1235", "jira.delta.sbrf.ru", issue)
        self.assertEqual("D-01.001.45.e2e-planner", snapshot["release_version"])
        self.assertEqual(snapshot["release_builds"][0]["artifact_url"], snapshot["release_dist_url"])

    def test_unrecognized_special_artifact_does_not_become_silent_primary(self):
        issue = {"fields": {
            "summary": "Релиз AI-Планировщик E2E(14061745)",
            "customfield_27011": [{
                "version": "D-01.001.99.unrelated-component",
                "url": "registry.ca.sbrf.ru/ci14061745/unrelated:D-01.001.99.unrelated-component",
                "PARENT_CI": "CI14061745",
            }],
        }}
        snapshot = _build_release_snapshot("AIGAS-1236", "jira.delta.sbrf.ru", issue)
        self.assertEqual([], snapshot["release_builds"])
        self.assertEqual("", snapshot["release_version"])

    def test_old_snapshot_hydrates_single_legacy_build(self):
        self.assertEqual([{
            "component": "legacy",
            "label": "Сборка",
            "version": "D-01.001.01-old",
            "artifact_url": "https://registry/old",
            "dpm_url": "",
        }], normalize_release_builds({
            "release_version": "D-01.001.01-old",
            "release_dist_url": "https://registry/old",
        }))


class ComponentRollbackLookupTests(unittest.TestCase):
    def test_page_hint_contains_both_previous_versions_before_jira_init(self):
        builds = resolve_ai_planner_builds([PLANNER, BUILDER], ke_id="14061745")["builds"]
        previous = {
            "row_key": "AIGAS-1200::1", "release_key": "AIGAS-1200", "release_number": 4,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_builds": builds,
        }
        current = {
            "row_key": "AIGAS-1235::1", "release_key": "AIGAS-1235", "release_number": 5,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_builds": [
                {**builds[0], "version": "D-01.001.47.e2e-planner"},
                {**builds[1], "version": "D-01.001.48.js-business-plan-builder"},
            ],
        }
        detection = {
            "found": True,
            "category": "AI_AGENTS",
            "release_clean": "Planner",
            "release_full": "Planner(14061745)",
        }
        with mock.patch(
            "routes.dashboard_routes.build_release_template_detection_context",
            return_value={"entries": [], "by_ke": {}},
        ), mock.patch(
            "routes.dashboard_routes.detect_release_template_from_values",
            return_value=detection,
        ):
            hints = _build_release_monitor_template_hints([previous, current])

        previous_versions = hints[current["row_key"]]["previous_build_versions"]
        self.assertEqual(PLANNER["version"], previous_versions["e2e_planner"])
        self.assertEqual(BUILDER["version"], previous_versions["business_plan_builder"])

    def test_previous_versions_are_resolved_per_component(self):
        previous_planner = {
            "row_key": "AIGAS-1199::1", "release_key": "AIGAS-1199", "release_number": 3,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_version": PLANNER["version"], "release_dist_url": PLANNER["url"],
        }
        previous_builder = {
            "row_key": "AIGAS-1200::1", "release_key": "AIGAS-1200", "release_number": 4,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_version": BUILDER["version"], "release_dist_url": BUILDER["url"],
        }
        current_builds = resolve_ai_planner_builds([
            {**PLANNER, "version": "D-01.001.47.e2e-planner"},
            {**BUILDER, "version": "D-01.001.48.js-business-plan-builder"},
        ], ke_id="14061745")["builds"]
        current = {
            "row_key": "AIGAS-1235::1", "release_key": "AIGAS-1235", "release_number": 5,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_builds": current_builds,
        }
        resolved = get_previous_build_versions_from_monitor_items(
            [previous_planner, previous_builder, current], current["row_key"], current["release_key"]
        )
        self.assertEqual(PLANNER["version"], resolved["e2e_planner"])
        self.assertEqual(BUILDER["version"], resolved["business_plan_builder"])

    def test_component_lookup_keeps_previous_year_fallback(self):
        builds = resolve_ai_planner_builds([PLANNER, BUILDER], ke_id="14061745")["builds"]
        previous = {
            "row_key": "AIGAS-999::1", "release_key": "AIGAS-999", "release_number": 99,
            "year": 2025, "ke_id": "14061745", "release_type": "planned",
            "release_builds": builds,
        }
        current = {
            "row_key": "AIGAS-1000::1", "release_key": "AIGAS-1000", "release_number": 1,
            "year": 2026, "ke_id": "14061745", "release_type": "planned",
            "release_builds": [
                {**builds[0], "version": "D-01.001.50.e2e-planner"},
                {**builds[1], "version": "D-01.001.51.js-business-plan-builder"},
            ],
        }
        resolved = get_previous_build_versions_from_monitor_items(
            [previous, current], current["row_key"], current["release_key"]
        )
        self.assertEqual(PLANNER["version"], resolved["e2e_planner"])
        self.assertEqual(BUILDER["version"], resolved["business_plan_builder"])


class AiPlannerDocumentTests(unittest.TestCase):
    def setUp(self):
        self.builds = resolve_ai_planner_builds([PLANNER, BUILDER], ke_id="14061745")["builds"]
        self.previous = {
            "e2e_planner": "D-01.001.44.e2e-planner",
            "business_plan_builder": "D-01.001.45.js-business-plan-builder",
        }

    def test_dual_plan_clones_install_and_rollback_rows(self):
        document = adapt_ai_planner_document(
            _plan_document(), template_name="План внедрения и отката.docx",
            release_builds=self.builds, previous_versions=self.previous,
        )
        document = _roundtrip(document)
        all_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn(BUILDER["version"], all_text)
        self.assertIn(self.previous["business_plan_builder"], all_text)
        self.assertEqual(2, all_text.count("21:15–23:05"))
        self.assertIn(PRIMARY_DPM_URL, _hyperlink_targets(document))
        self.assertIn(BUILDER_DPM_URL, _hyperlink_targets(document))
        numbers = [row.cells[0].text for row in document.tables[0].rows[1:]]
        self.assertEqual(["1.", "2.", "3.", "4.", "5."], numbers)

    def test_dual_checklist_clones_build_check(self):
        document = adapt_ai_planner_document(
            _checklist_document(), template_name="Чек-лист.docx",
            release_builds=self.builds, previous_versions=self.previous,
        )
        text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
        self.assertIn("RELEASE_VERSION", text)
        self.assertIn(BUILDER["version"], text)

    def test_single_build_does_not_add_rows_and_builder_only_uses_sub_link(self):
        planner_document = _plan_document()
        before = len(planner_document.tables[0].rows)
        adapt_ai_planner_document(
            planner_document, template_name="План внедрения и отката.docx",
            release_builds=[self.builds[0]], previous_versions={"e2e_planner": self.previous["e2e_planner"]},
        )
        self.assertEqual(before, len(planner_document.tables[0].rows))

        builder_document = adapt_ai_planner_document(
            _plan_document(), template_name="План внедрения и отката.docx",
            release_builds=[self.builds[1]],
            previous_versions={"business_plan_builder": self.previous["business_plan_builder"]},
        )
        self.assertIn(BUILDER_DPM_URL, _hyperlink_targets(builder_document))
        self.assertEqual(before, len(builder_document.tables[0].rows))

    def test_missing_builder_previous_version_fails_closed(self):
        with self.assertRaisesRegex(ValueError, "Builder"):
            adapt_ai_planner_document(
                _plan_document(), template_name="План внедрения и отката.docx",
                release_builds=self.builds, previous_versions={"e2e_planner": "old"},
            )

    def test_dtc_blocks_candidate_without_generator_dpm_anchor(self):
        with tempfile.TemporaryDirectory() as directory:
            template_dir = Path(directory) / "AEF Containers AI-Планировщик E2E(14061745)"
            template_dir.mkdir()
            active = template_dir / "План внедрения и отката.docx"
            candidate = template_dir / "candidate.docx"
            _plan_document().save(active)
            broken = _plan_document()
            for table in broken.tables:
                for row in table.rows:
                    for hyperlink in list(row._tr.iter(qn("w:hyperlink"))):
                        hyperlink.getparent().remove(hyperlink)
            broken.save(candidate)
            result = validate_candidate(active, candidate)
            self.assertFalse(result["ok"])
            self.assertIn("ai_planner_generator_contract", {item["code"] for item in result["errors"]})


class ReleaseMonitorFrontendContractTests(unittest.TestCase):
    def test_table_and_wizard_expose_dual_build_contract(self):
        source = Path("static/js/oplot_release.js").read_text(encoding="utf-8")
        self.assertIn("release-builds__list", source)
        self.assertNotIn('class="release-builds__label"', source)
        self.assertIn("requiresLiveBuildDetection", source)
        self.assertNotIn("release-doc-build-summary", source)
        self.assertIn("detection?.previous_build_versions || {}", source)
        self.assertIn("releaseDocumentSecondaryPrevVersion", source)
        self.assertIn("secondary_prev_version: state.form.secondary_prev_version", source)
        self.assertIn("release_builds_ambiguous", source)


if __name__ == "__main__":
    unittest.main()
