from __future__ import annotations

import io
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from docx import Document
from flask import Flask

from tests._support import PROJECT_ROOT, prepare_config_import

prepare_config_import()

from routes import release_routes
from routes.document_template_routes import document_template_bp
from routes.sup_admin_session_routes import sup_admin_session_bp
from services.document_template_csrf_service import CSRF_COOKIE_NAME
from services.document_template_creation_service import kit_id_for_relative_dir
from services.document_template_read_service import build_document_whitelist
from services.document_template_storage_service import (
    cache_root,
    candidate_directory,
    create_history_version,
    data_root,
    write_uploaded_candidate,
)
from services.oplot_ui_service import register_oplot_ui
from services.release_template_catalog_service import clear_template_catalog_cache


def make_docx(
    path: Path,
    *,
    heading: str = "Template",
    placeholder: str = "RELEASE_VERSION",
    include_jira_table: bool = False,
) -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(f"Version {placeholder}")
    document.add_paragraph("Date DATE")
    if include_jira_table:
        table = document.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = "№"
        table.rows[0].cells[1].text = "ЗНИ/JIRA ID"
        table.rows[0].cells[2].text = "Issue"
        table.rows[0].cells[3].text = "Issue Type"
    document.save(path)
    return path.read_bytes()


def add_external_image_relation(path: Path) -> bytes:
    replacement = path.with_suffix(".external.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(replacement, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            payload = source.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                text = payload.decode("utf-8")
                relation = (
                    '<Relationship Id="rIdExternalImage" '
                    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                    'Target="https://example.invalid/legacy.png" TargetMode="External"/>'
                )
                text = text.replace("</Relationships>", relation + "</Relationships>")
                payload = text.encode("utf-8")
            target.writestr(item, payload)
    path.write_bytes(replacement.read_bytes())
    replacement.unlink()
    return path.read_bytes()


def build_app(root: Path, runtime: Path) -> Flask:
    app = Flask(
        __name__,
        template_folder=str(PROJECT_ROOT / "templates"),
        static_folder=str(PROJECT_ROOT / "static"),
    )
    app.config.update(
        TESTING=True,
        SECRET_KEY="dtc-create-secret-0123456789abcdef",
        DOCUMENT_TEMPLATE_CENTER_ROOT=root,
        DOCUMENT_TEMPLATE_CENTER_RUNTIME_ROOT=runtime,
    )
    app.add_url_rule("/", endpoint="main.index", view_func=lambda: "home")
    app.add_url_rule("/help", endpoint="main.help_page", view_func=lambda: "help")
    app.add_url_rule("/dashboard", endpoint="dashboard.dashboard", view_func=lambda: "dashboard")
    app.add_url_rule("/release-monitor", endpoint="dashboard.release_monitor_page", view_func=lambda: "monitor")
    app.add_url_rule("/mpr", endpoint="mpr.mpr_page", view_func=lambda: "mpr")
    app.register_blueprint(document_template_bp)
    app.register_blueprint(sup_admin_session_bp)
    register_oplot_ui(app)
    return app


class DocumentTemplateCreationTests(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        self.base = Path(self.temporary.name)
        self.root = self.base / "doc_templates"
        self.runtime = self.base / "runtime"
        self.active_payload = make_docx(self.root / "PLATFORM" / "Base Kit PL (10001)" / "План внедрения.docx")
        clear_template_catalog_cache()
        self.app = build_app(self.root, self.runtime)
        self.client = self.app.test_client()
        self.client.get("/dashboard/release-monitor/document-templates")
        self.csrf = self.client.get_cookie(
            CSRF_COOKIE_NAME,
            path="/dashboard/release-monitor/document-templates/",
        ).value

    def tearDown(self):
        clear_template_catalog_cache()
        self.temporary.cleanup()

    def _draft_url_from(self, response):
        self.assertEqual(303, response.status_code, response.get_data(as_text=True))
        return response.headers["Location"]

    def _publish(self, draft_url: str):
        draft_page = self.client.get(draft_url).get_data(as_text=True)
        marker = "/dashboard/release-monitor/document-templates/drafts/"
        start = draft_url.index(marker) + len(marker)
        draft_uuid = draft_url[start:].split("/", 1)[0].split("?", 1)[0]
        self.assertIn("Проверка пройдена", draft_page)
        return self.client.post(
            f"/dashboard/release-monitor/document-templates/drafts/{draft_uuid}/publish",
            data={"_csrf_token": self.csrf},
        )

    def test_create_kit_from_uploaded_docx_and_publish_atomically(self):
        payload = make_docx(self.base / "new.docx", heading="New kit")
        response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "name": "New Service PL",
                "kit_ke": "20002",
                "source_mode": "upload",
                "target_names": ["План внедрения.docx"],
                "files": (io.BytesIO(payload), "source.docx"),
            },
            content_type="multipart/form-data",
        )
        draft_url = self._draft_url_from(response)
        publish = self._publish(draft_url)
        self.assertEqual(303, publish.status_code, publish.get_data(as_text=True))
        self.assertTrue((self.root / "PLATFORM" / "PLATFORM(20002)" / "План внедрения.docx").is_file())
        self.assertIn("created=kit", publish.headers["Location"])

        catalog_page = self.client.get("/dashboard/release-monitor/document-templates").get_data(as_text=True)
        self.assertNotIn("Название без КЭ", catalog_page)
        self.assertNotIn('id="kit-name"', catalog_page)

    def test_create_kit_copy_requires_variant_when_ke_is_reused(self):
        response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "name": "Another Kit",
                "kit_ke": "10001",
                "source_mode": "copy",
            },
        )
        self.assertEqual(400, response.status_code)
        self.assertIn("КЭ уже используется", response.get_data(as_text=True))

    def test_copy_uses_variant_as_new_kit_name_instead_of_source_name(self):
        source_kit_id = kit_id_for_relative_dir("PLATFORM/Base Kit PL (10001)")
        response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "kit_ke": "12345678",
                "variant": "TESTTTT",
                "source_mode": "copy",
                "source_kit_id": source_kit_id,
            },
        )
        draft_url = self._draft_url_from(response)
        draft_page = self.client.get(draft_url).get_data(as_text=True)
        self.assertIn("TESTTTT(12345678)", draft_page)
        self.assertNotIn("Base Kit PL TESTTTT(12345678)", draft_page)
        publish = self._publish(draft_url)
        self.assertEqual(303, publish.status_code, publish.get_data(as_text=True))
        self.assertTrue((self.root / "PLATFORM" / "TESTTTT(12345678)" / "План внедрения.docx").is_file())

    def test_copy_existing_kit_keeps_legacy_external_images_as_warning(self):
        source_dir = self.root / "PLATFORM" / "Legacy Image Kit (55555)"
        payload = make_docx(source_dir / "Legacy.docx", heading="Legacy image")
        payload = add_external_image_relation(source_dir / "Legacy.docx")
        clear_template_catalog_cache()
        source_kit_id = kit_id_for_relative_dir("PLATFORM/Legacy Image Kit (55555)")

        copy_response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "name": "Copied Legacy Kit",
                "kit_ke": "55556",
                "source_mode": "copy",
                "source_kit_id": source_kit_id,
            },
        )
        copy_draft_url = self._draft_url_from(copy_response)
        copy_page = self.client.get(copy_draft_url).get_data(as_text=True)
        self.assertIn("external_image_not_loaded", copy_page)
        self.assertIn("Проверка пройдена", copy_page)
        copy_publish = self._publish(copy_draft_url)
        self.assertEqual(303, copy_publish.status_code, copy_publish.get_data(as_text=True))
        self.assertTrue((self.root / "PLATFORM" / "Legacy Image Kit(55556)" / "Legacy.docx").is_file())

        upload_response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "name": "Uploaded Legacy Kit",
                "kit_ke": "55557",
                "source_mode": "upload",
                "target_names": ["Legacy.docx"],
                "files": (io.BytesIO(payload), "Legacy.docx"),
            },
            content_type="multipart/form-data",
        )
        upload_page = self.client.get(self._draft_url_from(upload_response)).get_data(as_text=True)
        self.assertIn("external_image_not_loaded", upload_page)
        self.assertNotIn("Документ нельзя опубликовать", upload_page)
        self.assertIn("Проверка пройдена", upload_page)
        upload_publish = self._publish(self._draft_url_from(upload_response))
        self.assertEqual(303, upload_publish.status_code, upload_publish.get_data(as_text=True))
        self.assertTrue((self.root / "PLATFORM" / "PLATFORM(55557)" / "Legacy.docx").is_file())

    def test_unknown_placeholder_blocks_publication_and_creates_no_partial_kit(self):
        payload = make_docx(self.base / "invalid.docx", placeholder="{{UNKNOWN_KEY}}")
        response = self.client.post(
            "/dashboard/release-monitor/document-templates/kits/drafts",
            data={
                "_csrf_token": self.csrf,
                "category": "PLATFORM",
                "name": "Invalid Kit",
                "kit_ke": "30003",
                "source_mode": "upload",
                "target_names": ["bad.docx"],
                "files": (io.BytesIO(payload), "bad.docx"),
            },
            content_type="multipart/form-data",
        )
        draft_url = self._draft_url_from(response)
        draft_page = self.client.get(draft_url).get_data(as_text=True)
        self.assertIn("UNKNOWN_KEY", draft_page)
        marker = "/dashboard/release-monitor/document-templates/drafts/"
        draft_uuid = draft_url[draft_url.index(marker) + len(marker):].split("/", 1)[0]
        publish = self.client.post(
            f"/dashboard/release-monitor/document-templates/drafts/{draft_uuid}/publish",
            data={"_csrf_token": self.csrf},
        )
        self.assertEqual(409, publish.status_code)
        self.assertFalse((self.root / "PLATFORM" / "Invalid Kit(30003)").exists())

    def test_add_document_rejects_duplicate_name_and_allows_second_unique_document(self):
        kit_id = kit_id_for_relative_dir("PLATFORM/Base Kit PL (10001)")
        duplicate = self.client.post(
            f"/dashboard/release-monitor/document-templates/kits/{kit_id}/documents/drafts",
            data={
                "_csrf_token": self.csrf,
                "target_filename": "план внедрения.docx",
                "file": (io.BytesIO(self.active_payload), "duplicate.docx"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(409, duplicate.status_code)

        payload = make_docx(self.base / "second.docx", heading="Second plan")
        response = self.client.post(
            f"/dashboard/release-monitor/document-templates/kits/{kit_id}/documents/drafts",
            data={
                "_csrf_token": self.csrf,
                "target_filename": "План внедрения 2.docx",
                "file": (io.BytesIO(payload), "second.docx"),
            },
            content_type="multipart/form-data",
        )
        publish = self._publish(self._draft_url_from(response))
        self.assertEqual(303, publish.status_code)
        self.assertTrue((self.root / "PLATFORM" / "Base Kit PL (10001)" / "План внедрения 2.docx").is_file())

    def test_delete_document_requires_only_dtc_csrf_and_exact_confirmation(self):
        make_docx(self.root / "PLATFORM" / "Base Kit PL (10001)" / "Чек-лист.docx")
        clear_template_catalog_cache()
        document = next(iter(build_document_whitelist(self.root).values()))
        url = f"/dashboard/release-monitor/document-templates/documents/{document.document_id}/delete"
        missing_csrf = self.client.post(
            url,
            data={"confirmation": document.filename},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(403, missing_csrf.status_code)
        self.assertTrue(document.path.exists())

        wrong_name = self.client.post(
            url,
            data={"_csrf_token": self.csrf, "confirmation": "другой документ.docx"},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(400, wrong_name.status_code)
        self.assertTrue(document.path.exists())

        deleted = self.client.post(
            url,
            data={"_csrf_token": self.csrf, "confirmation": document.filename},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(200, deleted.status_code)
        self.assertTrue(deleted.get_json()["success"])
        self.assertFalse(document.path.exists())

    def test_delete_document_is_blocked_by_active_candidate_without_admin_session(self):
        document = next(iter(build_document_whitelist(self.root).values()))
        url = f"/dashboard/release-monitor/document-templates/documents/{document.document_id}/delete"

        with self.app.app_context():
            write_uploaded_candidate(
                io.BytesIO(self.active_payload),
                document_id=document.document_id,
                source_filename="new.docx",
                active_filename=document.filename,
                active_sha=document.sha256,
                uploaded_by="tester",
                comment="active draft",
            )
        blocked = self.client.post(
            url,
            data={"_csrf_token": self.csrf, "confirmation": document.filename},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(409, blocked.status_code)
        self.assertTrue(document.path.exists())

    def test_delete_kit_uses_exact_confirmation_without_admin_session(self):
        kit_id = kit_id_for_relative_dir("PLATFORM/Base Kit PL (10001)")
        url = f"/dashboard/release-monitor/document-templates/kits/{kit_id}/delete"
        wrong_name = self.client.post(
            url,
            data={"_csrf_token": self.csrf, "confirmation": "Base Kit PL"},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(400, wrong_name.status_code)
        self.assertTrue((self.root / "PLATFORM" / "Base Kit PL (10001)").exists())

        deleted = self.client.post(
            url,
            data={"_csrf_token": self.csrf, "confirmation": "Base Kit PL (10001)"},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(200, deleted.status_code)
        self.assertTrue(deleted.get_json()["success"])
        self.assertFalse((self.root / "PLATFORM" / "Base Kit PL (10001)").exists())

    def test_deleted_kit_is_not_resurrected_by_stale_release_catalog(self):
        kit_id = kit_id_for_relative_dir("PLATFORM/Base Kit PL (10001)")
        stale_entry = {
            "category": "PLATFORM",
            "release_clean": "Base Kit PL",
            "release_full": "Base Kit PL (10001)",
            "ke": "10001",
            "variant": "PL",
            "requires_playbooks": False,
        }
        deleted = self.client.post(
            f"/dashboard/release-monitor/document-templates/kits/{kit_id}/delete",
            data={"_csrf_token": self.csrf, "confirmation": "Base Kit PL (10001)"},
            headers={"Accept": "application/json"},
        )
        self.assertEqual(200, deleted.status_code)

        with (
            mock.patch.object(release_routes, "DOC_TEMPLATES_ROOT", self.root),
            mock.patch.object(release_routes, "find_template_entries_by_ke", return_value=[stale_entry]),
            mock.patch.object(release_routes, "ID_MAP", {"10001": [("PLATFORM", "Base Kit PL")]}),
            mock.patch.object(
                release_routes,
                "RELEASE_STRUCTURE",
                {"PLATFORM": [("Base Kit PL", "Base Kit PL (10001)")]},
            ),
        ):
            detection = release_routes.detect_release_template_from_values("10001", "Base Kit PL")

        self.assertFalse(detection["found"])
        self.assertEqual([], detection["candidates"])

    def test_delete_kit_cascades_history_candidates_and_creation_drafts(self):
        kit_path = self.root / "PLATFORM" / "Base Kit PL (10001)"
        document = next(
            item
            for item in build_document_whitelist(self.root).values()
            if item.relative_directory.casefold() == "platform/base kit pl (10001)"
        )
        kit_id = kit_id_for_relative_dir("PLATFORM/Base Kit PL (10001)")
        with self.app.app_context():
            candidate = write_uploaded_candidate(
                io.BytesIO(self.active_payload),
                document_id=document.document_id,
                source_filename="new.docx",
                active_filename=document.filename,
                active_sha=document.sha256,
                uploaded_by="tester",
                comment="active draft",
            )
            history = create_history_version(
                document.document_id,
                self.active_payload,
                {
                    "state": "committed",
                    "sha256": document.sha256,
                    "created_at": "2026-08-21T10:00:00+00:00",
                },
            )
            creation_uuid = "11111111-1111-4111-8111-111111111111"
            creation_dir = cache_root() / "creation-drafts" / creation_uuid
            creation_dir.mkdir(parents=True)
            (creation_dir / "metadata.json").write_text(
                '{"draft_uuid":"' + creation_uuid + '","target_kit_id":"' + kit_id + '"}',
                encoding="utf-8",
            )

        deleted = self.client.post(
            f"/dashboard/release-monitor/document-templates/kits/{kit_id}/delete",
            data={"_csrf_token": self.csrf, "confirmation": "Base Kit PL (10001)"},
            headers={"Accept": "application/json"},
        )

        self.assertEqual(200, deleted.status_code)
        self.assertFalse(kit_path.exists())
        with self.app.app_context():
            self.assertFalse(candidate_directory(candidate["candidate_uuid"]).exists())
            self.assertFalse((data_root() / "history" / document.document_id / history["version_uuid"]).exists())
            self.assertFalse(creation_dir.exists())

    def test_guide_contains_public_placeholder_reference(self):
        response = self.client.get("/dashboard/release-monitor/document-templates/guide")
        text = response.get_data(as_text=True)
        self.assertEqual(200, response.status_code)
        for key in ("RELEASE_VERSION", "PREV_VERSION", "RELEASE_ID", "RELNUMBER", "OPLOT", "CHECKER", "DATE", "PLUS_1", "POB", "PLAYBOOKS", "INSTRUCTION_BLOCK"):
            self.assertIn(key, text)
        self.assertNotIn("<code>ИНСТРУКЦИЯ</code>", text)
        self.assertNotIn("Jira-таблица", text)
        self.assertNotIn(str(self.root), text)


if __name__ == "__main__":
    unittest.main()
