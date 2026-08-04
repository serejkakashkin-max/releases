"""Local synthetic app for the Stage 2 desktop visual gate.

It never reads or mutates the repository's real doc_templates directory.
"""
from __future__ import annotations

import io
import json
import os
import sys
import zipfile
from pathlib import Path

from docx import Document
from flask import Flask

from tests._support import PROJECT_ROOT, prepare_config_import

prepare_config_import()

from routes.document_template_routes import document_template_bp
from services.document_template_candidate_service import validate_staged_candidate
from services.document_template_publish_service import publish_candidate
from services.document_template_read_service import build_document_whitelist, resolve_document
from services.oplot_ui_service import register_oplot_ui
from services.document_template_storage_service import (
    create_history_version,
    sha256_bytes,
    update_candidate,
    utc_now,
    write_uploaded_candidate,
)
from services.document_template_validation_service import build_contract
from services.release_template_catalog_service import clear_template_catalog_cache


def _docx(path: Path, heading: str, *, complete: bool = True, jira: bool = True) -> bytes:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(); document.add_heading(heading, level=1)
    document.add_paragraph("Версия RELEASE_VERSION")
    if complete:
        document.add_paragraph("Дата DATE")
    if jira:
        table = document.add_table(rows=1, cols=4)
        headers = table.rows[0].cells
        headers[0].text = "№"; headers[1].text = "ЗНИ/JIRA ID"; headers[2].text = "Issue"; headers[3].text = "Issue Type"
    document.save(path)
    return path.read_bytes()


def create_visual_app(base: Path) -> tuple[Flask, dict]:
    root = base / "doc_templates"; runtime = base / "runtime"; source = base / "uploads"
    _docx(root / "PLATFORM" / "Платформа PL (10001)" / "План внедрения.docx", "План внедрения")
    _docx(root / "PLATFORM" / "Сервис BH (10002)" / "Чек-лист проверки.docx", "Чек-лист")
    _docx(root / "AI" / "AI-помощник PL (10003)" / "Описание агента.docx", "Описание AI-агента")
    _docx(root / "PLATFORM" / "Сервис без Jira PL (10004)" / "Регламент без Jira.docx", "Регламент без Jira", jira=False)
    app = Flask("document-template-visual", template_folder=str(PROJECT_ROOT / "templates"), static_folder=str(PROJECT_ROOT / "static"))
    app.config.update(TESTING=False, DOCUMENT_TEMPLATE_CENTER_ROOT=root, DOCUMENT_TEMPLATE_CENTER_RUNTIME_ROOT=runtime)
    app.add_url_rule("/", endpoint="main.index", view_func=lambda: "Oplot visual fixture")
    app.add_url_rule("/help", endpoint="main.help_page", view_func=lambda: "help")
    app.add_url_rule("/dashboard", endpoint="dashboard.dashboard", view_func=lambda: "dashboard")
    app.add_url_rule("/release-monitor", endpoint="dashboard.release_monitor_page", view_func=lambda: "monitor")
    app.add_url_rule("/mpr", endpoint="mpr.mpr_page", view_func=lambda: "mpr")
    app.register_blueprint(document_template_bp)
    register_oplot_ui(app)
    clear_template_catalog_cache()
    manifest = {"documents": {}}
    with app.app_context():
        documents = sorted(build_document_whitelist(root).values(), key=lambda item: item.filename)
        by_name = {item.filename: item for item in documents}
        validation_doc = by_name["Чек-лист проверки.docx"]
        valid_payload = _docx(source / "valid.docx", "Новая версия чек-листа")
        valid = write_uploaded_candidate(io.BytesIO(valid_payload), document_id=validation_doc.document_id, source_filename="Чек-лист 2.0.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Анна Редактор", comment="Уточнены шаги проверки и визуальная структура")
        validate_staged_candidate(validation_doc.document_id, valid["candidate_uuid"], validation_doc.path)
        uploaded = write_uploaded_candidate(io.BytesIO(valid_payload), document_id=validation_doc.document_id, source_filename="Новый кандидат.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Ольга Редактор", comment="Файл ожидает первичной проверки")
        validating = write_uploaded_candidate(io.BytesIO(valid_payload), document_id=validation_doc.document_id, source_filename="Кандидат в проверке.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Иван Редактор", comment="Проверяем обновлённые формулировки")
        update_candidate(validating["candidate_uuid"], {"state": "validating"})
        invalid_payload = _docx(source / "invalid.docx", "Неполная версия", complete=False)
        invalid = write_uploaded_candidate(io.BytesIO(invalid_payload), document_id=validation_doc.document_id, source_filename="Версия с ошибкой.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Анна Редактор", comment="Тестовое состояние ошибки контракта")
        validate_staged_candidate(validation_doc.document_id, invalid["candidate_uuid"], validation_doc.path)
        bomb_payload = io.BytesIO()
        with zipfile.ZipFile(bomb_payload, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "<Types/>"); archive.writestr("_rels/.rels", "<Relationships/>"); archive.writestr("word/document.xml", b"A" * (2 * 1024 * 1024))
        invalid_security = write_uploaded_candidate(io.BytesIO(bomb_payload.getvalue()), document_id=validation_doc.document_id, source_filename="Небезопасный кандидат.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Ольга Редактор", comment="Синтетическое небезопасное состояние")
        validate_staged_candidate(validation_doc.document_id, invalid_security["candidate_uuid"], validation_doc.path)
        conflict = write_uploaded_candidate(io.BytesIO(valid_payload), document_id=validation_doc.document_id, source_filename="Конфликтный кандидат.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Пётр Редактор", comment="Версия основана на предыдущем active SHA")
        validate_staged_candidate(validation_doc.document_id, conflict["candidate_uuid"], validation_doc.path)
        update_candidate(conflict["candidate_uuid"], {"state": "conflict"})
        recovery_blocked = write_uploaded_candidate(io.BytesIO(valid_payload), document_id=validation_doc.document_id, source_filename="Recovery blocked.docx", active_filename=validation_doc.filename, active_sha=validation_doc.sha256, uploaded_by="Служебная проверка", comment="Требуется контролируемое восстановление")
        validate_staged_candidate(validation_doc.document_id, recovery_blocked["candidate_uuid"], validation_doc.path)
        update_candidate(recovery_blocked["candidate_uuid"], {"state": "publish_failed", "recovery_blocking": True, "error_code": "visual_recovery_block"})
        no_jira_doc = by_name["Регламент без Jira.docx"]
        no_jira_payload = _docx(source / "no-jira.docx", "Новая версия без Jira", jira=False)
        no_jira = write_uploaded_candidate(io.BytesIO(no_jira_payload), document_id=no_jira_doc.document_id, source_filename="Регламент без Jira 2.0.docx", active_filename=no_jira_doc.filename, active_sha=no_jira_doc.sha256, uploaded_by="Анна Редактор", comment="Шаблон без Jira-таблицы успешно проверен")
        validate_staged_candidate(no_jira_doc.document_id, no_jira["candidate_uuid"], no_jira_doc.path)
        blocked_history = create_history_version(no_jira_doc.document_id, no_jira_doc.path.read_bytes(), {
            "created_at": utc_now(), "updated_at": utc_now(), "state": "prepared",
            "source_filename": no_jira_doc.filename, "sha256": sha256_bytes(no_jira_doc.path.read_bytes()),
            "actor": "Служебная проверка", "action": "publish_previous",
            "comment": "Визуальная проверка заблокированной истории", "contract": build_contract(no_jira_doc.path),
        })
        history_doc = by_name["Описание агента.docx"]
        history_payload = _docx(source / "published.docx", "Опубликованная версия AI-агента")
        published = write_uploaded_candidate(io.BytesIO(history_payload), document_id=history_doc.document_id, source_filename="Описание агента 2.0.docx", active_filename=history_doc.filename, active_sha=history_doc.sha256, uploaded_by="Мария Редактор", comment="Обновлено описание поведения агента")
        validate_staged_candidate(history_doc.document_id, published["candidate_uuid"], history_doc.path)
        publish_candidate(history_doc, published["candidate_uuid"], "Мария Редактор")
        manifest["documents"] = {name: item.document_id for name, item in by_name.items()}
        manifest["candidates"] = {"uploaded": uploaded["candidate_uuid"], "valid": valid["candidate_uuid"], "validating": validating["candidate_uuid"], "invalid_security": invalid_security["candidate_uuid"], "invalid_contract": invalid["candidate_uuid"], "conflict": conflict["candidate_uuid"], "recovery_blocked": recovery_blocked["candidate_uuid"], "no_jira": no_jira["candidate_uuid"]}
        manifest["blocked_history"] = {"document_id": no_jira_doc.document_id, "version_uuid": blocked_history["version_uuid"]}
    (base / "visual_manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return app, manifest


def main() -> int:
    base_value = os.environ.get("OPLOT_VISUAL_FIXTURE_ROOT")
    if not base_value:
        raise SystemExit("OPLOT_VISUAL_FIXTURE_ROOT is required")
    base = Path(base_value).resolve(); base.mkdir(parents=True, exist_ok=True)
    app, manifest = create_visual_app(base)
    port = int(os.environ.get("OPLOT_VISUAL_FIXTURE_PORT", "5127"))
    print(json.dumps({"url": f"http://127.0.0.1:{port}/dashboard/release-monitor/document-templates", **manifest}, ensure_ascii=False), flush=True)
    app.run(host="127.0.0.1", port=port, debug=False, use_reloader=False, threaded=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
