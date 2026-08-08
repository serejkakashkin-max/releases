from __future__ import annotations

import csv
import io
import os
import re
import tempfile
import unittest
from pathlib import Path
from unittest import mock
from zipfile import ZipFile

from docx import Document
from flask import Flask
from openpyxl import Workbook
from werkzeug.datastructures import FileStorage

from tests._support import PROJECT_ROOT, prepare_config_import

prepare_config_import()

from routes.mpr_routes import mpr_bp
from services import mpr_service
from services.mpr_service import MprError
from services.mpr_ui_service import build_mpr_ui_config
from services.oplot_ui_service import register_oplot_ui


MPR_TEMPLATE = PROJECT_ROOT / "templates" / "mpr.html"
MPR_CSS = PROJECT_ROOT / "static" / "css" / "oplot_mpr.css"
MPR_JS = PROJECT_ROOT / "static" / "js" / "oplot_mpr.js"


def build_app() -> Flask:
    app = Flask(
        __name__,
        template_folder=str(PROJECT_ROOT / "templates"),
        static_folder=str(PROJECT_ROOT / "static"),
    )
    app.config.update(TESTING=True, SECRET_KEY="mpr-tests")
    app.add_url_rule("/", endpoint="main.index", view_func=lambda: "home")
    app.register_blueprint(mpr_bp)
    register_oplot_ui(app)
    return app


def upload(data: bytes, filename: str) -> FileStorage:
    return FileStorage(stream=io.BytesIO(data), filename=filename)


def workbook_bytes(rows: list[dict], sheet_name: str = "История лимитов") -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    sheet.append(mpr_service.REQUIRED_COLUMNS)
    for row in rows:
        sheet.append([row.get(column, "") for column in mpr_service.REQUIRED_COLUMNS])
    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def source_row(**values) -> dict:
    row = {
        "Имя": "host-01",
        "Наименование услуги": "Сервис",
        "Имя дата-центра ВМ": "МегаЦОД",
        "Имя AC": "AC-1",
        "ID КЭ сервера": "100",
        "Платформа": "linux",
        "Статус стенда": "Работает",
    }
    row.update(values)
    return row


class MprUiConfigTests(unittest.TestCase):
    def test_builder_uses_first_actual_template_and_endpoint_urls(self):
        app = build_app()
        templates = [{"code": "first"}, {"code": "second"}]
        with app.test_request_context("/mpr"):
            config = build_mpr_ui_config(templates=templates)
        self.assertEqual("first", config["initial_template_code"])
        self.assertEqual("/mpr/preview", config["urls"]["preview"])
        self.assertEqual("/mpr/generate", config["urls"]["generate"])
        source = (PROJECT_ROOT / "services" / "mpr_ui_service.py").read_text(encoding="utf-8")
        self.assertNotIn("os_update", source)
        self.assertNotIn('"/mpr/', source)

    def test_builder_empty_and_prefix_matrix(self):
        app = build_app()
        cases = (
            ({}, "/mpr/preview"),
            ({"headers": {"X-Forwarded-Prefix": "/proxy"}}, "/proxy/mpr/preview"),
            ({"environ_overrides": {"SCRIPT_NAME": "/script"}}, "/script/mpr/preview"),
        )
        with mock.patch.dict(os.environ, {"BASE_PATH": "", "PUBLIC_BASE_PATH": "", "APP_BASE_PATH": "", "APPLICATION_ROOT": ""}, clear=False):
            for options, expected in cases:
                with self.subTest(expected=expected), app.test_request_context("/mpr", **options):
                    self.assertEqual(expected, build_mpr_ui_config(templates=[])["urls"]["preview"])
                    self.assertEqual("", build_mpr_ui_config(templates=[])["initial_template_code"])
            with mock.patch.dict(os.environ, {"BASE_PATH": "/base"}, clear=False):
                with app.test_request_context("/mpr"):
                    self.assertEqual("/base/mpr/preview", build_mpr_ui_config(templates=[])["urls"]["preview"])
                with app.test_request_context("/mpr", environ_overrides={"SCRIPT_NAME": "/base"}):
                    self.assertEqual("/base/mpr/preview", build_mpr_ui_config(templates=[])["urls"]["preview"])


class MprTemplateContractTests(unittest.TestCase):
    def _render(self, templates):
        app = build_app()
        with mock.patch("routes.mpr_routes.list_mpr_templates", return_value=templates):
            return app.test_client().get("/mpr")

    def test_shell_assets_ids_and_single_template(self):
        response = self._render([{"code": "alpha", "name": "Шаблон Альфа", "filename": "alpha/template.docx"}])
        self.assertEqual(200, response.status_code)
        text = response.get_data(as_text=True)
        self.assertIn("oplot-shell oplot-shell--no-sidebar", text)
        self.assertIn("oplot-topbar--core", text)
        self.assertIn('class="oplot-body oplot-shell-mode-app oplot-topbar-variant-core oplot-mpr"', text)
        self.assertEqual(1, len(re.findall(r"<h1(?:\s|>)", text)))
        self.assertNotIn("oplot-breadcrumbs", text)
        self.assertIn('href="/"', text)
        self.assertIn("data-oplot-theme-toggle", text)
        self.assertIn("css/oplot_mpr.css", text)
        self.assertIn("js/oplot_mpr.js", text)
        self.assertIn("defer", text)
        self.assertIn("is-readonly", text)
        self.assertIn('value="alpha"', text)
        for element_id in (
            "mprTemplateList", "mprFiles", "mprFileList", "mprAlert", "mprStatus",
            "mprGenerateBtn", "mprPackageModal", "mprPackageSummary", "mprPackageOptions",
            "mprUnmapped", "mprConfirmGenerateBtn", "mprConfirmGenerateLabel", "mprResult",
            "mprResultFilename", "mprResultType", "mprResultTemplate", "mprResultSources",
            "mprResultRows", "mprResultPackages", "mprDownloadAgainBtn", "mprNewGenerationBtn",
        ):
            self.assertIn(f'id="{element_id}"', text)
        self.assertIn('name="files"', text)
        self.assertIn('accept=".xlsx,.csv"', text)
        self.assertIn("multiple", text)
        self.assertIn('name="template_code"', text)

    def test_multiple_and_empty_template_states(self):
        multiple = self._render([
            {"code": "first", "name": "Первый", "filename": "first/template.docx"},
            {"code": "second", "name": "Второй", "filename": "second/template.docx"},
        ]).get_data(as_text=True)
        self.assertEqual(2, multiple.count("data-template-code="))
        self.assertIn('"initial_template_code": "first"', multiple)
        empty = self._render([]).get_data(as_text=True)
        self.assertIn("Шаблон МПР недоступен", empty)
        self.assertRegex(empty, r'id="mprGenerateBtn"[^>]*disabled')
        self.assertIn('"initial_template_code": ""', empty)

    def test_legacy_shell_and_inline_business_code_are_absent(self):
        source = MPR_TEMPLATE.read_text(encoding="utf-8")
        self.assertIn("extends 'layouts/oplot_base.html'", source)
        self.assertIn("oplot_show_sidebar = false", source)
        self.assertIn("oplot_topbar_variant = 'core'", source)
        self.assertIn("oplot_show_breadcrumbs = false", source)
        self.assertIn("oplot_show_page_header = false", source)
        for forbidden in ("base_styles.html", "bootstrap.min.css", "bootstrap.bundle", "location.hostname", "BASE_PATH", "new bootstrap.Modal"):
            self.assertNotIn(forbidden, source)
        self.assertNotRegex(source, r"<style\b")
        executable = re.findall(r'<script(?![^>]*type="application/json")[^>]*>(.*?)</script>', source, re.S)
        self.assertFalse(any(block.strip() for block in executable))
        for foreign in ("document_templates.js", "oplot_release.js", "oplot_duty_dashboard.js", "docx-preview", "jszip", "htmx"):
            self.assertNotIn(foreign, source)

    def test_css_is_scoped_and_uses_dna_tokens(self):
        css = MPR_CSS.read_text(encoding="utf-8")
        self.assertRegex(css, r"\.oplot-mpr\s*\{[^}]*background:\s*transparent")
        self.assertIn("var(--oplot-dna-panel-gradient)", css)
        self.assertIn("prefers-reduced-motion", css)
        for line in css.splitlines():
            stripped = line.strip()
            if stripped.endswith("{") and not stripped.startswith((".oplot-mpr", "@media")):
                self.fail(f"Unscoped selector: {stripped}")
        self.assertNotRegex(css, r"(?m)^\s*(?::root|body|\[data-theme|\.card|\.btn|\.modal|\.alert|\.table|\.form-control)\b")


class MprJavascriptContractTests(unittest.TestCase):
    def test_config_init_modal_and_blob_lifecycle_contracts(self):
        script = MPR_JS.read_text(encoding="utf-8")
        self.assertIn('"use strict"', script)
        self.assertIn("function initOplotMprPage", script)
        self.assertIn('dataset.oplotMprInitialized === "true"', script)
        self.assertIn("JSON.parse", script)
        self.assertIn("isLocalUrl", script)
        self.assertIn("new URL(value, window.location.origin)", script)
        self.assertIn("window.tabler && window.tabler.Modal", script)
        self.assertIn("getOrCreateInstance", script)
        self.assertIn("getInstance", script)
        self.assertIn("new WeakMap", script)
        self.assertNotIn("bootstrap.Modal", script)
        self.assertNotIn("location.hostname", script)
        self.assertNotIn("BASE_PATH", script)
        self.assertNotIn("/mpr/preview", script)
        self.assertNotIn("/mpr/generate", script)
        self.assertNotIn("window.initOplotMprPage", script)
        self.assertIn('formData.append("template_code"', script)
        self.assertIn('formData.append("files"', script)
        self.assertIn('formData.append("packages"', script)
        self.assertIn("URL.createObjectURL", script)
        self.assertIn("URL.revokeObjectURL", script)
        self.assertIn('window.addEventListener("pagehide"', script)
        self.assertIn("triggerDownload(resultObjectUrl, resultFilename)", script)
        self.assertIn("resetGeneration", script)


class MprRouteContractTests(unittest.TestCase):
    def setUp(self):
        self.app = build_app()
        self.client = self.app.test_client()
        self.template = {"code": "synthetic", "name": "Синтетический", "filename": "synthetic/template.docx"}

    def test_preview_success_and_safe_errors(self):
        with mock.patch("routes.mpr_routes.resolve_mpr_template"), mock.patch("routes.mpr_routes.build_mpr_rows", return_value=[{"ЦОД": "МегаЦОД"}]), mock.patch("routes.mpr_routes.build_mpr_package_preview", return_value={"rows_count": 1, "packages": [], "unmapped": []}):
            response = self.client.post("/mpr/preview", data={"template_code": "synthetic", "files": (io.BytesIO(b"x"), "source.csv")})
        self.assertEqual(200, response.status_code)
        self.assertEqual(1, response.get_json()["rows_count"])
        with mock.patch("routes.mpr_routes.resolve_mpr_template", side_effect=MprError("Некорректный файл", ["detail"])):
            response = self.client.post("/mpr/preview", data={"template_code": "bad"})
        self.assertEqual(400, response.status_code)
        self.assertEqual({"success": False, "error": "Некорректный файл", "details": ["detail"]}, response.get_json())
        with mock.patch("routes.mpr_routes.resolve_mpr_template", side_effect=RuntimeError("private path")):
            response = self.client.post("/mpr/preview", data={"template_code": "bad"})
        self.assertEqual(500, response.status_code)
        self.assertEqual({"success": False, "error": "Не удалось проверить данные МПР"}, response.get_json())

    def test_single_docx_and_multiple_zip_download_contracts(self):
        common = (
            mock.patch("routes.mpr_routes.resolve_mpr_template", return_value=(Path("synthetic.docx"), self.template)),
            mock.patch("routes.mpr_routes.build_mpr_rows", return_value=[{"ЦОД": "МегаЦОД"}]),
            mock.patch("routes.mpr_routes.select_mpr_package_rows", return_value={"mcod": [{}], "scod_vavilova": [{}]}),
            mock.patch("routes.mpr_routes.generate_mpr_docx", return_value=io.BytesIO(b"docx")),
        )
        with common[0], common[1], common[2], common[3], mock.patch("routes.mpr_routes.normalize_mpr_package_codes", return_value=["mcod"]), mock.patch("routes.mpr_routes.build_output_filename", return_value="МПР_проверка.docx"):
            response = self.client.post("/mpr/generate", data={"template_code": "synthetic", "packages": "mcod", "files": (io.BytesIO(b"x"), "source.csv")})
        self.assertEqual(200, response.status_code)
        self.assertEqual("application/vnd.openxmlformats-officedocument.wordprocessingml.document", response.mimetype)
        self.assertIn("attachment", response.headers["Content-Disposition"])
        self.assertRegex(response.headers["Content-Disposition"], r"filename\*=UTF-8''")
        common = (
            mock.patch("routes.mpr_routes.resolve_mpr_template", return_value=(Path("synthetic.docx"), self.template)),
            mock.patch("routes.mpr_routes.build_mpr_rows", return_value=[{"ЦОД": "МегаЦОД"}]),
            mock.patch("routes.mpr_routes.select_mpr_package_rows", return_value={"mcod": [{}], "scod_vavilova": [{}]}),
            mock.patch("routes.mpr_routes.generate_mpr_docx", return_value=io.BytesIO(b"docx")),
        )
        with common[0], common[1], common[2], common[3], mock.patch("routes.mpr_routes.normalize_mpr_package_codes", return_value=["mcod", "scod_vavilova"]), mock.patch("routes.mpr_routes.build_output_filename", side_effect=["one.docx", "two.docx"]), mock.patch("routes.mpr_routes.build_archive_filename", return_value="МПР_пакет.zip"):
            response = self.client.post("/mpr/generate", data={"template_code": "synthetic", "packages": ["mcod", "scod_vavilova"], "files": (io.BytesIO(b"x"), "source.csv")})
        self.assertEqual(200, response.status_code)
        self.assertEqual("application/zip", response.mimetype)
        with ZipFile(io.BytesIO(response.data)) as archive:
            self.assertEqual(["one.docx", "two.docx"], archive.namelist())


class MprParsingAndGenerationCharacterizationTests(unittest.TestCase):
    def test_xlsx_csv_encodings_filter_mapping_deduplication_and_sorting(self):
        rows = [
            source_row(**{"Имя": "host-b", "Имя AC": "AC-2"}),
            source_row(**{"Имя": "host-a", "Имя AC": "AC-1"}),
            source_row(**{"Имя": "host-a", "Имя AC": "AC-1"}),
            source_row(**{"Имя": "ignored", "Статус стенда": "Остановлен"}),
        ]
        parsed = mpr_service.build_mpr_rows([upload(workbook_bytes(rows), "limits.xlsx")])
        self.assertEqual(["host-a", "host-b"], [item["КТС"] for item in parsed])
        self.assertEqual("Сервис", parsed[0]["Наименование"])
        self.assertEqual("МегаЦОД", parsed[0]["ЦОД"])
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=mpr_service.REQUIRED_COLUMNS, delimiter=";")
        writer.writeheader(); writer.writerow(source_row())
        for encoding in ("utf-8-sig", "cp1251"):
            with self.subTest(encoding=encoding):
                parsed = mpr_service.build_mpr_rows([upload(output.getvalue().encode(encoding), "limits.csv")])
                self.assertEqual("host-01", parsed[0]["КТС"])

    def test_missing_sheet_columns_unknown_datacenter_and_packages(self):
        with self.assertRaisesRegex(MprError, "Не удалось обработать") as missing_sheet:
            mpr_service.build_mpr_rows([upload(workbook_bytes([source_row()], "Другой лист"), "bad.xlsx")])
        self.assertTrue(any("История лимитов" in detail for detail in missing_sheet.exception.details))
        stream = io.StringIO(); writer = csv.DictWriter(stream, fieldnames=["Имя"], delimiter=";"); writer.writeheader(); writer.writerow({"Имя": "host"})
        with self.assertRaisesRegex(MprError, "Не удалось обработать") as missing_columns:
            mpr_service.build_mpr_rows([upload(stream.getvalue().encode("utf-8"), "bad.csv")])
        self.assertTrue(any("обязательные колонки" in detail for detail in missing_columns.exception.details))
        rows = [{"КТС": "host", "Наименование": "svc", "ЦОД": "Неизвестный", "Имя AC": "AC", "ID КЭ сервера": "1", "Платформа": "linux"}]
        preview = mpr_service.build_mpr_package_preview(rows)
        self.assertEqual([{"datacenter": "Неизвестный", "rows_count": 1}], preview["unmapped"])
        with self.assertRaisesRegex(MprError, "нераспределенные"):
            mpr_service.select_mpr_package_rows(rows, ["mcod"])
        self.assertEqual(["mcod", "scod_vavilova"], mpr_service.normalize_mpr_package_codes([]))

    def _template(self, path: Path, include_first=True, include_second=True, include_heading=False):
        document = Document()
        document.add_paragraph("Сохраняемый текст")
        if include_heading: document.add_heading("ПРИЛОЖЕНИЕ 1", level=1)
        if include_first: document.add_paragraph("{{APPENDIX_1_TABLE}}")
        if include_second: document.add_paragraph("{{APPENDIX_2_TABLE}}")
        document.save(path)

    def test_exact_placeholders_existing_content_and_third_appendix(self):
        rows = [{"КТС": "host", "Наименование": "svc", "ЦОД": "Сколково", "Имя AC": "AC", "ID КЭ сервера": "1", "Платформа": "linux"}]
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "template.docx"
            self._template(source)
            generated = mpr_service.generate_mpr_docx(source, rows, package_code="scod_vavilova")
            result_path = Path(directory) / "result.docx"; result_path.write_bytes(generated.getvalue())
            document = Document(result_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Сохраняемый текст", text)
            self.assertIn("Приложение 3 - СЦОД", text)
            self.assertGreaterEqual(len(document.tables), 3)

    def test_heading_is_not_a_placeholder_fallback(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "template.docx"
            self._template(source, include_first=False, include_second=True, include_heading=True)
            with self.assertRaisesRegex(MprError, r"APPENDIX_1_TABLE"):
                mpr_service.generate_mpr_docx(source, [source_row()], package_code="mcod")


if __name__ == "__main__":
    unittest.main()
