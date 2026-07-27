import csv
import re
from copy import deepcopy
from datetime import datetime
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook

from config import MPR_TEMPLATES_ROOT


APPENDIX_PLACEHOLDER = "{{APPENDIX_1_TABLE}}"
APPENDIX_2_PLACEHOLDER = "{{APPENDIX_2_TABLE}}"
LOCATION_PLACEHOLDER = "{{MPR_LOCATION}}"
APPENDIX_1_SUFFIX_PLACEHOLDER = "{{APPENDIX_1_SUFFIX}}"
APPENDIX_2_SUFFIX_PLACEHOLDER = "{{APPENDIX_2_SUFFIX}}"
MCOD_PACKAGE_CODE = "mcod"
SCOD_VAVILOVA_PACKAGE_CODE = "scod_vavilova"
MCOD_DATACENTERS = ("МегаЦОД",)
SCOD_DATACENTERS = ("Сколково",)
VAVILOVA_DATACENTERS = ("Вавилова", "Вавилова (observer)")
DB_HOST_KEYWORDS = ("postgres", "postgre", "_db_", "pangolin")
LOADBALANCER_PLATFORM = "loadbalancer_osse_20_2l"
SCOD_EXCLUDED_APP_2_PLATFORMS = ("ignite_se",)
MCOD_EXCLUDED_APP_2_SERVICES = (
    "clm_const_pprb_lb_o_prom",
    "focus_admin_lboss_o_prom",
)
MCOD_EXCLUDED_APP_2_HOSTS = (
    "pslsb-efs002173",
    "pslsb-efs002176",
    "pslsb-efs002178",
)
SCOD_EXCLUDED_APP_3_SERVICES = ("aef_ai_lbosse_o_prom",)
HOST_PLACEHOLDERS = {
    "{{MPR_SOWA_HOSTS}}": ("sowa",),
    "{{MPR_POSTGRES_HOSTS}}": DB_HOST_KEYWORDS,
    "{{MPR_SYNGX_HOSTS}}": ("syngx", "syng"),
}
MPR_PACKAGES = {
    MCOD_PACKAGE_CODE: {
        "label": "МЦОД",
        "datacenters": MCOD_DATACENTERS,
    },
    SCOD_VAVILOVA_PACKAGE_CODE: {
        "label": "СЦОД и Вавилова",
        "datacenters": (*SCOD_DATACENTERS, *VAVILOVA_DATACENTERS),
    },
}
MPR_TEMPLATE_FILENAME = "template.docx"
MPR_TEMPLATE_NAMES = {
    "os_update": "Обновление ОС",
}

REQUIRED_COLUMNS = [
    "Имя",
    "Наименование услуги",
    "Имя дата-центра ВМ",
    "Имя AC",
    "ID КЭ сервера",
    "Платформа",
    "Статус стенда",
]

SOURCE_COLUMNS = [
    "Имя",
    "Наименование услуги",
    "Имя дата-центра ВМ",
    "Имя AC",
    "ID КЭ сервера",
    "Платформа",
]

DOCX_HEADERS = ["№", "Имя AC", "Наименование услуги", "Имя", "ЦОД"]


class MprError(Exception):
    """User-facing MPR validation/generation error."""

    def __init__(self, message, details=None):
        super().__init__(message)
        self.message = message
        self.details = details or []


def _normalize_header(value):
    return str(value or "").replace("\ufeff", "").strip()


def _normalize_cell(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _safe_template_code(value):
    code = (value or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]+", code):
        return ""
    return code


def list_mpr_templates():
    MPR_TEMPLATES_ROOT.mkdir(parents=True, exist_ok=True)
    templates = []
    for folder in sorted(MPR_TEMPLATES_ROOT.iterdir()):
        if not folder.is_dir():
            continue
        path = folder / MPR_TEMPLATE_FILENAME
        if not path.exists() or path.name.startswith("~$"):
            continue
        code = folder.name
        templates.append({
            "code": code,
            "name": MPR_TEMPLATE_NAMES.get(code, code.replace("_", " ").strip() or code),
            "filename": f"{code}/{MPR_TEMPLATE_FILENAME}",
        })
    return templates


def resolve_mpr_template(template_code):
    code = _safe_template_code(template_code)
    if not code:
        raise MprError("Не выбран шаблон МПР")

    candidates = {item["code"]: item for item in list_mpr_templates()}
    if code not in candidates:
        raise MprError("Шаблон DOCX не найден")

    path = (MPR_TEMPLATES_ROOT / code / MPR_TEMPLATE_FILENAME).resolve()
    root = MPR_TEMPLATES_ROOT.resolve()
    if root not in path.parents or not path.exists():
        raise MprError("Шаблон DOCX не найден")
    return path, candidates[code]


def build_mpr_rows(files):
    if not files:
        raise MprError("Не загружены файлы")

    rows = []
    errors = []

    for uploaded in files:
        filename = uploaded.filename or "без имени"
        suffix = Path(filename).suffix.lower()
        if suffix not in {".xlsx", ".csv"}:
            errors.append(f"{filename}: неподдерживаемый формат файла")
            continue

        try:
            file_rows = _read_xlsx(uploaded, filename) if suffix == ".xlsx" else _read_csv(uploaded, filename)
            rows.extend(file_rows)
        except MprError as exc:
            errors.append(exc.message)

    if errors:
        raise MprError("Не удалось обработать загруженные файлы", errors)

    filtered = []
    for row in rows:
        if row.get("Статус стенда", "").strip().casefold() != "работает":
            continue
        filtered.append({column: row.get(column, "") for column in SOURCE_COLUMNS})

    if not filtered:
        raise MprError("После фильтрации нет строк со статусом «Работает»")

    deduped = []
    seen = set()
    for row in filtered:
        marker = tuple(row.get(column, "") for column in SOURCE_COLUMNS)
        if marker in seen:
            continue
        seen.add(marker)
        deduped.append({
            "КТС": row.get("Имя", ""),
            "Наименование": row.get("Наименование услуги", ""),
            "ЦОД": row.get("Имя дата-центра ВМ", ""),
            "Имя AC": row.get("Имя AC", ""),
            "ID КЭ сервера": row.get("ID КЭ сервера", ""),
            "Платформа": row.get("Платформа", ""),
        })

    return sorted(
        deduped,
        key=lambda item: (
            item.get("Имя AC", "").casefold(),
            item.get("Наименование", "").casefold(),
            item.get("КТС", "").casefold(),
        ),
    )


def _read_xlsx(uploaded, filename):
    data = BytesIO(uploaded.read())
    try:
        workbook = load_workbook(data, read_only=True, data_only=True)
    except Exception as exc:
        raise MprError(f"{filename}: не удалось прочитать XLSX") from exc

    if "История лимитов" not in workbook.sheetnames:
        raise MprError(f"{filename}: отсутствует лист «История лимитов»")

    sheet = workbook["История лимитов"]
    rows_iter = sheet.iter_rows(values_only=True)
    try:
        headers = [_normalize_header(value) for value in next(rows_iter)]
    except StopIteration as exc:
        raise MprError(f"{filename}: лист «История лимитов» пуст") from exc

    _validate_columns(headers, filename)
    indexes = {header: headers.index(header) for header in REQUIRED_COLUMNS}
    result = []
    for values in rows_iter:
        row = {}
        for column in REQUIRED_COLUMNS:
            index = indexes[column]
            row[column] = _normalize_cell(values[index] if index < len(values) else "")
        result.append(row)
    return result


def _read_csv(uploaded, filename):
    raw = uploaded.read()
    text = None
    for encoding in ("utf-8-sig", "cp1251"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise MprError(f"{filename}: не удалось определить кодировку CSV")

    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=";,|\t,")
    except csv.Error:
        dialect = csv.excel
        dialect.delimiter = ";"

    reader = csv.DictReader(StringIO(text), dialect=dialect)
    headers = [_normalize_header(header) for header in (reader.fieldnames or [])]
    _validate_columns(headers, filename)

    result = []
    for raw_row in reader:
        normalized = {_normalize_header(key): _normalize_cell(value) for key, value in raw_row.items()}
        result.append({column: normalized.get(column, "") for column in REQUIRED_COLUMNS})
    return result


def _validate_columns(headers, filename):
    present = set(headers)
    missing = [column for column in REQUIRED_COLUMNS if column not in present]
    if missing:
        raise MprError(f"{filename}: отсутствуют обязательные колонки: {', '.join(missing)}")


def normalize_mpr_package_codes(values):
    requested = [str(value or "").strip() for value in (values or []) if str(value or "").strip()]
    if not requested:
        return list(MPR_PACKAGES)

    result = []
    unknown = []
    for code in requested:
        if code not in MPR_PACKAGES:
            unknown.append(code)
            continue
        if code not in result:
            result.append(code)
    if unknown:
        raise MprError(f"Неизвестные пакеты МПР: {', '.join(unknown)}")
    if not result:
        raise MprError("Не выбран ни один документ МПР")
    return result


def build_mpr_package_preview(rows):
    grouped, unmapped = _split_mpr_rows(rows)
    packages = []
    for code, config in MPR_PACKAGES.items():
        packages.append({
            "code": code,
            "label": config["label"],
            "datacenters": list(config["datacenters"]),
            "rows_count": len(grouped[code]),
            "available": bool(grouped[code]),
        })
    return {
        "rows_count": len(rows),
        "packages": packages,
        "unmapped": [
            {"datacenter": datacenter, "rows_count": count}
            for datacenter, count in unmapped.items()
        ],
    }


def select_mpr_package_rows(rows, package_codes):
    codes = normalize_mpr_package_codes(package_codes)
    grouped, unmapped = _split_mpr_rows(rows)
    if unmapped:
        details = [
            f"{datacenter}: {count} строк"
            for datacenter, count in unmapped.items()
        ]
        raise MprError("В загруженных данных есть нераспределенные значения ЦОД", details)

    empty = [MPR_PACKAGES[code]["label"] for code in codes if not grouped[code]]
    if empty:
        raise MprError(f"Для выбранных пакетов не найдены хосты: {', '.join(empty)}")
    return {code: grouped[code] for code in codes}


def generate_mpr_docx(template_path, rows, location_label=None, package_code=None):
    try:
        document = Document(template_path)
    except Exception as exc:
        raise MprError("Не удалось открыть DOCX-шаблон") from exc

    package_code = _normalize_generated_package_code(package_code, location_label)
    appendix_plan = _build_appendix_plan(rows, package_code)
    _apply_work_plan_notes(document, package_code)

    replacements = _build_host_placeholder_values(rows)
    replacements[APPENDIX_1_SUFFIX_PLACEHOLDER] = (
        appendix_plan["suffixes"].get(APPENDIX_PLACEHOLDER, "")
    )
    replacements[APPENDIX_2_SUFFIX_PLACEHOLDER] = (
        appendix_plan["suffixes"].get(APPENDIX_2_PLACEHOLDER, "")
    )
    if location_label is not None:
        if not _document_contains_placeholder(document, LOCATION_PLACEHOLDER):
            raise MprError(f"Плейсхолдер {LOCATION_PLACEHOLDER} не найден в DOCX-шаблоне")
        replacements[LOCATION_PLACEHOLDER] = location_label
    _replace_host_placeholders(document, replacements)

    first_table = _insert_appendix_table(
        document,
        APPENDIX_PLACEHOLDER,
        appendix_plan["tables"][0]["rows"],
        required=True,
    )
    second_table = _insert_appendix_table(
        document,
        APPENDIX_2_PLACEHOLDER,
        appendix_plan["tables"][1]["rows"],
        required=True,
    )
    if len(appendix_plan["tables"]) > 2:
        _insert_appendix_after_table(
            document,
            second_table or first_table,
            appendix_plan["tables"][2]["title"],
            appendix_plan["tables"][2]["rows"],
        )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def build_output_filename(template_info, package_label=None, timestamp=None):
    name = template_info.get("name") or template_info.get("code") or "МПР"
    safe_name = re.sub(r'[<>:"/\\|?*]+', " ", name).strip() or "МПР"
    if package_label:
        safe_package = re.sub(r'[<>:"/\\|?*]+', " ", package_label).strip()
        if safe_package:
            safe_name = f"{safe_name}_{safe_package}"
    timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"МПР_{safe_name}_{timestamp}.docx"


def build_archive_filename(template_info, timestamp=None):
    name = template_info.get("name") or template_info.get("code") or "МПР"
    safe_name = re.sub(r'[<>:"/\\|?*]+', " ", name).strip() or "МПР"
    timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"МПР_{safe_name}_{timestamp}.zip"


def _split_mpr_rows(rows):
    datacenter_map = {}
    for code, config in MPR_PACKAGES.items():
        for datacenter in config["datacenters"]:
            datacenter_map[_normalize_datacenter(datacenter)] = code

    grouped = {code: [] for code in MPR_PACKAGES}
    unmapped = {}
    for item in rows:
        raw_datacenter = str(item.get("ЦОД", "") or "").strip()
        code = datacenter_map.get(_normalize_datacenter(raw_datacenter))
        if code:
            grouped[code].append(item)
            continue
        label = raw_datacenter or "Пустое значение ЦОД"
        unmapped[label] = unmapped.get(label, 0) + 1
    return grouped, unmapped


def _normalize_datacenter(value):
    return re.sub(r"\s+", " ", str(value or "").strip()).casefold()


def _normalize_generated_package_code(package_code, location_label):
    code = str(package_code or "").strip()
    if code in MPR_PACKAGES:
        return code
    label = str(location_label or "").strip()
    for candidate, config in MPR_PACKAGES.items():
        if label == config["label"]:
            return candidate
    return code


def _build_appendix_plan(rows, package_code):
    if package_code == MCOD_PACKAGE_CODE:
        return {
            "suffixes": {
                APPENDIX_PLACEHOLDER: "",
                APPENDIX_2_PLACEHOLDER: "",
            },
            "tables": [
                {
                    "title": "ПРИЛОЖЕНИЕ 1",
                    "rows": _filter_mpr_rows(
                        rows,
                        datacenters=MCOD_DATACENTERS,
                        exclude_platforms=(LOADBALANCER_PLATFORM,),
                    ),
                },
                {
                    "title": "ПРИЛОЖЕНИЕ 2",
                    "rows": _filter_mpr_rows(
                        rows,
                        datacenters=MCOD_DATACENTERS,
                        platforms=(LOADBALANCER_PLATFORM,),
                        exclude_services=MCOD_EXCLUDED_APP_2_SERVICES,
                        exclude_hosts=MCOD_EXCLUDED_APP_2_HOSTS,
                    ),
                },
            ],
        }

    if package_code == SCOD_VAVILOVA_PACKAGE_CODE:
        return {
            "suffixes": {
                APPENDIX_PLACEHOLDER: " — Вавилова (observer)",
                APPENDIX_2_PLACEHOLDER: " — СЦОД",
            },
            "tables": [
                {
                    "title": "ПРИЛОЖЕНИЕ 1 — Вавилова (observer)",
                    "rows": _filter_mpr_rows(
                        rows,
                        datacenters=VAVILOVA_DATACENTERS,
                        keywords=DB_HOST_KEYWORDS,
                    ),
                },
                {
                    "title": "ПРИЛОЖЕНИЕ 2 — СЦОД",
                    "rows": _filter_mpr_rows(
                        rows,
                        datacenters=SCOD_DATACENTERS,
                        exclude_platforms=(
                            LOADBALANCER_PLATFORM,
                            *SCOD_EXCLUDED_APP_2_PLATFORMS,
                        ),
                    ),
                },
                {
                    "title": "Приложение 3 - СЦОД",
                    "rows": _filter_mpr_rows(
                        rows,
                        datacenters=SCOD_DATACENTERS,
                        platforms=(LOADBALANCER_PLATFORM,),
                        exclude_services=SCOD_EXCLUDED_APP_3_SERVICES,
                    ),
                },
            ],
        }

    return {
        "suffixes": {
            APPENDIX_PLACEHOLDER: "",
            APPENDIX_2_PLACEHOLDER: "",
        },
        "tables": [
            {"title": "ПРИЛОЖЕНИЕ 1", "rows": rows},
            {"title": "ПРИЛОЖЕНИЕ 2", "rows": []},
        ],
    }


def _build_host_placeholder_values(rows):
    values = {}
    for placeholder, keywords in HOST_PLACEHOLDERS.items():
        hosts = []
        seen = set()
        for item in rows:
            name = str(item.get("КТС", "") or "").strip()
            if not name or not _row_matches_keywords(item, keywords):
                continue
            marker = name.casefold()
            if marker in seen:
                continue
            seen.add(marker)
            hosts.append(name)
        values[placeholder] = "\n".join(hosts) if hosts else "—"
    return values


def _apply_work_plan_notes(document, package_code):
    if not document.tables:
        return
    table = document.tables[0]
    if len(table.rows) < 10 or len(table.columns) < 4:
        return

    if package_code == SCOD_VAVILOVA_PACKAGE_CODE:
        all_appendices_note = "Список КТС взять из приложения 1 и приложение 2 и приложение 3"
        _set_table_cell_text(table, 2, 3, all_appendices_note)
        _set_table_cell_text(table, 5, 3, "Список КТС взять из приложения 1 ")
        _insert_scod_update_row(table, source_row_index=5)
        _set_table_cell_text(table, 10, 3, f"\n{all_appendices_note}")
        return

    if package_code == MCOD_PACKAGE_CODE:
        appendices_note = "Список КТС взять из приложения 1 и приложения 2"
        _set_table_cell_text(table, 2, 3, appendices_note)
        _set_table_cell_text(table, 5, 3, appendices_note)
        _set_table_cell_text(table, 9, 3, f"\n{appendices_note}")


def _insert_scod_update_row(table, source_row_index):
    source_row = table.rows[source_row_index]
    new_row = deepcopy(source_row._tr)
    source_row._tr.addnext(new_row)
    _set_table_cell_text(table, source_row_index + 1, 3, "Список КТС взять из приложения 2 и приложения 3")


def _filter_mpr_rows(
    rows,
    datacenters=None,
    keywords=None,
    platforms=None,
    exclude_platforms=None,
    exclude_services=None,
    exclude_hosts=None,
):
    datacenter_markers = {
        _normalize_datacenter(value)
        for value in (datacenters or [])
        if str(value or "").strip()
    }
    platform_markers = {
        _normalize_platform(value)
        for value in (platforms or [])
        if str(value or "").strip()
    }
    excluded_platform_markers = {
        _normalize_platform(value)
        for value in (exclude_platforms or [])
        if str(value or "").strip()
    }
    excluded_service_markers = {
        _normalize_service(value)
        for value in (exclude_services or [])
        if str(value or "").strip()
    }
    excluded_host_markers = {
        _normalize_host(value)
        for value in (exclude_hosts or [])
        if str(value or "").strip()
    }
    result = []
    for item in rows:
        if datacenter_markers:
            datacenter = _normalize_datacenter(item.get("ЦОД", ""))
            if datacenter not in datacenter_markers:
                continue
        platform = _normalize_platform(item.get("Платформа", ""))
        if platform_markers and platform not in platform_markers:
            continue
        if excluded_platform_markers and platform in excluded_platform_markers:
            continue
        if _normalize_service(item.get("Наименование", "")) in excluded_service_markers:
            continue
        if _normalize_host(item.get("КТС", "")) in excluded_host_markers:
            continue
        if keywords and not _row_matches_keywords(item, keywords):
            continue
        result.append(item)
    return result


def _normalize_platform(value):
    return str(value or "").strip().casefold()


def _normalize_service(value):
    return str(value or "").strip().casefold()


def _normalize_host(value):
    return str(value or "").strip().casefold()


def _row_matches_keywords(item, keywords):
    haystack = f"{item.get('Наименование', '')} {item.get('КТС', '')}".casefold()
    return any(str(keyword or "").casefold() in haystack for keyword in keywords)


def _replace_host_placeholders(document, values):
    for paragraph in _iter_document_paragraphs(document):
        if not paragraph.text:
            continue
        needs_fallback = False
        for placeholder, value in values.items():
            if placeholder not in paragraph.text:
                continue
            matching_runs = [run for run in paragraph.runs if placeholder in run.text]
            if matching_runs:
                for run in matching_runs:
                    _replace_run_placeholder(run, placeholder, value)
            else:
                needs_fallback = True
                break
        if needs_fallback:
            updated = paragraph.text
            for placeholder, value in values.items():
                updated = updated.replace(placeholder, value)
            _replace_paragraph_multiline(paragraph, updated)


def _document_contains_placeholder(document, placeholder):
    return any(placeholder in paragraph.text for paragraph in _iter_document_paragraphs(document))


def _iter_document_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _replace_paragraph_multiline(paragraph, text):
    source_run = paragraph.runs[0] if paragraph.runs else None
    for run in paragraph.runs:
        run.text = ""

    lines = str(text or "").splitlines() or [""]
    target_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    _copy_run_style(source_run, target_run)
    target_run.text = lines[0]
    for line in lines[1:]:
        target_run.add_break()
        target_run.add_text(line)


def _replace_run_placeholder(run, placeholder, value):
    before, after = run.text.split(placeholder, 1)
    lines = str(value or "").splitlines() or [""]
    run.text = before + lines[0]
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)
    run.add_text(after)


def _copy_run_style(source, target):
    if source is None:
        return
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    if source.font.color and source.font.color.rgb:
        target.font.color.rgb = source.font.color.rgb


def _find_placeholder_paragraph(document, placeholder=APPENDIX_PLACEHOLDER):
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def _insert_appendix_table(document, placeholder, rows, required=True):
    paragraph = _find_placeholder_paragraph(document, placeholder)
    if paragraph is None:
        if required:
            raise MprError(f"Плейсхолдер {placeholder} не найден в DOCX-шаблоне")
        return False
    table = _build_appendix_table(document, rows)
    _insert_table_at_placeholder(paragraph, table, placeholder=placeholder)
    return table


def _remove_optional_appendix_section(document, placeholder):
    paragraphs = list(document.paragraphs)
    target_index = None
    for index, paragraph in enumerate(paragraphs):
        if placeholder in paragraph.text:
            target_index = index
            break
    if target_index is None:
        return

    indexes_to_remove = {target_index}
    for index in range(target_index - 1, -1, -1):
        text = paragraphs[index].text.strip()
        if not text:
            indexes_to_remove.add(index)
            continue
        if text.startswith("ПРИЛОЖЕНИЕ 2"):
            indexes_to_remove.add(index)
        break

    for index in sorted(indexes_to_remove, reverse=True):
        element = paragraphs[index]._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _build_appendix_table(document, rows):
    table = document.add_table(rows=1, cols=len(DOCX_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    try:
        normal_style = document.styles["Normal"]
    except Exception:
        normal_style = None
    font_name = normal_style.font.name if normal_style and normal_style.font.name else None
    font_size = normal_style.font.size if normal_style and normal_style.font.size else Pt(9)

    header_cells = table.rows[0].cells
    for index, header in enumerate(DOCX_HEADERS):
        _set_cell_text(header_cells[index], header, bold=True, font_name=font_name, font_size=font_size)
        _shade_cell(header_cells[index], "EAF2F8")

    for number, item in enumerate(rows, start=1):
        cells = table.add_row().cells
        values = [
            str(number),
            item.get("Имя AC", ""),
            item.get("Наименование", ""),
            item.get("КТС", ""),
            item.get("ЦОД", ""),
        ]
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value, font_name=font_name, font_size=font_size)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=80, start=100, bottom=80, end=100)

    _apply_table_geometry(document, table)
    _repeat_header_row(table.rows[0])
    return table


def _insert_table_at_placeholder(paragraph, table, placeholder=APPENDIX_PLACEHOLDER):
    remaining_text = paragraph.text.replace(placeholder, "").strip()
    paragraph._p.addnext(table._tbl)
    if remaining_text:
        _replace_paragraph_text(paragraph, remaining_text)
    else:
        parent = paragraph._element.getparent()
        parent.remove(paragraph._element)


def _insert_appendix_after_table(document, previous_table, title, rows):
    if previous_table is None:
        return
    paragraph = document.add_paragraph(str(title or ""))
    table = _build_appendix_table(document, rows)
    previous_table._tbl.addnext(paragraph._p)
    paragraph._p.addnext(table._tbl)


def _replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _set_table_cell_text(table, row_index, column_index, text):
    if row_index >= len(table.rows) or column_index >= len(table.columns):
        return
    _replace_cell_text(table.rows[row_index].cells[column_index], text)


def _replace_cell_text(cell, text):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = str(text or "")
    else:
        paragraph.add_run(str(text or ""))


def _set_cell_text(cell, text, bold=False, font_name=None, font_size=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size


def _shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _apply_table_geometry(document, table):
    section = document.sections[0]
    total_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    weights = [0.07, 0.24, 0.34, 0.22, 0.13]
    widths = [int(total_width * weight) for weight in weights]
    widths[-1] = int(total_width) - sum(widths[:-1])

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(total_width)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        table._tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    table._tbl.insert(table._tbl.index(tbl_pr) + 1, tbl_grid)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def _repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
