import re
import logging
from datetime import datetime, timedelta
from docx import Document
from docx.oxml import OxmlElement, parse_xml  # ДОБАВЛЕН parse_xml
from docx.oxml.ns import qn
from config import DEFAULT_BH_PLAYBOOKS, OPLOT_VALUES
from utils.common import normalize_text
from services.jira_service import get_jira_domain_and_token


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def replace_keys_in_doc(doc, context, issues, release_id, instruction_url=None):
    def replace_in_paragraph(paragraph):
        runs = list(paragraph.runs)
        text = ''.join(run.text for run in runs)
        if not any(key in text for key in context):
            return
        
        keys_positions = []
        for key in context:
            start = 0
            while start < len(text):
                pos = text.find(key, start)
                if pos == -1:
                    break
                keys_positions.append((pos, pos + len(key), key))
                start = pos + 1
        
        if not keys_positions:
            return
        
        keys_positions.sort(key=lambda x: x[0])
        new_runs = []
        current_pos = 0
        
        for start, end, key in keys_positions:
            if current_pos < start:
                pre_text = text[current_pos:start]
                if new_runs:
                    new_runs[-1].text += pre_text
                else:
                    new_run = paragraph.add_run(pre_text)
                    new_runs.append(new_run)
            
            replacement = context[key]
            new_run = paragraph.add_run(replacement)
            
            for run in runs:
                if run.text and key in run.text:
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.font.highlight_color = run.font.highlight_color
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    break
            
            new_runs.append(new_run)
            current_pos = end
        
        if current_pos < len(text):
            post_text = text[current_pos:]
            if new_runs:
                new_runs[-1].text += post_text
            else:
                paragraph.add_run(post_text)
        
        for run in runs:
            p = run._element
            p.getparent().remove(p)
    
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Обработка таблицы задач
    table_found = False
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "ЗНИ/JIRA ID" in headers:
            table_found = True
            # Удаляем старые строки (кроме заголовка)
            for _ in range(len(table.rows) - 1, 0, -1):
                table._tbl.remove(table.rows[_]._tr)
            
            # Добавляем задачи из Jira
            for idx, issue in enumerate(issues, start=1):
                row = table.add_row().cells
                row[0].text = str(idx)
                p = row[1].paragraphs[0]
                add_hyperlink(p, f"{get_jira_domain_and_token(release_id)[0]}/browse/{issue['key']}", issue['key'])
                row[2].text = issue['summary']
                row[3].text = issue['type']
            
            # Добавляем границы таблице
            tbl = table._tbl
            if tbl.tblPr is None:
                tbl.tblPr = OxmlElement('w:tblPr')
            tblPr = tbl.tblPr
            tblBorders = parse_xml('''
                <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                </w:tblBorders>
            ''')
            tblPr.append(tblBorders)
    
    # Обработка инструкции
    if instruction_url:
        for para in doc.paragraphs:
            full_text = para.text
            if "ИНСТРУКЦИЯ" in full_text:
                runs_to_remove = []
                for run in para.runs:
                    if "ИНСТРУКЦИЯ" in run.text:
                        run.text = run.text.replace("ИНСТРУКЦИЯ", "")
                        runs_to_remove.append(run)
                add_hyperlink(para, instruction_url, "инструкция")
                if runs_to_remove:
                    original_run = runs_to_remove[0]
                    hyperlink_run = para.runs[-1]
                    hyperlink_run.font.color.rgb = original_run.font.color.rgb
                    hyperlink_run.font.bold = original_run.font.bold
                    hyperlink_run.font.italic = original_run.font.italic
                    hyperlink_run.font.underline = original_run.font.underline
                    hyperlink_run.font.highlight_color = original_run.font.highlight_color
                break
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        if "ИНСТРУКЦИЯ" in full_text:
                            runs_to_remove = []
                            for run in paragraph.runs:
                                if "ИНСТРУКЦИЯ" in run.text:
                                    run.text = run.text.replace("ИНСТРУКЦИЯ", "")
                                    runs_to_remove.append(run)
                            add_hyperlink(paragraph, instruction_url, "инструкция")
                            if runs_to_remove:
                                original_run = runs_to_remove[0]
                                hyperlink_run = paragraph.runs[-1]
                                hyperlink_run.font.color.rgb = original_run.font.color.rgb
                                hyperlink_run.font.bold = original_run.font.bold
                                hyperlink_run.font.italic = original_run.font.italic
                                hyperlink_run.font.underline = original_run.font.underline
                                hyperlink_run.font.highlight_color = original_run.font.highlight_color
                            break
    
    return doc

def check_document(document_path, context, issues, release_id):
    errors = []
    try:
        doc = Document(document_path)
        mandatory_keys = [
            "RELEASE_VERSION", "PREV_VERSION", "RELEASE_ID",
            "OPLOT", "CHECKER", "DATE", "PLUS_1", "INSTRUCTION_BLOCK",
            "POB", "RELNUMBER"
        ]
        
        playbooks_should_be_checked = False
        playbooks_key_found = False
        
        for key in mandatory_keys:
            skip_check = (key == "DATE" and "PLAYBOOKS" in context and
                          any(pb in context["PLAYBOOKS"] for pb in DEFAULT_BH_PLAYBOOKS))
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if skip_check and key in text:
                    if not any(pb in text for pb in DEFAULT_BH_PLAYBOOKS if key in pb):
                        errors.append(f"Найден шаблонный ключ '{key}' в тексте")
                elif key in text:
                    errors.append(f"Найден шаблонный ключ '{key}' в тексте")
                
                if key == "PLAYBOOKS" and "PLAYBOOKS" in text:
                    playbooks_key_found = True
            
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells] if len(table.rows) > 0 else []
                is_issues_table = "ЗНИ/JIRA ID" in headers
                
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            if is_issues_table and row_idx > 0:
                                continue
                            
                            if skip_check and key in text:
                                if not any(pb in text for pb in DEFAULT_BH_PLAYBOOKS if key in pb):
                                    errors.append(f"Найден шаблонный ключ '{key}' в таблице")
                            elif key in text:
                                errors.append(f"Найден шаблонный ключ '{key}' в таблице")
                            
                            if key == "PLAYBOOKS" and "PLAYBOOKS" in text:
                                playbooks_key_found = True
        
        if playbooks_key_found:
            if "SOWA" not in document_path.name.upper() and context["PLAYBOOKS"] == "":
                errors.append("PLAYBOOKS: не выбраны плейбуки")
            playbooks_should_be_checked = True
        
        # Проверка таблицы задач
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if "ЗНИ/JIRA ID" in headers:
                if len(table.rows) < 2:
                    errors.append("Таблица задач пуста")
                    continue
                
                if issues:
                    table_issues = []
                    for row in table.rows[1:]:
                        cells = row.cells
                        if len(cells) >= 4:
                            issue_key = cells[1].text.strip()
                            issue_summary = cells[2].text.strip()
                            issue_type = cells[3].text.strip()
                            table_issues.append({"key": issue_key, "summary": issue_summary, "type": issue_type})
                    
                    if len(table_issues) != len(issues):
                        errors.append(f"Количество задач в таблице ({len(table_issues)}) не совпадает с Jira ({len(issues)})")
                    
                    jira_issues_map = {issue['key']: issue for issue in issues}
                    for table_issue in table_issues:
                        jira_issue = jira_issues_map.get(table_issue['key'])
                        if not jira_issue:
                            errors.append(f"Задача {table_issue['key']} не найдена в Jira")
                            continue
                        
                        if normalize_text(table_issue['summary']) != normalize_text(jira_issue['summary']):
                            errors.append(f"Описание задачи {table_issue['key']} не совпадает с Jira\nВ документе: '{table_issue['summary']}'\nВ Jira: '{jira_issue['summary']}'")
                        
                        if table_issue['type'].strip() != jira_issue['type'].strip():
                            errors.append(f"Тип задачи {table_issue['key']} не совпадает с Jira")
                break
        
        # Проверка контекста
        if context["RELEASE_VERSION"] == "":
            errors.append("RELEASE_VERSION: не получено значение из Jira")
        if context["PREV_VERSION"] == "":
            errors.append("PREV_VERSION: не заполнено значение")
        if context["RELEASE_ID"] == "":
            errors.append("RELEASE_ID: не заполнено значение")
        if context["OPLOT"] not in OPLOT_VALUES:
            errors.append(f"OPLOT: недопустимое значение '{context['OPLOT']}'")
        if context["CHECKER"] == "":
            errors.append("CHECKER: не заполнено значение")
        
        try:
            datetime.strptime(context["DATE"], "%d.%m.%Y")
        except ValueError:
            errors.append("DATE: неверный формат даты")
        
        try:
            next_day = (datetime.strptime(context["DATE"], "%d.%m.%Y") + timedelta(days=1)).strftime("%d.%m.%Y")
            if context["PLUS_1"] != next_day:
                errors.append(f"PLUS_1: ожидалось {next_day}, получено {context['PLUS_1']}")
        except:
            errors.append("PLUS_1: ошибка вычисления следующего дня")
        
        if playbooks_should_be_checked and context["PLAYBOOKS"] == "":
            errors.append("PLAYBOOKS: не выбраны плейбуки")
        
        if context["INSTRUCTION_BLOCK"] not in ["Выполнить пункты инструкции по внедрению ИНСТРУКЦИЯ", "Отсутствуют"]:
            errors.append("INSTRUCTION_BLOCK: некорректное значение")
        
        if context["POB"] == "":
            errors.append("POB: не получено значение из Jira")
        if context["RELNUMBER"] == "":
            errors.append("RELNUMBER: не заполнено значение")
        elif context["RELNUMBER"] != release_id:
            errors.append(f"RELNUMBER: ожидалось {release_id}, получено {context['RELNUMBER']}")
            
    except Exception as e:
        errors.append(f"Ошибка при проверке документа: {str(e)}")
    
    return errors