from __future__ import annotations

import posixpath
import re
import unicodedata
import zipfile
from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import unquote, urlsplit

from docx import Document
from lxml import etree

from services.ai_planner_document_service import validate_ai_planner_template_structure


MAX_ENTRIES = 2048
MAX_UNCOMPRESSED = 100 * 1024 * 1024
MAX_ENTRY = 32 * 1024 * 1024
MAX_RATIO = 100
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
REQUIRED_PARTS = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
LEGACY_PLACEHOLDERS = (
    "RELEASE_VERSION", "PREV_VERSION", "RELEASE_ID", "OPLOT", "CHECKER",
    "DATE", "PLUS_1", "PLAYBOOKS", "INSTRUCTION_BLOCK", "ИНСТРУКЦИЯ", "POB", "RELNUMBER",
)
PLACEHOLDER_RE = re.compile(
    r"\{\{[^{}]+\}\}|\$\{[^{}]+\}|<<[^<>]+>>|\[\[[^\[\]]+\]\]|\b(?:"
    + "|".join(map(re.escape, LEGACY_PLACEHOLDERS)) + r")\b"
)
FORBIDDEN_PREFIXES = (
    "word/activeX/", "word/embeddings/", "word/vbaProject", "word/oleObject",
)
FORBIDDEN_SUFFIXES = (".exe", ".dll", ".com", ".bat", ".cmd", ".js", ".vbs", ".ps1", ".msi", ".scr")
JIRA_ID_HEADER = "ЗНИ/JIRA ID"
JIRA_SUMMARY_HEADERS = {"issue", "summary", "issue summary", "суть доработки", "описание задачи"}
JIRA_TYPE_HEADERS = {"issue type", "type"}


@dataclass(frozen=True)
class ValidationFailure:
    code: str
    message: str
    group: str

    def as_dict(self) -> dict[str, str]:
        return {"code": self.code, "message": self.message, "group": self.group}


def _unsafe_zip_name(name: str) -> bool:
    normalized = unicodedata.normalize("NFC", name)
    if not normalized or "\0" in normalized or "\\" in normalized or re.match(r"^[A-Za-z]:", normalized) or normalized.startswith("/"):
        return True
    return any(part in {"", ".", ".."} for part in PurePosixPath(normalized).parts)


def _relationship_files(archive: zipfile.ZipFile) -> list[str]:
    return [name for name in archive.namelist() if name.casefold().endswith(".rels")]


def _unique_failures(errors: list[ValidationFailure]) -> list[ValidationFailure]:
    unique = {(item.code, item.message, item.group): item for item in errors}
    return list(unique.values())


def _read_bounded(archive: zipfile.ZipFile, name: str, *, limit: int = MAX_ENTRY) -> bytes:
    """Read one already-vetted ZIP member without trusting its declared size."""
    payload = bytearray()
    with archive.open(name, "r") as source:
        remaining = limit + 1
        while remaining:
            chunk = source.read(min(1024 * 1024, remaining))
            if not chunk:
                break
            payload.extend(chunk)
            remaining -= len(chunk)
    if len(payload) > limit:
        raise ValueError("zip_entry_too_large")
    return bytes(payload)


def _safe_xml(payload: bytes) -> etree._Element:
    if re.search(br"<!\s*(?:DOCTYPE|ENTITY)\b", payload, re.I):
        raise ValueError("xml_dtd")
    parser = etree.XMLParser(resolve_entities=False, no_network=True, load_dtd=False, recover=False, huge_tree=False)
    return etree.fromstring(payload, parser=parser)


def _owner_for_relationship(path: str) -> str:
    if path == "_rels/.rels":
        return ""
    match = re.match(r"^(.*?)/_rels/([^/]+)\.rels$", path)
    return f"{match.group(1)}/{match.group(2)}" if match else ""


def inspect_docx(path: Path | bytes) -> tuple[list[ValidationFailure], set[str]]:
    errors: list[ValidationFailure] = []
    external_images: set[str] = set()
    try:
        archive_source = BytesIO(path) if isinstance(path, bytes) else path
        with zipfile.ZipFile(archive_source) as archive:
            infos = archive.infolist()
            names = [info.filename for info in infos]
            if len(infos) > MAX_ENTRIES:
                errors.append(ValidationFailure("zip_too_many_entries", "В документе слишком много внутренних частей.", "security"))
            seen: set[str] = set()
            total = 0
            total_compressed = 0
            for info in infos:
                key = unicodedata.normalize("NFC", info.filename).casefold()
                if key in seen:
                    errors.append(ValidationFailure("zip_duplicate_name", "DOCX содержит конфликтующие имена внутренних частей.", "security"))
                seen.add(key)
                if _unsafe_zip_name(info.filename):
                    errors.append(ValidationFailure("zip_unsafe_path", "DOCX содержит небезопасный внутренний путь.", "security"))
                total += info.file_size
                total_compressed += info.compress_size
                if info.flag_bits & 1:
                    errors.append(ValidationFailure("zip_encrypted", "Зашифрованные DOCX не поддерживаются.", "security"))
                if info.file_size > MAX_ENTRY:
                    errors.append(ValidationFailure("zip_entry_too_large", "Одна из внутренних частей DOCX слишком велика.", "security"))
                if info.compress_size == 0 and info.file_size > 0 or info.compress_size and info.file_size / info.compress_size > MAX_RATIO:
                    errors.append(ValidationFailure("zip_compression_ratio", "DOCX имеет небезопасную степень сжатия.", "security"))
                lower = info.filename.casefold()
                if lower.startswith(tuple(item.casefold() for item in FORBIDDEN_PREFIXES)) or lower.endswith(FORBIDDEN_SUFFIXES):
                    errors.append(ValidationFailure("forbidden_embedded_content", "Документ содержит неподдерживаемое встроенное содержимое.", "security"))
            if total > MAX_UNCOMPRESSED:
                errors.append(ValidationFailure("zip_uncompressed_too_large", "Распакованный DOCX превышает безопасный размер.", "security"))
            if total > 0 and (total_compressed == 0 or total / total_compressed > MAX_RATIO):
                errors.append(ValidationFailure("zip_overall_compression_ratio", "DOCX имеет небезопасную общую степень сжатия.", "security"))
            if not REQUIRED_PARTS.issubset(set(names)):
                errors.append(ValidationFailure("ooxml_required_parts", "Файл не является полноценным документом Word.", "structure"))

            # Phase A is central-directory-only. Never decompress any member when a
            # size/path/encryption/structure gate has already failed.
            if errors:
                return _unique_failures(errors), external_images

            for rel_path in _relationship_files(archive):
                try:
                    root = _safe_xml(_read_bounded(archive, rel_path))
                except (KeyError, etree.XMLSyntaxError, ValueError):
                    errors.append(ValidationFailure("relationship_xml_invalid", "Связи DOCX повреждены или небезопасны.", "security"))
                    continue
                owner = _owner_for_relationship(rel_path)
                owner_dir = posixpath.dirname(owner)
                for rel in root.xpath("//*[local-name()='Relationship']"):
                    target = str(rel.get("Target") or "")
                    relation_type = str(rel.get("Type") or "").casefold()
                    external = str(rel.get("TargetMode") or "").casefold() == "external"
                    if external:
                        scheme = urlsplit(target).scheme.casefold()
                        if relation_type.endswith("/image"):
                            if scheme not in {"http", "https"}:
                                errors.append(ValidationFailure("external_image_scheme", "Внешнее изображение использует небезопасный адрес.", "security"))
                            external_images.add(unicodedata.normalize("NFC", target))
                        elif relation_type.endswith("/hyperlink") and scheme in {"http", "https", "mailto"}:
                            pass
                        elif relation_type.endswith("/attachedtemplate"):
                            errors.append(ValidationFailure("external_attached_template", "Внешний подключённый шаблон запрещён.", "security"))
                        else:
                            errors.append(ValidationFailure("external_relationship", "DOCX содержит запрещённую внешнюю связь.", "security"))
                    else:
                        decoded_target = unquote(target)
                        resolved = posixpath.normpath(posixpath.join(owner_dir, decoded_target))
                        if "\\" in decoded_target or re.match(r"^[A-Za-z]:", decoded_target) or decoded_target.startswith(("/", "//")) or resolved.startswith("../") or resolved == "..":
                            errors.append(ValidationFailure("relationship_escape", "Внутренняя связь DOCX выходит за границы файла.", "security"))
                        if relation_type.endswith("/afchunk") or relation_type.endswith("/altchunk"):
                            errors.append(ValidationFailure("altchunk", "Встроенные внешние фрагменты не поддерживаются.", "security"))
            for name in names:
                if name.casefold().endswith(".xml") or name.casefold().endswith(".rels"):
                    try:
                        _safe_xml(_read_bounded(archive, name))
                    except (etree.XMLSyntaxError, ValueError):
                        errors.append(ValidationFailure("xml_unsafe", "Одна из XML-частей DOCX повреждена или небезопасна.", "security"))
                        break
    except (zipfile.BadZipFile, OSError):
        errors.append(ValidationFailure("not_docx", "Выбранный файл не является корректным DOCX.", "structure"))
    if not errors:
        try:
            Document(BytesIO(path) if isinstance(path, bytes) else path)
        except Exception:
            errors.append(ValidationFailure("python_docx_reopen", "Документ не открывается стандартным обработчиком Word.", "structure"))
    return _unique_failures(errors), external_images


def _damaged_recognized_placeholders(text: str) -> set[str]:
    damaged: set[str] = set()
    for match in re.finditer(r"[A-ZА-ЯЁ][A-ZА-ЯЁ0-9]*(?:(?:\s+_?|_\s+)[A-ZА-ЯЁ0-9]+)+", text):
        raw = match.group(0)
        canonical = re.sub(r"\s+", "", raw)
        if canonical in LEGACY_PLACEHOLDERS and raw != canonical:
            damaged.add(canonical)
    return damaged


def _paragraph_placeholders(document: Document) -> tuple[Counter, dict[str, Counter], set[str]]:
    total: Counter = Counter()
    areas: dict[str, Counter] = {"body": Counter(), "table": Counter(), "unsupported": Counter()}
    damaged: set[str] = set()
    for paragraph in document.paragraphs:
        values = PLACEHOLDER_RE.findall(paragraph.text)
        total.update(values); areas["body"].update(values)
        damaged.update(_damaged_recognized_placeholders(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    values = PLACEHOLDER_RE.findall(paragraph.text)
                    total.update(values); areas["table"].update(values)
                    damaged.update(_damaged_recognized_placeholders(paragraph.text))
    for section in document.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                values = PLACEHOLDER_RE.findall(paragraph.text)
                total.update(values); areas["unsupported"].update(values)
                damaged.update(_damaged_recognized_placeholders(paragraph.text))
    return total, areas, damaged


def _normalize_jira_header(value: str) -> str:
    return " ".join(unicodedata.normalize("NFC", str(value or "")).split()).casefold()


def _header_index(headers: list[str], accepted: set[str]) -> int | None:
    return next((index for index, header in enumerate(headers) if header in accepted), None)


def _jira_table_signatures(document: Document) -> list[dict[str, Any]]:
    signatures: list[dict[str, Any]] = []
    normalized_id = _normalize_jira_header(JIRA_ID_HEADER)
    for table in document.tables:
        if not table.rows:
            continue
        raw_headers = [unicodedata.normalize("NFC", cell.text).strip() for cell in table.rows[0].cells]
        normalized_headers = [_normalize_jira_header(item) for item in raw_headers]
        jira_id_index = next((index for index, header in enumerate(normalized_headers) if header == normalized_id), None)
        if jira_id_index is None:
            continue
        summary_index = _header_index(normalized_headers, JIRA_SUMMARY_HEADERS)
        type_index = _header_index(normalized_headers, JIRA_TYPE_HEADERS)
        generator_supported = len(raw_headers) >= 4 and jira_id_index == 1 and raw_headers[1] == JIRA_ID_HEADER
        modern_supported = generator_supported and summary_index == 2 and type_index == 3
        normalized_format_changed = any(raw.casefold() != normalized for raw, normalized in zip(raw_headers, normalized_headers))
        signatures.append({
            "normalized_headers": normalized_headers,
            "column_count": len(raw_headers),
            "jira_id_index": jira_id_index,
            "summary_index": summary_index,
            "type_index": type_index,
            "generator_supported": generator_supported,
            "modern_supported": modern_supported,
            "legacy": generator_supported and (not modern_supported or normalized_format_changed),
        })
    return signatures


def _legacy_jira_warning() -> ValidationFailure:
    return ValidationFailure("legacy_jira_table", "Сохранена поддерживаемая историческая структура Jira-таблицы.", "jira")


def _jira_contract_comparison(active: dict[str, Any], candidate: dict[str, Any]) -> tuple[list[ValidationFailure], list[ValidationFailure]]:
    errors: list[ValidationFailure] = []
    warnings: list[ValidationFailure] = []
    active_signatures = list(active.get("jira_signatures") or [])
    candidate_signatures = list(candidate.get("jira_signatures") or [])
    if len(candidate_signatures) < len(active_signatures):
        errors.append(ValidationFailure("jira_table_count", "Нельзя удалять Jira-таблицы действующего шаблона.", "jira"))
    for index, active_signature in enumerate(active_signatures):
        if index >= len(candidate_signatures):
            break
        candidate_signature = candidate_signatures[index]
        active_headers = Counter(active_signature.get("normalized_headers") or [])
        candidate_headers = Counter(candidate_signature.get("normalized_headers") or [])
        loses_active_header = bool(active_headers - candidate_headers)
        if (
            not candidate_signature.get("generator_supported")
            or int(candidate_signature.get("column_count") or 0) < int(active_signature.get("column_count") or 0)
            or loses_active_header
            or (active_signature.get("modern_supported") and not candidate_signature.get("modern_supported"))
        ):
            errors.append(ValidationFailure("jira_table_columns", "Структура Jira-таблицы несовместима с действующим шаблоном или генератором.", "jira"))
        elif candidate_signature.get("legacy"):
            warnings.append(_legacy_jira_warning())
    for candidate_signature in candidate_signatures[len(active_signatures):]:
        if not candidate_signature.get("modern_supported"):
            errors.append(ValidationFailure("jira_table_columns", "Новая Jira-таблица должна соответствовать современной поддерживаемой структуре.", "jira"))
    return _unique_failures(errors), _unique_failures(warnings)


def build_contract(path: Path) -> dict[str, Any]:
    document = Document(path)
    total, areas, damaged = _paragraph_placeholders(document)
    try:
        with zipfile.ZipFile(path) as archive:
            for part_name in ("word/document.xml", "word/footnotes.xml", "word/endnotes.xml"):
                if part_name not in archive.namelist():
                    continue
                xml_root = _safe_xml(_read_bounded(archive, part_name))
                if part_name == "word/document.xml":
                    nodes = xml_root.xpath("//*[local-name()='txbxContent']")
                else:
                    nodes = [xml_root]
                unsupported_values = []
                for node in nodes:
                    logical_text = "".join(node.itertext())
                    unsupported_values.extend(PLACEHOLDER_RE.findall(logical_text))
                    damaged.update(_damaged_recognized_placeholders(logical_text))
                if unsupported_values:
                    total.update(unsupported_values); areas["unsupported"].update(unsupported_values)
    except (OSError, KeyError, zipfile.BadZipFile, etree.XMLSyntaxError, ValueError):
        pass
    jira_signatures = _jira_table_signatures(document)
    jira_warnings = [_legacy_jira_warning()] if any(item.get("legacy") for item in jira_signatures) else []
    return {
        "placeholders": dict(total),
        "areas": {key: dict(value) for key, value in areas.items()},
        "damaged_placeholders": sorted(damaged),
        "jira_table_count": len(jira_signatures),
        "jira_signatures": jira_signatures,
        "jira_blocking_errors": [],
        "jira_warnings": [item.as_dict() for item in jira_warnings],
        "jira_errors": [],
    }


def compare_contract(active_path: Path, candidate_path: Path) -> tuple[list[ValidationFailure], dict[str, Any]]:
    active = build_contract(active_path)
    candidate = build_contract(candidate_path)
    errors: list[ValidationFailure] = []
    # A template revision may intentionally add, remove or move supported
    # placeholders between ordinary paragraphs and table cells. Requiring an
    # exact match with the active document made legitimate edits impossible
    # even though both locations are handled by the generator. Safety is
    # enforced below against actually unsupported areas and then by synthetic
    # generation of the complete candidate.
    if candidate["areas"].get("unsupported"):
        errors.append(ValidationFailure("placeholder_unsupported_area", "Служебные поля в колонтитулах и специальных областях не поддерживаются.", "placeholders"))
    if candidate.get("damaged_placeholders"):
        errors.append(ValidationFailure("placeholder_whitespace_damaged", "Распознанное служебное поле повреждено пробелами.", "placeholders"))
    jira_errors, jira_warnings = _jira_contract_comparison(active, candidate)
    errors.extend(jira_errors)
    candidate["jira_blocking_errors"] = [item.as_dict() for item in jira_errors]
    candidate["jira_warnings"] = [item.as_dict() for item in jira_warnings]
    return errors, candidate


def validate_candidate(active_path: Path, candidate_path: Path) -> dict[str, Any]:
    active_errors, active_external_images = inspect_docx(active_path)
    candidate_errors, candidate_external_images = inspect_docx(candidate_path)
    errors = list(candidate_errors)
    if active_errors:
        errors.append(ValidationFailure("active_template_unavailable", "Действующий шаблон не прошёл контрольное чтение.", "structure"))
    added_images = candidate_external_images - active_external_images
    warnings: list[ValidationFailure] = []
    if added_images:
        errors.append(ValidationFailure("new_external_image", "Новая версия добавляет внешние изображения.", "security"))
    elif candidate_external_images:
        warnings.append(ValidationFailure("grandfathered_external_image", "Сохранены ранее существовавшие внешние изображения; они не загружаются автоматически.", "security"))
    contract: dict[str, Any] = {}
    if not errors:
        try:
            contract_errors, contract = compare_contract(active_path, candidate_path)
            errors.extend(contract_errors)
            warnings.extend(ValidationFailure(**item) for item in contract.get("jira_warnings") or [])
        except Exception:
            errors.append(ValidationFailure("contract_read", "Не удалось проверить структуру служебных полей.", "structure"))
    if not errors and "14061745" in active_path.parent.name:
        try:
            generator_errors = validate_ai_planner_template_structure(
                Document(candidate_path), template_name=active_path.name
            )
            errors.extend(
                ValidationFailure(
                    "ai_planner_generator_contract",
                    message,
                    "generation",
                )
                for message in generator_errors
            )
        except Exception:
            errors.append(ValidationFailure(
                "ai_planner_generator_contract",
                "Не удалось проверить базовые строки Planner/Builder для генерации документов.",
                "generation",
            ))
    return {
        "ok": not errors,
        "errors": [item.as_dict() for item in errors],
        "warnings": [item.as_dict() for item in warnings],
        "contract": contract,
        "checks": {
            "security": not any(item.group == "security" for item in errors),
            "structure": not any(item.group == "structure" for item in errors),
            "placeholders": not any(item.group == "placeholders" for item in errors),
            "jira": not any(item.group == "jira" for item in errors),
            "generation": not any(item.group == "generation" for item in errors),
        },
    }
