import html
import json
import re
import zipfile
from collections import Counter, defaultdict
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Dict, Iterable, List, Optional, Tuple
import xml.etree.ElementTree as ET

from docx import Document
from docx.table import Table

from config import DOC_TEMPLATES_ROOT


CRITICAL_PLACEHOLDERS = [
    "RELEASE_VERSION",
    "PREV_VERSION",
    "OPLOT",
    "CHECKER",
    "DATE",
    "PLUS_1",
    "POB",
    "RELNUMBER",
]

OPTIONAL_PLACEHOLDERS = [
    "RELEASE_ID",
    "INSTRUCTION_BLOCK",
    "PLAYBOOKS",
]

ENHANCEMENT_DEFINITIONS = {
    "instruction_block": {
        "title": "Добавить ключ инструкции",
        "description": "В строках инструкции будет использоваться INSTRUCTION_BLOCK.",
    },
    "rollback_plan": {
        "title": "Добавить план возврата",
        "description": "Раздел будет собран копией плана внедрения с заменой версии на PREV_VERSION.",
    },
    "emergency_rollback_plan": {
        "title": "Добавить план аварийного возврата",
        "description": "Раздел будет собран копией плана внедрения с заменой версии на PREV_VERSION.",
    },
    "success_notification": {
        "title": "Добавить SMS об успешном внедрении",
        "description": "Будет добавлен эталонный текст успешного внедрения.",
    },
    "failure_notification": {
        "title": "Добавить SMS о неуспешном внедрении",
        "description": "Будет добавлен эталонный текст неуспешного внедрения.",
    },
}

SPECIAL_VARIANT_KE = {
    "2256008": ("GREEN", "BLUE"),
    "3894421": ("BH", "PL"),
}

DOC_TYPE_LABELS = {
    "plan": "План внедрения/возврата",
    "checklist": "Чек-лист",
    "checks": "Проверки ПСИ/перечень проверок",
    "unknown": "Не определен",
}

_WORD_XML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_FOLDER_ID_RE = re.compile(r"\s*\((\d{5,})\)\s*$")
_DATE_RE = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")
_RELEASE_KEY_RE = re.compile(r"\b[A-ZА-ЯЁ]{2,12}-\d{3,}\b", re.IGNORECASE)
_VERSION_RE = re.compile(r"\b(?:D|R)-\d{2}\.\d{3}\.\d{2}(?:[-.]\d+)?\b", re.IGNORECASE)
_FIO_RE = re.compile(r"\b[А-ЯЁ][а-яё]+ [А-ЯЁ]\.[А-ЯЁ]\.")
_KE_IN_TEXT_RE = re.compile(r"\((\d{5,})\)")
_URL_RE = re.compile(r"https?://[^\s<>\"]+", re.IGNORECASE)
_INSTRUCTION_ROW_RE = re.compile(r"инструкции\s+(?:до|после)\s+установки\s+релиза", re.IGNORECASE)
_ROLLBACK_PLAN_RE = re.compile(r"(?:план(?:овый)?\s+возврат|план\s+возврата)", re.IGNORECASE)
_EMERGENCY_ROLLBACK_PLAN_RE = re.compile(r"(?:аварийн\w*\s+возврат|план\s+аварийн\w*\s+возврата)", re.IGNORECASE)
_SUCCESS_NOTIFICATION_RE = re.compile(r"оповещени\w+\s+об\s+успешн\w+\s+внедрени\w+\s+релиза", re.IGNORECASE)
_FAILURE_NOTIFICATION_RE = re.compile(r"оповещени\w+\s+о\s*неуспешн\w+\s+внедрени\w+\s+релиза", re.IGNORECASE)


def _safe_filename(filename: str) -> str:
    name = str(filename or "").replace("\\", "/")
    name = PurePosixPath(name).name.strip()
    return name or "document.docx"


def _read_docx_text(data: bytes) -> str:
    parts = []
    with zipfile.ZipFile(BytesIO(data)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            if not any(marker in name for marker in ("document", "header", "footer", "footnotes", "endnotes")):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for node in root.iter(_WORD_XML_NS + "t"):
                if node.text:
                    parts.append(node.text)
    try:
        document = Document(BytesIO(data))
        for paragraph in document.paragraphs:
            if paragraph.text:
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
    except Exception:
        pass
    return "\n".join(parts)


def _normalize_placeholder_text(value: str) -> str:
    text = str(value or "").upper()
    text = re.sub(r"\s*_\s*", "_", text)
    return re.sub(r"\s+", " ", text)


def _strip_folder_id(value: str) -> str:
    return _FOLDER_ID_RE.sub("", str(value or "")).strip()


def _extract_folder_ke(value: str) -> str:
    match = _FOLDER_ID_RE.search(str(value or "").strip())
    return match.group(1) if match else ""


def _normalize_variant(value: str) -> str:
    raw = re.sub(r"\s+", " ", str(value or "").strip()).upper()
    if raw in ("", "ОБЫЧНЫЙ", "ORDINARY", "DEFAULT", "NONE", "NO"):
        return ""
    return raw


def _infer_variant(name: str) -> str:
    normalized = re.sub(r"[^A-ZА-ЯЁ0-9]+", " ", str(name or "").upper()).strip()
    tokens = set(normalized.split())
    for variant in ("GREEN", "BLUE", "BH", "PL"):
        if variant in tokens:
            return variant
    return ""


def _detect_doc_type(filename: str, text: str) -> str:
    searchable = _searchable_text(f"{filename}\n{text}".replace("_", " "))
    if "план внедр" in searchable or "план внедрения" in searchable or "план внедрения и возврата" in searchable:
        return "plan"
    if "чек-лист" in searchable or "чек лист" in searchable or "чеклист" in searchable:
        return "checklist"
    if (
        "выполненные провер" in searchable
        or "перечень провер" in searchable
        or "проверок при проведении" in searchable
    ):
        return "checks"
    return "unknown"


def _find_placeholders(text: str) -> List[str]:
    normalized = _normalize_placeholder_text(text)
    placeholders = []
    for key in CRITICAL_PLACEHOLDERS + OPTIONAL_PLACEHOLDERS:
        if key in normalized:
            placeholders.append(key)
    return placeholders


def _directory_text_and_files(directory: Path) -> Tuple[str, List[Path]]:
    files = sorted(path for path in directory.glob("*.docx") if path.is_file())
    texts = []
    for path in files:
        try:
            texts.append(_read_docx_text(path.read_bytes()))
        except Exception:
            continue
    return "\n".join(texts), files


def _load_manifest(directory: Path) -> Dict:
    manifest_path = directory / "manifest.json"
    if not manifest_path.exists():
        return {}
    try:
        with manifest_path.open("r", encoding="utf-8-sig") as handle:
            data = json.load(handle)
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _load_global_catalog(root: Path) -> Dict[Tuple[str, str], Dict]:
    catalog_path = root / "template_catalog.json"
    if not catalog_path.exists():
        return {}
    try:
        with catalog_path.open("r", encoding="utf-8-sig") as handle:
            raw = json.load(handle)
    except Exception:
        return {}

    entries = raw.get("templates") if isinstance(raw, dict) else raw
    if not isinstance(entries, list):
        return {}

    result = {}
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        category = str(entry.get("category") or "").strip()
        release_full = str(entry.get("release_full") or entry.get("folder") or "").strip()
        if category and release_full:
            result[(category, release_full)] = entry
    return result


def build_template_catalog(root: Path = DOC_TEMPLATES_ROOT) -> List[Dict]:
    catalog = []
    if not root.exists():
        return catalog
    global_catalog = _load_global_catalog(root)

    for directory in sorted(path for path in root.rglob("*") if path.is_dir()):
        files = sorted(directory.glob("*.docx"))
        if not files:
            continue

        relative = directory.relative_to(root)
        if len(relative.parts) < 2:
            continue

        base_category = relative.parts[0]
        release_full = relative.parts[-1]
        file_manifest = _load_manifest(directory)
        manifest = {
            **file_manifest,
            **global_catalog.get((str(file_manifest.get("category") or base_category).strip(), release_full), {}),
        }
        category = str(manifest.get("category") or relative.parts[0]).strip()
        folder_ke = _extract_folder_ke(release_full)
        clean_name = str(manifest.get("name") or _strip_folder_id(release_full) or release_full).strip()
        ke = str(manifest.get("ke") or folder_ke).strip()
        variant = _normalize_variant(manifest.get("variant") or _infer_variant(clean_name))
        text, _ = _directory_text_and_files(directory)
        normalized_text = _normalize_placeholder_text(text)
        aliases = manifest.get("aliases") if isinstance(manifest.get("aliases"), list) else []
        aliases = [str(alias).strip() for alias in aliases if str(alias or "").strip()]

        catalog.append({
            "category": category,
            "release_clean": clean_name,
            "release_full": release_full,
            "ke": ke,
            "variant": variant,
            "requires_playbooks": bool(manifest.get("requires_playbooks"))
            if "requires_playbooks" in manifest
            else "PLAYBOOKS" in normalized_text,
            "requires_instruction": bool(manifest.get("requires_instruction"))
            if "requires_instruction" in manifest
            else "INSTRUCTION_BLOCK" in normalized_text,
            "aliases": list(dict.fromkeys([clean_name, release_full, variant, *aliases])),
            "doc_count": len(files),
            "source": "manifest" if manifest else "folder",
        })

    return catalog


def get_catalog_release_structure() -> Dict[str, List[Tuple[str, str]]]:
    grouped = defaultdict(list)
    for entry in build_template_catalog():
        grouped[entry["category"]].append((entry["release_clean"], entry["release_full"]))
    return {
        category: sorted(values, key=lambda item: item[0].lower())
        for category, values in sorted(grouped.items())
    }


def summarize_template_catalog() -> Dict:
    catalog = build_template_catalog()
    by_category = Counter(entry["category"] for entry in catalog)
    return {
        "total_types": len(catalog),
        "total_documents": sum(entry.get("doc_count", 0) for entry in catalog),
        "with_playbooks": sum(1 for entry in catalog if entry.get("requires_playbooks")),
        "without_playbooks": sum(1 for entry in catalog if not entry.get("requires_playbooks")),
        "categories": dict(sorted(by_category.items())),
    }


def find_template_entries_by_ke(ke: str) -> List[Dict]:
    normalized_ke = str(ke or "").strip()
    if not normalized_ke:
        return []
    return [entry for entry in build_template_catalog() if str(entry.get("ke") or "") == normalized_ke]


def template_requires_playbooks(release_full: str = "", category: str = "") -> Optional[bool]:
    release_full = str(release_full or "").strip()
    category = str(category or "").strip()
    for entry in build_template_catalog():
        if release_full and entry.get("release_full") != release_full:
            continue
        if category and entry.get("category") != category:
            continue
        return bool(entry.get("requires_playbooks"))
    return None


def select_template_by_summary(entries: List[Dict], summary: str) -> Optional[Dict]:
    summary_text = str(summary or "").lower()
    if not entries or not summary_text:
        return None

    def has_token(token: str) -> bool:
        if len(token) <= 3:
            return bool(re.search(rf"(?<![a-zа-яё0-9]){re.escape(token.lower())}(?![a-zа-яё0-9])", summary_text))
        return token.lower() in summary_text

    scored = []
    for entry in entries:
        score = 0
        aliases = [entry.get("release_clean", ""), entry.get("release_full", ""), entry.get("variant", "")]
        aliases.extend(entry.get("aliases") or [])
        for alias in aliases:
            alias_text = str(alias or "").strip().lower()
            if not alias_text:
                continue
            if alias_text in summary_text:
                score += 8
            for token in re.findall(r"[a-zа-яё0-9]+", alias_text, flags=re.IGNORECASE):
                if token in {"обычный", "default"}:
                    continue
                if has_token(token):
                    score += 3 if len(token) <= 3 else 2

        variant = str(entry.get("variant") or "").lower()
        if variant and has_token(variant):
            score += 10

        scored.append((score, entry))

    scored.sort(key=lambda item: item[0], reverse=True)
    if scored and scored[0][0] > 0 and (len(scored) == 1 or scored[0][0] > scored[1][0]):
        return scored[0][1]
    return None


def _metadata_bool(value, default: bool = False) -> bool:
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "on", "да", "y"}


def _find_suspicious_values(text: str) -> Dict[str, List[str]]:
    return {
        "dates": sorted(set(_DATE_RE.findall(text)))[:12],
        "release_keys": sorted(set(match.upper() for match in _RELEASE_KEY_RE.findall(text)))[:12],
        "versions": sorted(set(_VERSION_RE.findall(text)))[:12],
        "fio": sorted(set(_FIO_RE.findall(text)))[:12],
        "ke_like": sorted(set(_KE_IN_TEXT_RE.findall(text)))[:12],
    }


def _compact_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _searchable_text(value: str) -> str:
    text = _compact_text(value).lower()
    text = re.sub(r"\b([а-яё])\s+([а-яё]{2,})\b", r"\1\2", text)
    return text


def _context_contains(text: str, start: int, end: int, markers: Iterable[str], radius: int = 140) -> bool:
    context = text[max(0, start - radius): min(len(text), end + radius)].lower()
    return any(marker.lower() in context for marker in markers)


def _add_replacement_suggestion(
    suggestions: List[Dict],
    source: str,
    target: str,
    reason: str,
    confidence: str = "medium",
    occurrences: int = 0,
) -> None:
    source = str(source or "").strip()
    target = str(target or "").strip()
    if not source or not target or source == target:
        return
    for existing in suggestions:
        if existing["source"] == source:
            return
        if existing["target"] == target and existing["source"] == source:
            return
    suggestions.append({
        "source": source,
        "target": target,
        "reason": reason,
        "confidence": confidence,
        "occurrences": occurrences,
    })


def _extract_contextual_fio(text: str, label_markers: Iterable[str]) -> Optional[str]:
    compact = _compact_text(text)
    normalized_candidates = []
    for match in re.finditer(r"\b[А-ЯЁ][а-яё]+ [А-ЯЁ]\.\s*[А-ЯЁ]\.", compact):
        value = re.sub(r"\.\s+", ".", match.group(0))
        if _context_contains(compact, match.start(), match.end(), label_markers, radius=80):
            normalized_candidates.append(value)
    if normalized_candidates:
        return Counter(normalized_candidates).most_common(1)[0][0]
    return None


def _extract_fio_after_label(text: str, labels: Iterable[str]) -> Optional[str]:
    compact = _compact_text(text)
    values = []
    for label in labels:
        pattern = (
            rf"{re.escape(label)}\s*:?\s*(?:[•\-\u2022]\s*)?"
            r"([А-ЯЁ][а-яё]+ [А-ЯЁ]\.\s*[А-ЯЁ]\.)"
        )
        for match in re.finditer(pattern, compact, flags=re.IGNORECASE):
            values.append(re.sub(r"\.\s+", ".", match.group(1)))
    if values:
        return Counter(values).most_common(1)[0][0]
    return None


def _extract_checker_from_functional_check(text: str) -> Optional[str]:
    compact = _compact_text(text)
    values = []
    pattern = (
        r"Проверка внедренного функционала"
        r"(?P<body>.*?)(?:(?:\s\d+\.\s)|(?:Приложение\s+\d+)|$)"
    )
    for section in re.finditer(pattern, compact, flags=re.IGNORECASE):
        body = section.group("body")
        match = re.search(
            r"Исполнители\s*:?\s*(?:[•\-\u2022]\s*)?([А-ЯЁ][а-яё]+ [А-ЯЁ]\.\s*[А-ЯЁ]\.)",
            body,
            flags=re.IGNORECASE,
        )
        if match:
            values.append(re.sub(r"\.\s+", ".", match.group(1)))
    if values:
        return Counter(values).most_common(1)[0][0]
    return None


def suggest_template_replacements(text: str, found_placeholders: Optional[List[str]] = None) -> List[Dict]:
    found = set(found_placeholders or [])
    suggestions = []
    compact = _compact_text(text)

    versions = Counter(_VERSION_RE.findall(text))
    if versions:
        previous_versions = []
        for match in _VERSION_RE.finditer(text):
            if _context_contains(text, match.start(), match.end(), ("предыдущ", "откат", "возврат")):
                previous_versions.append(match.group(0))
        prev_version = Counter(previous_versions).most_common(1)[0][0] if previous_versions else ""
        release_version = ""
        for version, _count in versions.most_common():
            if version != prev_version:
                release_version = version
                break
        if not release_version:
            release_version = versions.most_common(1)[0][0]
        if "RELEASE_VERSION" not in found:
            _add_replacement_suggestion(
                suggestions,
                release_version,
                "RELEASE_VERSION",
                "самая вероятная версия устанавливаемого дистрибутива",
                "high",
                versions[release_version],
            )
        if prev_version and "PREV_VERSION" not in found:
            _add_replacement_suggestion(
                suggestions,
                prev_version,
                "PREV_VERSION",
                "версия рядом с текстом про откат/предыдущую сборку",
                "high",
                versions[prev_version],
            )
        elif len(versions) > 1 and "PREV_VERSION" not in found:
            for version, _count in versions.most_common():
                if version != release_version:
                    _add_replacement_suggestion(
                        suggestions,
                        version,
                        "PREV_VERSION",
                        "вторая найденная версия, вероятно версия отката",
                        "medium",
                        versions[version],
                    )
                    break

    dates = sorted(
        set(_DATE_RE.findall(text)),
        key=lambda value: tuple(reversed(value.split("."))),
    )
    if dates and "DATE" not in found:
        _add_replacement_suggestion(
            suggestions,
            dates[0],
            "DATE",
            "первая дата работ в документах",
            "medium",
            text.count(dates[0]),
        )
    if len(dates) > 1 and "PLUS_1" not in found:
        _add_replacement_suggestion(
            suggestions,
            dates[1],
            "PLUS_1",
            "следующая дата работ, используется для перехода окна на следующий день",
            "medium",
            text.count(dates[1]),
        )

    release_key_scores = defaultdict(int)
    release_key_counts = Counter(match.upper() for match in _RELEASE_KEY_RE.findall(text))
    for match in _RELEASE_KEY_RE.finditer(text):
        value = match.group(0).upper()
        release_key_scores[value] += 2
        if _context_contains(text, match.start(), match.end(), ("релиз", "релиза", "release")):
            release_key_scores[value] += 8
        if "ЗНИ/JIRA ID" in text[max(0, match.start() - 180):match.start()]:
            release_key_scores[value] -= 3
    if release_key_scores and "RELNUMBER" not in found:
        release_key = max(release_key_scores, key=lambda key: (release_key_scores[key], release_key_counts[key]))
        _add_replacement_suggestion(
            suggestions,
            release_key,
            "RELNUMBER",
            "ключ релиза, найденный в тексте про релиз",
            "medium" if release_key_scores[release_key] < 10 else "high",
            release_key_counts[release_key],
        )

    oplot_fio = _extract_fio_after_label(compact, ("Исполнители", "Исполнитель"))
    if not oplot_fio:
        oplot_fio = _extract_contextual_fio(compact, ("исполнители", "исполнитель"))
    if oplot_fio and "OPLOT" not in found:
        _add_replacement_suggestion(
            suggestions,
            oplot_fio,
            "OPLOT",
            "ФИО рядом с блоком исполнителя",
            "medium",
            compact.count(oplot_fio),
        )

    checker_fio = _extract_checker_from_functional_check(compact)
    checker_reason = "исполнитель шага «Проверка внедренного функционала»"
    checker_confidence = "high"
    if not checker_fio:
        checker_fio = _extract_fio_after_label(compact, ("Проверяющий", "Проверил", "Проверяет"))
        checker_reason = "ФИО рядом с блоком проверяющего"
        checker_confidence = "medium"
    if not checker_fio:
        checker_fio = _extract_contextual_fio(compact, ("проверяющий", "проверил", "проверяет"))
        checker_reason = "ФИО рядом с блоком проверяющего"
        checker_confidence = "medium"
    if checker_fio and "CHECKER" not in found:
        _add_replacement_suggestion(
            suggestions,
            checker_fio,
            "CHECKER",
            checker_reason,
            checker_confidence,
            compact.count(checker_fio),
        )

    pob_candidates = [
        match.group(1).upper()
        for match in re.finditer(
            r"(?:\bРОВ\b|\bPOB\b)[^A-ZА-ЯЁ0-9]{0,30}([A-ZА-ЯЁ]{2,12}-\d{3,})",
            compact,
            flags=re.IGNORECASE,
        )
    ]
    if pob_candidates and "POB" not in found:
        pob = Counter(pob_candidates).most_common(1)[0][0]
        _add_replacement_suggestion(
            suggestions,
            pob,
            "POB",
            "ключ РОВ/POB рядом с соответствующим текстом",
            "medium",
            pob_candidates.count(pob),
        )

    urls = [url.rstrip(".,);") for url in _URL_RE.findall(text)]
    confluence_urls = [url for url in urls if "confluence" in url.lower()]
    if confluence_urls and "INSTRUCTION_BLOCK" not in found:
        url = Counter(confluence_urls).most_common(1)[0][0]
        _add_replacement_suggestion(
            suggestions,
            url,
            "INSTRUCTION_BLOCK",
            "ссылка Confluence относится к блоку инструкции",
            "low",
            confluence_urls.count(url),
        )

    return suggestions


def _build_block_status(text: str, found: Iterable[str]) -> Dict[str, bool]:
    searchable = _searchable_text(text)
    found_set = set(found or [])

    return {
        "has_instruction_rows": bool(_INSTRUCTION_ROW_RE.search(searchable)),
        "has_instruction_key": "INSTRUCTION_BLOCK" in found_set,
        "has_rollback_plan": _has_rollback_heading_text(text, emergency=False),
        "has_emergency_rollback_plan": _has_rollback_heading_text(text, emergency=True),
        "has_success_notification": _has_notification_template_text(searchable, _SUCCESS_NOTIFICATION_RE),
        "has_failure_notification": _has_notification_template_text(searchable, _FAILURE_NOTIFICATION_RE),
    }


def _build_enhancement_suggestions(status: Dict[str, bool]) -> List[Dict]:
    suggestions = []

    def add(key: str, reason: str) -> None:
        definition = ENHANCEMENT_DEFINITIONS[key]
        suggestions.append({
            "key": key,
            "title": definition["title"],
            "description": definition["description"],
            "reason": reason,
            "default": True,
        })

    if not status.get("has_instruction_key"):
        if status.get("has_instruction_rows"):
            add("instruction_block", "Строки инструкции найдены, но в них нет ключа INSTRUCTION_BLOCK.")
        else:
            add("instruction_block", "В плане не найден отдельный блок инструкции. Он должен быть в начале плана.")
    if not status.get("has_rollback_plan"):
        add("rollback_plan", "Не найден отдельный раздел «План возврата».")
    if not status.get("has_emergency_rollback_plan"):
        add("emergency_rollback_plan", "Не найден отдельный раздел «План аварийного возврата».")
    if not status.get("has_success_notification"):
        add("success_notification", "Не найден текст SMS для успешного внедрения.")
    if not status.get("has_failure_notification"):
        add("failure_notification", "Не найден текст SMS для неуспешного внедрения.")
    return suggestions


def _format_replacements(replacements: List[Tuple[str, str]]) -> str:
    return "\n".join(f"{source} => {target}" for source, target in replacements)


def _apply_replacements_to_text(text: str, replacements: List[Tuple[str, str]]) -> str:
    result = str(text or "")
    for old, new in sorted(replacements or [], key=lambda item: len(item[0]), reverse=True):
        if old and new:
            result = result.replace(old, new)
    return result


def analyze_template_package(uploaded_docs: List[Dict], metadata: Optional[Dict] = None) -> Dict:
    metadata = metadata or {}
    documents = []
    errors = []
    warnings = []

    for uploaded in uploaded_docs:
        filename = _safe_filename(uploaded.get("filename"))
        data = uploaded.get("data") or b""
        if not filename.lower().endswith(".docx"):
            continue
        try:
            text = _read_docx_text(data)
        except Exception as exc:
            errors.append(f"{filename}: не удалось прочитать DOCX ({exc})")
            text = ""

        documents.append({
            "filename": filename,
            "type": _detect_doc_type(filename, text),
            "type_label": DOC_TYPE_LABELS.get(_detect_doc_type(filename, text), DOC_TYPE_LABELS["unknown"]),
            "placeholders": _find_placeholders(text),
            "size": len(data),
            "text": text,
        })

    if len(documents) != 3:
        errors.append(f"В пакете должно быть ровно 3 DOCX-документа, найдено: {len(documents)}")

    type_counts = Counter(doc["type"] for doc in documents)
    for doc_type in ("plan", "checklist", "checks"):
        if type_counts[doc_type] == 0:
            errors.append(f"Не найден документ типа: {DOC_TYPE_LABELS[doc_type]}")
        elif type_counts[doc_type] > 1:
            errors.append(f"Найдено несколько документов типа: {DOC_TYPE_LABELS[doc_type]}")
    if type_counts["unknown"]:
        errors.append("Не удалось определить тип одного или нескольких документов")

    combined_text = "\n".join(doc["text"] for doc in documents)
    normalized_text = _normalize_placeholder_text(combined_text)
    found = [key for key in CRITICAL_PLACEHOLDERS + OPTIONAL_PLACEHOLDERS if key in normalized_text]
    suggested_replacements = suggest_template_replacements(combined_text, found)
    covered_by_suggestions = {item["target"] for item in suggested_replacements}
    missing_critical = [
        key for key in CRITICAL_PLACEHOLDERS
        if key not in found and key not in covered_by_suggestions
    ]
    missing_optional = [key for key in OPTIONAL_PLACEHOLDERS if key not in found]
    analysis_replacements = [(item["source"], item["target"]) for item in suggested_replacements]
    analysis_replacements.extend(_parse_replacements(metadata.get("replacements") or ""))
    block_detection_text = _apply_replacements_to_text(combined_text, analysis_replacements)
    block_status = _build_block_status(block_detection_text, _find_placeholders(block_detection_text))
    enhancement_suggestions = _build_enhancement_suggestions(block_status)
    enhancement_keys = {item.get("key") for item in enhancement_suggestions}
    auto_added_placeholders = set()
    if {"rollback_plan", "emergency_rollback_plan"} & enhancement_keys:
        auto_added_placeholders.add("PREV_VERSION")
    if {"success_notification", "failure_notification"} & enhancement_keys:
        auto_added_placeholders.add("POB")
    missing_critical = [key for key in missing_critical if key not in auto_added_placeholders]
    if missing_critical:
        warnings.append(
            "Не удалось автоматически найти значения для маркеров: "
            + ", ".join(missing_critical)
            + ". При необходимости добавьте их в подтвержденные замены вручную."
        )
    if suggested_replacements:
        warnings.append("Конструктор подготовил автозамены. Проверьте их перед скачиванием ZIP-кандидата.")
    if "RELEASE_ID" in missing_optional:
        warnings.append("В шаблонах не найден RELEASE_ID. Это допустимо, если используется RELNUMBER.")

    requires_playbooks = "PLAYBOOKS" in normalized_text
    requires_instruction = True
    suspicious = _find_suspicious_values(combined_text)
    if suspicious["dates"]:
        warnings.append("Найдены похожие на фиксированные даты значения: " + ", ".join(suspicious["dates"]))
    if suspicious["release_keys"]:
        warnings.append("Найдены похожие на номера задач/релизов значения: " + ", ".join(suspicious["release_keys"]))
    if suspicious["versions"]:
        warnings.append("Найдены похожие на фиксированные версии/теги значения: " + ", ".join(suspicious["versions"]))
    if suspicious["fio"]:
        warnings.append("Найдены похожие на фиксированные ФИО значения: " + ", ".join(suspicious["fio"]))

    category = str(metadata.get("category") or "").strip()
    name = str(metadata.get("name") or "").strip()
    ke = str(metadata.get("ke") or "").strip()
    variant = _normalize_variant(metadata.get("variant") or _infer_variant(name))
    existing = find_template_entries_by_ke(ke) if ke else []
    if ke in SPECIAL_VARIANT_KE and variant not in SPECIAL_VARIANT_KE[ke]:
        errors.append(f"Для КЭ {ke} нужно указать вариант: {', '.join(SPECIAL_VARIANT_KE[ke])}")
    if existing:
        same = [
            entry for entry in existing
            if (not category or entry.get("category") == category)
            and (not name or _strip_folder_id(entry.get("release_full", "")).lower() == _strip_folder_id(name).lower())
        ]
        warnings.append(
            f"КЭ {ke} уже есть в каталоге: "
            + "; ".join(f"{entry['category']} / {entry['release_full']}" for entry in existing[:6])
        )
        if same:
            warnings.append("Название похоже на уже существующий шаблон. Проверьте, что это новый вариант, а не дубль.")

    return {
        "documents": [
            {key: value for key, value in doc.items() if key != "text"}
            for doc in documents
        ],
        "found_placeholders": found,
        "missing_critical": missing_critical,
        "missing_optional": missing_optional,
        "requires_playbooks": requires_playbooks,
        "requires_instruction": requires_instruction,
        "block_status": block_status,
        "enhancements": enhancement_suggestions,
        "suspicious": suspicious,
        "suggested_replacements": suggested_replacements,
        "suggested_replacements_text": _format_replacements([
            (item["source"], item["target"]) for item in suggested_replacements
        ]),
        "errors": errors,
        "warnings": list(dict.fromkeys(warnings)),
        "can_build": not errors,
        "summary": {
            "doc_count": len(documents),
            "known_doc_types": sum(1 for doc in documents if doc["type"] != "unknown"),
            "placeholder_count": len(found),
        },
    }


def _parse_replacements(raw: str) -> List[Tuple[str, str]]:
    replacements = []
    for line in str(raw or "").splitlines():
        line = line.strip()
        if not line or "=>" not in line:
            continue
        old, new = line.split("=>", 1)
        old = old.strip()
        new = new.strip()
        if old and new:
            replacements.append((old, new))
    return replacements


def _replace_text_in_paragraph(paragraph, replacements: List[Tuple[str, str]]) -> None:
    runs = list(paragraph.runs)
    if not runs:
        return
    text = "".join(run.text for run in runs)
    new_text = text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == text:
        return

    first_run = runs[0]
    for run in runs:
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(new_text)
    new_run.bold = first_run.bold
    new_run.italic = first_run.italic
    new_run.underline = first_run.underline
    new_run.font.name = first_run.font.name
    new_run.font.size = first_run.font.size
    new_run.font.color.rgb = first_run.font.color.rgb


def _apply_replacements_to_docx(data: bytes, replacements: List[Tuple[str, str]]) -> bytes:
    if not replacements:
        return data

    document = Document(BytesIO(data))
    for paragraph in document.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)

    buffer = BytesIO()
    document.save(buffer)
    return _apply_replacements_to_docx_xml(buffer.getvalue(), replacements)


def _apply_replacements_to_docx_xml(data: bytes, replacements: List[Tuple[str, str]]) -> bytes:
    output = BytesIO()
    with zipfile.ZipFile(BytesIO(data), "r") as source_archive:
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as target_archive:
            for item in source_archive.infolist():
                content = source_archive.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        text = content.decode("utf-8")
                        for old, new in replacements:
                            text = text.replace(old, new)
                        content = text.encode("utf-8")
                    except UnicodeDecodeError:
                        pass
                target_archive.writestr(item, content)
    return output.getvalue()


def _set_cell_text(cell, text: str) -> None:
    cell.text = str(text or "")


def _insert_row_after(table, after_index: int, values: List[str]):
    row = table.add_row()
    for index, value in enumerate(values[:len(row.cells)]):
        _set_cell_text(row.cells[index], value)
    table._tbl.remove(row._tr)
    table.rows[after_index]._tr.addnext(row._tr)
    return row


def _first_plan_table(document):
    for table in document.tables:
        if len(table.rows) < 1:
            continue
        if len(table.rows[0].cells) >= 5:
            return table
    return None


def _row_text(row) -> str:
    return " ".join(cell.text for cell in row.cells)


def _document_searchable_text(document) -> str:
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.append(_row_text(row))
    return _searchable_text("\n".join(parts))


def _has_notification_template_text(searchable: str, pattern: re.Pattern) -> bool:
    for match in pattern.finditer(searchable):
        context = searchable[match.start(): match.start() + 2000]
        normalized_context = re.sub(r"\s*_\s*", "_", context)
        if "relnumber" in normalized_context and "pob" in normalized_context and "release_version" in normalized_context:
            return True
    return False


def _has_rollback_heading_text(text: str, emergency: bool = False) -> bool:
    lines = [_searchable_text(line) for line in str(text or "").splitlines()]
    for line in lines:
        line = line.strip(" .:-")
        if not line:
            continue
        if emergency:
            if line in {"аварийный возврат", "план аварийного возврата"}:
                return True
        elif line in {"плановый возврат", "план возврата"}:
            return True
    return False


def _document_has_rollback_section(document, emergency: bool = False) -> bool:
    for paragraph in document.paragraphs:
        text = _searchable_text(paragraph.text).strip(" .:-")
        if emergency and text in {"аварийный возврат", "план аварийного возврата"}:
            return True
        if not emergency and text in {"плановый возврат", "план возврата"}:
            return True
    return False


def _remove_no_rollback_notes(document) -> None:
    for paragraph in list(document.paragraphs):
        text = _searchable_text(paragraph.text)
        if "откат не предусмотрен" in text or "возврат не предусмотрен" in text:
            paragraph._element.getparent().remove(paragraph._element)


def _ensure_instruction_block(document) -> None:
    changed = False
    for table in document.tables:
        for row in table.rows:
            row_search = _searchable_text(_row_text(row))
            if not _INSTRUCTION_ROW_RE.search(row_search):
                continue
            target_cell = row.cells[1] if len(row.cells) > 1 else row.cells[0]
            if "INSTRUCTION_BLOCK" in target_cell.text:
                changed = True
                continue
            replaced = False
            for paragraph in target_cell.paragraphs:
                if "Отсутствуют" in paragraph.text:
                    _replace_text_in_paragraph(paragraph, [("Отсутствуют", "INSTRUCTION_BLOCK")])
                    replaced = True
                    changed = True
            if not replaced:
                target_cell.add_paragraph("INSTRUCTION_BLOCK")
                changed = True

    if changed:
        return

    table = _first_plan_table(document)
    if not table:
        table = document.add_table(rows=1, cols=6)
        headers = ["№", "Работы", "Дата и время начала", "Дата и время окончания", "ФИО исполнителя", "Примечание"]
        for index, value in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], value)
    values = [
        "1.",
        "Инструкции ДО установки релиза\nINSTRUCTION_BLOCK",
        "DATE\n21:00\n(МСК)",
        "DATE\n21:15\n(МСК)",
        "Ответственный:\nOPLOT\nИсполнители:\nOPLOT",
        "",
    ]
    _insert_row_after(table, 0, values[:len(table.rows[0].cells)])


def _clone_table(document, source_table):
    copied = deepcopy(source_table._tbl)
    document._body._element.append(copied)
    return Table(copied, source_table._parent)


def _find_pre_contacts_anchor(document):
    for paragraph in document.paragraphs:
        text = _searchable_text(paragraph.text).strip(" .:-")
        if text == "контактная информация" or text.startswith("приложение 5"):
            return paragraph._p
    return None


def _insert_paragraph_before(document, anchor, text: str):
    paragraph = document.add_paragraph(text)
    element = paragraph._p
    element.getparent().remove(element)
    if anchor is not None:
        anchor.addprevious(element)
    else:
        document._body._element.append(element)
    return paragraph


def _clone_table_before(document, source_table, anchor):
    copied = deepcopy(source_table._tbl)
    if anchor is not None:
        anchor.addprevious(copied)
    else:
        document._body._element.append(copied)
    return Table(copied, source_table._parent)


def _replace_text_in_cell(cell, replacements: List[Tuple[str, str]]) -> None:
    for paragraph in cell.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)


def _set_notification_row_like_template(row, title: str, text: str) -> None:
    cells = row.cells
    if len(cells) >= 6:
        _set_cell_text(cells[0], "")
        _set_cell_text(cells[1], title)
        for cell in cells[2:]:
            _set_cell_text(cell, text)
        return
    _set_cell_text(cells[0], title)
    if len(cells) > 1:
        _set_cell_text(cells[1], text)


def _format_relative_time(minutes: int) -> str:
    if minutes <= 0:
        return "Х"
    return f"Х+{minutes} минут"


def _simplify_rollback_work_cell(cell) -> None:
    text = cell.text or ""
    searchable = _searchable_text(text)
    if "проверка работоспособности" not in searchable:
        return

    cut_markers = [
        "Ссылки на кластера",
        "Ссылка на кластера",
        "Критерии проверки",
        "все PODы",
    ]
    cut_at = None
    for marker in cut_markers:
        index = text.find(marker)
        if index >= 0 and (cut_at is None or index < cut_at):
            cut_at = index
    if cut_at is None:
        return
    compact = text[:cut_at].strip()
    compact = re.sub(r"\s+", " ", compact)
    if compact:
        _set_cell_text(cell, compact)


def _prepare_rollback_table(table, app_name: str) -> None:
    if not table.rows:
        return

    header = table.rows[0].cells
    if len(header) >= 4:
        _set_cell_text(header[2], "Время начала")
        _set_cell_text(header[3], "Время окончания")

    work_index = 0
    has_failure_notification = False
    for row in table.rows[1:]:
        row_text = _searchable_text(_row_text(row))
        if _SUCCESS_NOTIFICATION_RE.search(row_text) or _FAILURE_NOTIFICATION_RE.search(row_text):
            _set_notification_row_like_template(
                row,
                "Оповещение о неуспешном внедрении релиза:",
                _failure_notification_text(app_name),
            )
            has_failure_notification = True
            continue

        for cell in row.cells:
            _replace_text_in_cell(cell, [("RELEASE_VERSION", "PREV_VERSION")])
        if len(row.cells) > 1:
            _simplify_rollback_work_cell(row.cells[1])

        if len(row.cells) >= 4:
            start = work_index * 10
            _set_cell_text(row.cells[2], _format_relative_time(start))
            _set_cell_text(row.cells[3], _format_relative_time(start + 10))
        work_index += 1

    if not has_failure_notification:
        row = table.add_row()
        _set_notification_row_like_template(
            row,
            "Оповещение о неуспешном внедрении релиза:",
            _failure_notification_text(app_name),
        )


def _add_rollback_section_from_implementation(document, title: str, app_name: str, emergency: bool = False) -> None:
    if emergency and _document_has_rollback_section(document, emergency=True):
        return
    if not emergency and _document_has_rollback_section(document, emergency=False):
        return

    source_table = _first_plan_table(document)
    if source_table is None:
        return

    anchor = _find_pre_contacts_anchor(document)
    _insert_paragraph_before(document, anchor, "")
    _insert_paragraph_before(document, anchor, title)
    table = _clone_table_before(document, source_table, anchor)
    _prepare_rollback_table(table, app_name)


def _success_notification_text(app_name: str) -> str:
    return (
        f"Успешное внедрение релиза АС {app_name}. "
        "Дата: PLUS_1. Версия: RELEASE_VERSION. "
        "Релиз: RELNUMBER, РоВ: POB. "
        "Подсистема функционирует в штатном режиме. Отправитель: OPLOT"
    )


def _failure_notification_text(app_name: str) -> str:
    return (
        f"Произошла ошибка при установке в ПРОМ АС {app_name}. "
        "Дата: PLUS_1. Версия: RELEASE_VERSION. "
        "ПО возвращено на версию PREV_VERSION "
        "Релиз: RELNUMBER, РоВ: POB. "
        "Подсистема функционирует в штатном режиме. Отправитель: OPLOT"
    )


def _find_notification_table(document):
    for table in document.tables:
        for row in table.rows:
            text = _searchable_text(_row_text(row))
            if _SUCCESS_NOTIFICATION_RE.search(text) or _FAILURE_NOTIFICATION_RE.search(text):
                return table
    return None


def _add_notification_row(table, title: str, text: str) -> None:
    row = table.add_row()
    _set_notification_row_like_template(row, title, text)


def _fill_existing_notification_row(document, pattern: re.Pattern, text: str) -> bool:
    for table in document.tables:
        for row in table.rows:
            if not pattern.search(_searchable_text(_row_text(row))):
                continue
            target_index = None
            for index, cell in enumerate(row.cells):
                if pattern.search(_searchable_text(cell.text)):
                    target_index = min(index + 1, len(row.cells) - 1)
                    break
            if target_index is None:
                target_index = 1 if len(row.cells) > 1 else 0
            _set_cell_text(row.cells[target_index], text)
            return True
    return False


def _ensure_notification_blocks(
    document,
    enhancements: set,
    app_name: str,
    include_success: bool = True,
    include_failure: bool = True,
) -> None:
    searchable = _document_searchable_text(document)
    need_success = (
        include_success
        and "success_notification" in enhancements
        and not _has_notification_template_text(searchable, _SUCCESS_NOTIFICATION_RE)
    )
    need_failure = (
        include_failure
        and "failure_notification" in enhancements
        and not _has_notification_template_text(searchable, _FAILURE_NOTIFICATION_RE)
    )
    if not need_success and not need_failure:
        return

    table = _find_notification_table(document)
    if table is None or len(table.rows[0].cells) < 2:
        document.add_paragraph("")
        document.add_paragraph("Оповещения по результатам внедрения")
        table = document.add_table(rows=1, cols=2)
        _set_cell_text(table.rows[0].cells[0], "Событие")
        _set_cell_text(table.rows[0].cells[1], "Текст оповещения")

    if need_success and not _fill_existing_notification_row(
        document,
        _SUCCESS_NOTIFICATION_RE,
        _success_notification_text(app_name),
    ):
        _add_notification_row(
            table,
            "Оповещение об успешном внедрении релиза",
            _success_notification_text(app_name),
        )
    if need_failure and not _fill_existing_notification_row(
        document,
        _FAILURE_NOTIFICATION_RE,
        _failure_notification_text(app_name),
    ):
        _add_notification_row(
            table,
            "Оповещение о неуспешном внедрении релиза",
            _failure_notification_text(app_name),
        )


def _apply_enhancements_to_plan_docx(data: bytes, enhancements: Iterable[str], app_name: str) -> bytes:
    selected = {str(item or "").strip() for item in enhancements if str(item or "").strip()}
    if not selected:
        return data

    document = Document(BytesIO(data))
    if "instruction_block" in selected:
        _ensure_instruction_block(document)
    _ensure_notification_blocks(
        document,
        selected,
        app_name,
        include_success=True,
        include_failure=not ({"rollback_plan", "emergency_rollback_plan"} & selected),
    )
    if {"rollback_plan", "emergency_rollback_plan"} & selected:
        _remove_no_rollback_notes(document)
    if "rollback_plan" in selected:
        _add_rollback_section_from_implementation(document, "Плановый возврат", app_name, emergency=False)
    if "emergency_rollback_plan" in selected:
        _add_rollback_section_from_implementation(document, "Аварийный возврат", app_name, emergency=True)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _normalize_template_filename(filename: str) -> str:
    safe = _safe_filename(filename)
    stem = Path(safe).stem.strip() or "Шаблон"
    if "КЭ" not in stem.upper():
        stem = f"{stem} КЭ"
    return f"{stem}.docx"


def _candidate_folder_name(name: str, ke: str) -> str:
    clean_name = _strip_folder_id(name)
    clean_name = re.sub(r"[\\/]+", " ", clean_name).strip()
    clean_name = clean_name or "NEW_TEMPLATE"
    return f"{clean_name}({ke})"


def _safe_zip_part(value: str) -> str:
    return re.sub(r"[\\/]+", " ", str(value or "")).strip().strip(".") or "template"


def _build_report_html(analysis: Dict, manifest: Dict) -> str:
    def list_items(values: Iterable[str]) -> str:
        items = list(values or [])
        if not items:
            return "<li>Нет</li>"
        return "".join(f"<li>{html.escape(str(item))}</li>" for item in items)

    docs_html = "".join(
        "<tr>"
        f"<td>{html.escape(doc.get('filename', ''))}</td>"
        f"<td>{html.escape(doc.get('type_label', ''))}</td>"
        f"<td>{html.escape(', '.join(doc.get('placeholders') or []))}</td>"
        "</tr>"
        for doc in analysis.get("documents", [])
    )
    enhancements_html = "".join(
        "<li>"
        f"{html.escape(item.get('title', item.get('key', '')))}: "
        f"{html.escape(item.get('reason', ''))}"
        "</li>"
        for item in analysis.get("enhancements", [])
    ) or "<li>Нет</li>"
    return f"""<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <title>Отчет проверки шаблона</title>
  <style>
    body {{ font-family: Arial, sans-serif; color: #1f2937; margin: 32px; }}
    h1, h2 {{ margin-bottom: 8px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #d1d5db; padding: 8px 10px; text-align: left; vertical-align: top; }}
    th {{ background: #eef2ff; }}
    code, pre {{ background: #f3f4f6; padding: 2px 4px; border-radius: 4px; }}
    .ok {{ color: #047857; }}
    .warn {{ color: #b45309; }}
    .err {{ color: #b91c1c; }}
  </style>
</head>
<body>
  <h1>Отчет проверки кандидата шаблона</h1>
  <p>Сформировано: {html.escape(datetime.now().strftime('%d.%m.%Y %H:%M:%S'))}</p>
  <h2>Manifest</h2>
  <pre>{html.escape(json.dumps(manifest, ensure_ascii=False, indent=2))}</pre>
  <h2>Документы</h2>
  <table>
    <thead><tr><th>Файл</th><th>Тип</th><th>Маркеры</th></tr></thead>
    <tbody>{docs_html}</tbody>
  </table>
  <h2>Рекомендованные улучшения</h2>
  <ul>{enhancements_html}</ul>
  <h2>Ошибки</h2>
  <ul class="err">{list_items(analysis.get('errors'))}</ul>
  <h2>Предупреждения</h2>
  <ul class="warn">{list_items(analysis.get('warnings'))}</ul>
  <h2>Итог</h2>
  <p class="{'ok' if analysis.get('can_build') else 'err'}">
    {'Кандидат можно проверять и передавать админу.' if analysis.get('can_build') else 'Кандидат собран с ошибками анализа. Проверьте замечания.'}
  </p>
</body>
</html>"""


def build_template_candidate(uploaded_docs: List[Dict], metadata: Dict) -> BytesIO:
    category = str(metadata.get("category") or "").strip()
    name = str(metadata.get("name") or "").strip()
    ke = str(metadata.get("ke") or "").strip()
    variant = _normalize_variant(metadata.get("variant") or _infer_variant(name))
    aliases = [
        alias.strip()
        for alias in re.split(r"[,;\n]+", str(metadata.get("aliases") or ""))
        if alias.strip()
    ]
    replacements = _parse_replacements(metadata.get("replacements") or "")

    if not category:
        raise ValueError("Укажите категорию шаблона")
    if not name:
        raise ValueError("Укажите название типа релиза")
    if not re.fullmatch(r"\d{5,}", ke):
        raise ValueError("Укажите КЭ релиза числом из 5+ цифр")

    analysis = analyze_template_package(uploaded_docs, {
        "category": category,
        "name": name,
        "ke": ke,
        "variant": variant,
    })
    if analysis.get("errors"):
        raise ValueError("Исправьте ошибки анализа: " + "; ".join(analysis["errors"]))

    form_requires_playbooks = _metadata_bool(metadata.get("requires_playbooks"), analysis["requires_playbooks"])
    if form_requires_playbooks != analysis["requires_playbooks"]:
        if analysis["requires_playbooks"]:
            raise ValueError("В документах есть PLAYBOOKS, поэтому плейбуки должны быть включены")
        raise ValueError("В документах нет PLAYBOOKS, поэтому плейбуки нельзя включить для этого шаблона")

    if ke in SPECIAL_VARIANT_KE and variant not in SPECIAL_VARIANT_KE[ke]:
        raise ValueError(f"Для КЭ {ke} нужно указать вариант: {', '.join(SPECIAL_VARIANT_KE[ke])}")

    if not replacements:
        replacements = [
            (item["source"], item["target"])
            for item in analysis.get("suggested_replacements") or []
        ]
    selected_enhancements = {
        str(item or "").strip()
        for item in (metadata.get("enhancements") or [])
        if str(item or "").strip()
    }

    clean_name = _strip_folder_id(name)
    manifest = {
        "category": category,
        "name": clean_name,
        "ke": ke,
        "variant": variant or "обычный",
        "requires_playbooks": analysis["requires_playbooks"],
        "requires_instruction": True,
        "aliases": list(dict.fromkeys([clean_name, variant, *aliases])),
        "constructor_enhancements": sorted(selected_enhancements),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "created_by": "template-constructor",
    }

    folder = f"{_safe_zip_part(category)}/{_safe_zip_part(_candidate_folder_name(clean_name, ke))}"
    report_html = _build_report_html(analysis, manifest)

    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        used_names = set()
        for uploaded in uploaded_docs:
            filename = _safe_filename(uploaded.get("filename"))
            if not filename.lower().endswith(".docx"):
                continue
            normalized_name = _normalize_template_filename(filename)
            if normalized_name in used_names:
                normalized_name = f"{Path(normalized_name).stem} {len(used_names) + 1}.docx"
            used_names.add(normalized_name)
            docx_data = _apply_replacements_to_docx(uploaded.get("data") or b"", replacements)
            doc_type = _detect_doc_type(filename, _read_docx_text(docx_data))
            if doc_type == "plan":
                docx_data = _apply_enhancements_to_plan_docx(docx_data, selected_enhancements, clean_name)
            archive.writestr(f"{folder}/{normalized_name}", docx_data)

        archive.writestr(
            f"{folder}/manifest.json",
            json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"),
        )
        archive.writestr(
            "template_check_report.html",
            report_html.encode("utf-8"),
        )

    output.seek(0)
    return output
