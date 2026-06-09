import re
import logging
from pathlib import Path
from docx import Document
from config import SMS_TEMPLATES_ROOT

SMS_PROFILE_WHITELIST = ("CLM", "EMRM", "AIST", "AI")


def normalize_sms_profile(profile):
    normalized = str(profile or "").strip().upper()
    if normalized not in SMS_PROFILE_WHITELIST:
        raise ValueError(
            f"Неизвестный профиль SMS: {profile or 'не указан'}. "
            f"Допустимые профили: {', '.join(SMS_PROFILE_WHITELIST)}."
        )
    return normalized


def resolve_sms_profile_template(profile):
    """Resolve a whitelisted logical profile to a CSV inside SMS_TEMPLATES_ROOT."""
    normalized = normalize_sms_profile(profile)
    root = Path(SMS_TEMPLATES_ROOT).resolve()
    if not root.exists() or not root.is_dir():
        raise FileNotFoundError(f"Папка SMS-шаблонов не найдена: {root}")

    for csv_file in sorted(root.glob("*.csv"), key=lambda path: path.name.lower()):
        resolved = csv_file.resolve()
        try:
            resolved.relative_to(root)
        except ValueError:
            continue
        csv_key = str(extract_key_from_filename(csv_file.name) or "").strip().upper()
        if csv_key == normalized:
            return resolved

    raise FileNotFoundError(f"CSV-шаблон для профиля {normalized} не найден.")


def get_sms_profile_availability():
    availability = {}
    for profile in SMS_PROFILE_WHITELIST:
        try:
            resolve_sms_profile_template(profile)
            availability[profile] = True
        except (FileNotFoundError, ValueError):
            availability[profile] = False
    return availability


def find_matching_csv(key):
    """Находит CSV шаблон по ключу из имени файла"""
    if not SMS_TEMPLATES_ROOT.exists():
        logging.error(f"Папка sms не найдена по пути: {SMS_TEMPLATES_ROOT}")
        return None
    
    logging.debug(f"Ищем CSV для ключа ({key}) в папке: {SMS_TEMPLATES_ROOT}")
    for csv_file in SMS_TEMPLATES_ROOT.glob("*.csv"):
        csv_key = extract_key_from_filename(csv_file.name)
        logging.debug(f"Проверяем файл: {csv_file.name}, ключ в файле: {csv_key}")
        if csv_key == key:
            logging.info(f"Найден соответствующий CSV: {csv_file}")
            return csv_file
    logging.warning(f"Шаблон CSV для ключа ({key}) не найден в {SMS_TEMPLATES_ROOT}")
    return None

def extract_notification_text(doc_path, is_failure=False):
    """
    Извлекает текст оповещения из DOCX файла.
    Логика: ищет фразу "оповещение о... внедрении" в ячейке, 
    возвращает текст из следующей ячейки той же строки.
    """
    try:
        doc = Document(doc_path)
        search_phrase = "оповещение о неуспешном внедрении релиза" if is_failure else "оповещение об успешном внедрении релиза"
        logging.info(f"Извлечение текста ({'неуспешное' if is_failure else 'успешное'}) из {doc_path}")
        
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows, 1):
                notification_found = False
                notification_text = None
                
                for cell_idx, cell in enumerate(row.cells):
                    # Объединяем все параграфы ячейки в один текст
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    
                    if not cell_text:
                        continue
                    
                    logging.debug(f"Таблица {table_idx}, строка {row_idx}, ячейка {cell_idx}: {cell_text[:50]}...")
                    
                    # Ищем фразу-якорь
                    if search_phrase in cell_text.lower():
                        notification_found = True
                        logging.info(f"Найдена фраза '{search_phrase}' в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                    # Если фраза уже найдена в этой строке и текущая ячейка не пустая - берем её как текст
                    elif notification_found and not notification_text and cell_text:
                        notification_text = cell_text
                        logging.info(f"Найден текст оповещения в следующей ячейке: {notification_text[:80]}...")
                        break
                
                if notification_text:
                    # Очистка текста
                    notification_text = re.sub(r'\s+', ' ', notification_text).strip()
                    notification_text = notification_text.replace('"', '')  # Удаляем кавычки
                    logging.info(f"Успешно извлечен текст ({'неуспешное' if is_failure else 'успешное'}): {notification_text[:100]}...")
                    return notification_text
        
        logging.warning(f"Текст '{search_phrase}' не найден в: {doc_path}")
        return None
    except Exception as e:
        logging.error(f"Ошибка при чтении документа {doc_path}: {e}")
        return None

def extract_both_notification_texts(doc_path):
    """Извлекает оба текста (успешный и неуспешный) из DOCX файла"""
    return {
        'success': extract_notification_text(doc_path, is_failure=False),
        'failure': extract_notification_text(doc_path, is_failure=True)
    }

def process_csv(template_path, notification_text, output_path):
    """Обрабатывает CSV шаблон, заменяя текст в столбце B"""
    try:
        with open(template_path, 'r', encoding='cp1251') as f:
            lines = f.readlines()
        
        processed_lines = []
        for line in lines:
            clean_line = line.strip()
            if not clean_line:
                processed_lines.append(line)
                continue
            
            if ';' in clean_line:
                parts = clean_line.split(';', 1)
                phone = parts[0].strip()
                new_line = f"{phone};{notification_text}\n"
            else:
                new_line = f"{clean_line};{notification_text}\n"
            
            processed_lines.append(new_line)
        
        with open(output_path, 'w', encoding='cp1251') as f:
            f.writelines(processed_lines)
        logging.info(f"CSV сохранен: {output_path}")
    except Exception as e:
        logging.error(f"Ошибка обработки CSV: {e}")
        raise

def extract_key_from_filename(filename):
    """Извлекает ключ из имени файла (например, (BH) или просто буквы)"""
    # Сначала ищем ключ в скобках как в десктопной версии
    match = re.search(r'\((.*?)\)', filename)
    if match:
        return match.group(1)
    # Затем ищем заглавные буквы
    match = re.search(r'([A-Z]+)', filename)
    return match.group(1) if match else None
