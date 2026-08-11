"""Runtime-only DOCX adaptation for the 14061745 Planner/Builder release family."""

from __future__ import annotations

import re
from copy import deepcopy
from typing import Any

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

PRIMARY_DPM_URL = "https://dpm-sigma.sberbank.ru/dpm/front/main/key/CI14061745"
BUILDER_DPM_URL = "https://dpm-sigma.sberbank.ru/dpm/front/main/key/CI14061745-SUB"


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _replace_token(element, token: str, value: str) -> bool:
    replaced = False
    for node in element.iter(qn("w:t")):
        if token in (node.text or ""):
            node.text = (node.text or "").replace(token, value)
            replaced = True
    return replaced


def _relationship_target(document, relationship_id: str) -> str:
    relationship = document.part.rels.get(relationship_id)
    if relationship is None or relationship.reltype != RT.HYPERLINK:
        return ""
    return str(relationship.target_ref or "")


def _set_dpm_link(document, row_element, target_url: str) -> bool:
    replacement_id = None
    changed = False
    for hyperlink in row_element.iter(qn("w:hyperlink")):
        relationship_id = hyperlink.get(qn("r:id"))
        current_target = _relationship_target(document, relationship_id)
        if not current_target.startswith("https://dpm-sigma.sberbank.ru/dpm/front/main/key/CI14061745"):
            continue
        if replacement_id is None:
            replacement_id = document.part.relate_to(target_url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), replacement_id)
        changed = True
    return changed


def _renumber_table(table) -> None:
    next_number = 1
    for row in table.rows:
        if not row.cells:
            continue
        first_cell = row.cells[0]
        if not re.fullmatch(r"\s*\d+\.\s*", first_cell.text or ""):
            continue
        replacement = f"{next_number}."
        replaced = False
        for node in first_cell._tc.iter(qn("w:t")):
            if not replaced and re.fullmatch(r"\s*\d+\.\s*", node.text or ""):
                node.text = replacement
                replaced = True
        if not replaced and first_cell.paragraphs:
            first_cell.paragraphs[0].text = replacement
        next_number += 1


def _matching_rows(document, token: str, *, require_dpm: bool) -> list[tuple[Any, Any]]:
    matches = []
    for table in document.tables:
        for row in table.rows:
            if token not in _element_text(row._tr):
                continue
            has_dpm = any(
                _relationship_target(document, hyperlink.get(qn("r:id"))).startswith(
                    "https://dpm-sigma.sberbank.ru/dpm/front/main/key/CI14061745"
                )
                for hyperlink in row._tr.iter(qn("w:hyperlink"))
            )
            if has_dpm == require_dpm:
                matches.append((table, row))
    return matches


def _clone_after(table, row, *, token: str, value: str, document, dpm_url: str = "") -> None:
    clone = deepcopy(row._tr)
    if not _replace_token(clone, token, value):
        raise ValueError(f"Шаблон 14061745 повреждён: не найден служебный маркер {token}.")
    if dpm_url and not _set_dpm_link(document, clone, dpm_url):
        raise ValueError("Шаблон 14061745 повреждён: не найдена DPM-ссылка в строке установки.")
    row._tr.addnext(clone)
    _renumber_table(table)


def adapt_ai_planner_document(
    document,
    *,
    template_name: str,
    release_builds: list[dict[str, str]],
    previous_versions: dict[str, str],
):
    """Clone deployment/checklist rows in memory; never mutates source DOCX files."""

    builds = [dict(build) for build in release_builds or []]
    if not builds or not any(build.get("component") in {"e2e_planner", "business_plan_builder"} for build in builds):
        return document

    primary = builds[0]
    secondary = builds[1] if len(builds) > 1 else None
    is_plan = "План внедрения" in str(template_name or "")
    is_checklist = "Чек-лист" in str(template_name or "")

    if is_plan:
        rollout_rows = _matching_rows(document, "RELEASE_VERSION", require_dpm=True)
        rollback_rows = _matching_rows(document, "PREV_VERSION", require_dpm=True)
        if len(rollout_rows) != 1 or not rollback_rows:
            raise ValueError("Шаблон плана 14061745 несовместим с генератором строк установки.")

        primary_dpm = str(primary.get("dpm_url") or PRIMARY_DPM_URL)
        if not _set_dpm_link(document, rollout_rows[0][1]._tr, primary_dpm):
            raise ValueError("Шаблон плана 14061745 не содержит основную DPM-ссылку.")
        for _table, row in rollback_rows:
            if not _set_dpm_link(document, row._tr, primary_dpm):
                raise ValueError("Шаблон отката 14061745 не содержит основную DPM-ссылку.")

        if secondary:
            secondary_component = str(secondary.get("component") or "")
            secondary_version = str(secondary.get("version") or "").strip()
            secondary_previous = str(previous_versions.get(secondary_component) or "").strip()
            if not secondary_version or not secondary_previous:
                raise ValueError("Для Builder необходимо указать текущую и предыдущую версии сборки.")
            secondary_dpm = str(secondary.get("dpm_url") or BUILDER_DPM_URL)
            _clone_after(
                rollout_rows[0][0], rollout_rows[0][1], token="RELEASE_VERSION",
                value=secondary_version, document=document, dpm_url=secondary_dpm,
            )
            for table, row in reversed(rollback_rows):
                _clone_after(
                    table, row, token="PREV_VERSION", value=secondary_previous,
                    document=document, dpm_url=secondary_dpm,
                )

    elif is_checklist and secondary:
        checklist_rows = _matching_rows(document, "RELEASE_VERSION", require_dpm=False)
        if len(checklist_rows) != 1:
            raise ValueError("Чек-лист 14061745 несовместим с генератором двух сборок.")
        _clone_after(
            checklist_rows[0][0], checklist_rows[0][1], token="RELEASE_VERSION",
            value=str(secondary.get("version") or ""), document=document,
        )

    return document


def validate_ai_planner_template_structure(document, *, template_name: str) -> list[str]:
    """DTC-compatible semantic check for the runtime cloning anchors."""

    if "План внедрения" in str(template_name or ""):
        if len(_matching_rows(document, "RELEASE_VERSION", require_dpm=True)) != 1:
            return ["План должен содержать одну строку установки с RELEASE_VERSION и DPM-ссылкой."]
        if not _matching_rows(document, "PREV_VERSION", require_dpm=True):
            return ["План должен содержать строки отката с PREV_VERSION и DPM-ссылкой."]
    if "Чек-лист" in str(template_name or ""):
        if len(_matching_rows(document, "RELEASE_VERSION", require_dpm=False)) != 1:
            return ["Чек-лист должен содержать одну строку проверки с RELEASE_VERSION."]
    return []
