import html
import json
import re
import time
import zipfile
from collections import Counter, defaultdict
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Dict, Iterable, List, Optional, Tuple
import xml.etree.ElementTree as ET

from docx import Document
from docx.table import Table

from config import DOC_TEMPLATES_ROOT


CRITICAL_PLACEHOLDERS = [
    "RELEASE_VERSION",
    "PREV_VERSION",
    "OPLOT",
    "CHECKER",
    "DATE",
    "PLUS_1",
    "POB",
    "RELNUMBER",
]

OPTIONAL_PLACEHOLDERS = [
    "RELEASE_ID",
    "INSTRUCTION_BLOCK",
    "PLAYBOOKS",
]

ENHANCEMENT_DEFINITIONS = {
    "instruction_block": {
        "title": "Добавить ключ инструкции",
        "description": "В строках инструкции будет использоваться INSTRUCTION_BLOCK.",
    },
    "rollback_plan": {
        "title": "Добавить план возврата",
        "description": "Раздел будет собран копией плана внедрения с заменой версии на PREV_VERSION.",
    },
    "emergency_rollback_plan": {
        "title": "Добавить план аварийного возврата",
        "description": "Раздел будет собран копией плана внедрения с заменой версии на PREV_VERSION.",
    },
    "success_notification": {
        "title": "Добавить SMS об успешном внедрении",
        "description": "Будет добавлен эталонный текст успешного внедрения.",
    },
    "failure_notification": {
        "title": "Добавить SMS о неуспешном внедрении",
        "description": "Будет добавлен эталонный текст неуспешного внедрения.",
    },
}

SPECIAL_VARIANT_KE = {
    "2256008": ("GREEN", "BLUE"),
    "3894421": ("BH", "PL"),
}

DOC_TYPE_LABELS = {
    "plan": "План внедрения/возврата",
    "checklist": "Чек-лист",
    "checks": "Проверки ПСИ/перечень проверок",
    "unknown": "Не определен",
}

_WORD_XML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_FOLDER_ID_RE = re.compile(r"\s*\((\d{5,})\)\s*$")
_DATE_RE = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")
_RELEASE_KEY_RE = re.compile(r"\b[A-ZА-ЯЁ]{2,12}-\d{3,}\b", re.IGNORECASE)
_VERSION_RE = re.compile(r"\b(?:D|R)-\d{2}\.\d{3}\.\d{2}(?:[-.]\d+)?\b", re.IGNORECASE)
_FIO_RE = re.compile(r"\b[А-ЯЁ][а-яё]+ [А-ЯЁ]\.[А-ЯЁ]\.")
_KE_IN_TEXT_RE = re.compile(r"\((\d{5,})\)")
_URL_RE = re.compile(r"https?://[^\s<>\"]+", re.IGNORECASE)
_INSTRUCTION_ROW_RE = re.compile(r"инструкции\s+(?:до|после)\s+установки\s+релиза", re.IGNORECASE)
_ROLLBACK_PLAN_RE = re.compile(r"(?:план(?:овый)?\s+возврат|план\s+возврата)", re.IGNORECASE)
_EMERGENCY_ROLLBACK_PLAN_RE = re.compile(r"(?:аварийн\w*\s+возврат|план\s+аварийн\w*\s+возврата)", re.IGNORECASE)
_SUCCESS_NOTIFICATION_RE = re.compile(r"оповещени\w+\s+об\s+успешн\w+\s+внедрени\w+\s+релиза", re.IGNORECASE)
_FAILURE_NOTIFICATION_RE = re.compile(r"оповещени\w+\s+о\s*неуспешн\w+\s+внедрени\w+\s+релиза", re.IGNORECASE)
_CATALOG_CACHE = {}
_RUNTIME_CATALOG_TTL_SECONDS = 300


def is_ai_agents_template_category(category: str = "") -> bool:
    normalized = re.sub(r"[^A-Z0-9]+", "_", str(category or "").upper()).strip("_")
    return normalized in {"AI_AGENTS", "AI_AGENT", "AI_AGENTS_TEMPLATES"}


def _safe_filename(filename: str) -> str:
    name = str(filename or "").replace("\\", "/")
    name = PurePosixPath(name).name.strip()
    return name or "document.docx"


def _read_docx_text(data: bytes, deep: bool = True) -> str:
    parts = []
    with zipfile.ZipFile(BytesIO(data)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            if not any(marker in name for marker in ("document", "header", "footer", "footnotes", "endnotes")):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for node in root.iter(_WORD_XML_NS + "t"):
                if node.text:
                    parts.append(node.text)
    if not deep:
        return "\n".join(parts)

    try:
        document = Document(BytesIO(data))
        for paragraph in document.paragraphs:
            if paragraph.text:
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
    except Exception:
        pass
    return "\n".join(parts)


def _normalize_placeholder_text(value: str) -> str:
    text = str(value or "").upper()
    text = re.sub(r"\s*_\s*", "_", text)
    return re.sub(r"\s+", " ", text)


def _strip_folder_id(value: str) -> str:
    return _FOLDER_ID_RE.sub("", str(value or "")).strip()


def _extract_folder_ke(value: str) -> str:
    match = _FOLDER_ID_RE.search(str(value or "").strip())
    return match.group(1) if match else ""


def _normalize_variant(value: str) -> str:
    raw = re.sub(r"\s+", " ", str(value or "").strip()).upper()
    if raw in ("", "ОБЫЧНЫЙ", "ORDINARY", "DEFAULT", "NONE", "NO"):
        return ""
    return raw


def _infer_variant(name: str) -> str:
    normalized = re.sub(r"[^A-ZА-ЯЁ0-9]+", " ", str(name or "").upper()).strip()
    tokens = set(normalized.split())
    for variant in ("GREEN", "BLUE", "BH", "PL"):
        if variant in tokens:
            return variant
    return ""


def _detect_doc_type(filename: str, text: str) -> str:
    searchable = _searchable_text(f"{filename}\n{text}".replace("_", " "))
    if "план внедр" in searchable or "план внедрения" in searchable or "план внедрения и возврата" in searchable:
        return "plan"
    if "чек-лист" in searchable or "чек лист" in searchable or "чеклист" in searchable:
        return "checklist"
    if (
        "выполненные провер" in searchable
        or "перечень провер" in searchable
        or "проверок при проведении" in searchable
    ):
        return "checks"
    return "unknown"


def _find_placeholders(text: str) -> List[str]:
    normalized = _normalize_placeholder_text(text)
    placeholders = []
    for key in CRITICAL_PLACEHOLDERS + OPTIONAL_PLACEHOLDERS:
        if key in normalized:
            placeholders.append(key)
    return placeholders


def _directory_text_and_files(directory: Path) -> Tuple[str, List[Path]]:
    files = sorted(path for path in directory.glob("*.docx") if path.is_file())
    texts = []
    for path in files:
        try:
            texts.append(_read_docx_text(path.read_bytes(), deep=False))
        except Exception:
            continue
    return "\n".join(texts), files


def _load_manifest(directory: Path) -> Dict:
    manifest_path = directory / "manifest.json"
    if not manifest_path.exists():
        return {}
    try:
        with manifest_path.open("r", encoding="utf-8-sig") as handle:
            data = json.load(handle)
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _load_global_catalog(root: Path) -> Dict[Tuple[str, str], Dict]:
    catalog_path = root / "template_catalog.json"
    if not catalog_path.exists():
        return {}
    try:
        with catalog_path.open("r", encoding="utf-8-sig") as handle:
            raw = json.load(handle)
    except Exception:
        return {}

    entries = raw.get("templates") if isinstance(raw, dict) else raw
    if not isinstance(entries, list):
        return {}

    result = {}
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        category = str(entry.get("category") or "").strip()
        release_full = str(entry.get("release_full") or entry.get("folder") or "").strip()
        if category and release_full:
            result[(category, release_full)] = entry
    return result


def _catalog_signature(root: Path) -> Tuple:
    if not root.exists():
        return ()
    signature = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in {".docx", ".json"}:
            continue
        try:
            stat = path.stat()
            relative = str(path.relative_to(root)).replace("\\", "/")
            signature.append((relative, stat.st_mtime_ns, stat.st_size))
        except OSError:
            continue
    return tuple(sorted(signature))


def clear_template_catalog_cache() -> None:
    _CATALOG_CACHE.clear()


def _count_docx_files_shallow(directory: Path) -> int:
    try:
        return sum(1 for path in directory.iterdir() if path.is_file() and path.suffix.lower() == ".docx")
    except OSError:
        return 0


def build_template_catalog(root: Path = DOC_TEMPLATES_ROOT, deep: bool = True) -> List[Dict]:
    cache_key = f"{'deep' if deep else 'runtime'}:{root.resolve()}"
    cached = _CATALOG_CACHE.get(cache_key)
    if deep:
        signature = _catalog_signature(root)
        if cached and cached.get("signature") == signature:
            return [dict(entry) for entry in cached.get("catalog", [])]
    else:
        now = time.time()
        if cached and cached.get("expires_at", 0) > now:
            return [dict(entry) for entry in cached.get("catalog", [])]
        signature = None

    catalog = []
    if not root.exists():
        return catalog
    global_catalog = _load_global_catalog(root)

    for directory in sorted(path for path in root.rglob("*") if path.is_dir()):
        doc_count = len(sorted(directory.glob("*.docx"))) if deep else _count_docx_files_shallow(directory)
        if not doc_count:
            continue

        relative = directory.relative_to(root)
        if len(relative.parts) < 2:
            continue

        base_category = relative.parts[0]
        release_full = relative.parts[-1]
        file_manifest = _load_manifest(directory)
        manifest = {
            **file_manifest,
            **global_catalog.get((str(file_manifest.get("category") or base_category).strip(), release_full), {}),
        }
        category = str(manifest.get("category") or relative.parts[0]).strip()
        folder_ke = _extract_folder_ke(release_full)
        clean_name = str(manifest.get("name") or _strip_folder_id(release_full) or release_full).strip()
        ke = str(manifest.get("ke") or folder_ke).strip()
        variant = _normalize_variant(manifest.get("variant") or _infer_variant(clean_name))
        normalized_text = ""
        if deep:
            text, _ = _directory_text_and_files(directory)
            normalized_text = _normalize_placeholder_text(text)
        aliases = manifest.get("aliases") if isinstance(manifest.get("aliases"), list) else []
        aliases = [str(alias).strip() for alias in aliases if str(alias or "").strip()]
        if is_ai_agents_template_category(base_category) or is_ai_agents_template_category(category):
            requires_playbooks = False
        elif "requires_playbooks" in manifest:
            requires_playbooks = bool(manifest.get("requires_playbooks"))
        elif deep:
            requires_playbooks = "PLAYBOOKS" in normalized_text
        else:
            requires_playbooks = None

        if "requires_instruction" in manifest:
            requires_instruction = bool(manifest.get("requires_instruction"))
        elif deep:
            requires_instruction = "INSTRUCTION_BLOCK" in normalized_text
        else:
            requires_instruction = None

        catalog.append({
            "category": category,
            "release_clean": clean_name,
            "release_full": release_full,
            "ke": ke,
            "variant": variant,
            "requires_playbooks": requires_playbooks,
            "requires_instruction": requires_instruction,
            "aliases": list(dict.fromkeys([clean_name, release_full, variant, *aliases])),
            "doc_count": doc_count,
            "source": "manifest" if manifest else "folder",
        })

    cache_payload = {"catalog": [dict(entry) for entry in catalog]}
    if deep:
        cache_payload["signature"] = signature
    else:
        cache_payload["expires_at"] = time.time() + _RUNTIME_CATALOG_TTL_SECONDS
    _CATALOG_CACHE[cache_key] = cache_payload
    return catalog


def build_runtime_template_catalog(root: Path = DOC_TEMPLATES_ROOT) -> List[Dict]:
    return build_template_catalog(root=root, deep=False)


def get_catalog_release_structure() -> Dict[str, List[Tuple[str, str]]]:
    grouped = defaultdict(list)
    for entry in build_runtime_template_catalog():
        grouped[entry["category"]].append((entry["release_clean"], entry["release_full"]))
    return {
        category: sorted(values, key=lambda item: item[0].lower())
        for category, values in sorted(grouped.items())
    }


def summarize_template_catalog() -> Dict:
    catalog = build_template_catalog()
    by_category = Counter(entry["category"] for entry in catalog)
    return {
        "total_types": len(catalog),
        "total_documents": sum(entry.get("doc_count", 0) for entry in catalog),
        "with_playbooks": sum(1 for entry in catalog if entry.get("requires_playbooks")),
        "without_playbooks": sum(1 for entry in catalog if not entry.get("requires_playbooks")),
        "categories": dict(sorted(by_category.items())),
    }


def find_template_entries_by_ke(ke: str) -> List[Dict]:
    normalized_ke = str(ke or "").strip()
    if not normalized_ke:
        return []
    return [entry for entry in build_runtime_template_catalog() if str(entry.get("ke") or "") == normalized_ke]


def template_requires_playbooks(release_full: str = "", category: str = "") -> Optional[bool]:
    release_full = str(release_full or "").strip()
    category = str(category or "").strip()
    if is_ai_agents_template_category(category):
        return False
    for entry in build_runtime_template_catalog():
        if release_full and entry.get("release_full") != release_full:
            continue
        if category and entry.get("category") != category:
            continue
        if entry.get("requires_playbooks") is None:
            return None
        return bool(entry.get("requires_playbooks"))
    return None


def select_template_by_summary(entries: List[Dict], summary: str) -> Optional[Dict]:
    summary_text = str(summary or "").lower()
    if not entries or not summary_text:
        return None

    def has_token(token: str) -> bool:
        if len(token) <= 3:
            return bool(re.search(rf"(?<![a-zа-яё0-9]){re.escape(token.lower())}(?![a-zа-яё0-9])", summary_text))
        return token.lower() in summary_text

    scored = []
    for entry in entries:
        score = 0
        aliases = [entry.get("release_clean", ""), entry.get("release_full", ""), entry.get("variant", "")]
        aliases.extend(entry.get("aliases") or [])
        for alias in aliases:
            alias_text = str(alias or "").strip().lower()
            if not alias_text:
                continue
            if alias_text in summary_text:
                score += 8
            for token in re.findall(r"[a-zа-яё0-9]+", alias_text, flags=re.IGNORECASE):
                if token in {"обычный", "default"}:
                    continue
                if has_token(token):
                    score += 3 if len(token) <= 3 else 2

        variant = str(entry.get("variant") or "").lower()
        if variant and has_token(variant):
            score += 10

        scored.append((score, entry))

    scored.sort(key=lambda item: item[0], reverse=True)
    if scored and scored[0][0] > 0 and (len(scored) == 1 or scored[0][0] > scored[1][0]):
        return scored[0][1]
    return None
