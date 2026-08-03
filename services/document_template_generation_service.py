from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from services.docx_service import replace_keys_in_doc
from services.document_template_runtime_service import atomic_write_bytes
from services.document_template_storage_service import candidate_directory
from services.document_template_validation_service import PLACEHOLDER_RE, inspect_docx


SYNTHETIC_JIRA_BASE = "https://example.invalid/jira"
SYNTHETIC_ISSUE = {
    "key": "SYNTHETIC-1",
    "summary": "Синтетическая задача без сетевых запросов",
    "type": "Task",
}


def _synthetic_context(document: Document) -> dict[str, str]:
    keys: set[str] = set()
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        keys.update(PLACEHOLDER_RE.findall(paragraph.text))
    today = datetime(2030, 1, 15)
    conventional = {
        "RELEASE_VERSION": "99.99.99",
        "PREV_VERSION": "99.99.98",
        "RELEASE_ID": "SYNTHETIC-1",
        "OPLOT": "Тестовый редактор Oplot",
        "CHECKER": "Тестовый проверяющий",
        "DATE": today.strftime("%d.%m.%Y"),
        "PLUS_1": (today + timedelta(days=1)).strftime("%d.%m.%Y"),
        "PLAYBOOKS": "SYNTHETIC_PLAYBOOK",
        "INSTRUCTION_BLOCK": "Отсутствуют",
        "ИНСТРУКЦИЯ": "",
        "POB": "Синтетический ПОВ",
        "RELNUMBER": "SYNTHETIC-1",
    }
    context = dict(conventional)
    for key in keys:
        context.setdefault(key, f"Тестовое значение {key[:40]}")
    return context


def _body_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _verify_synthetic_jira(document: Document) -> bool:
    expected_url = f"{SYNTHETIC_JIRA_BASE}/browse/{SYNTHETIC_ISSUE['key']}"
    matching_row = False
    for table in document.tables:
        if not table.rows or "ЗНИ/JIRA ID" not in [cell.text.strip() for cell in table.rows[0].cells]:
            continue
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 4 and cells[1] == SYNTHETIC_ISSUE["key"] and cells[2] == SYNTHETIC_ISSUE["summary"] and cells[3] == SYNTHETIC_ISSUE["type"]:
                matching_row = True
                break
    relationships = [
        rel for rel in document.part.rels.values()
        if rel.reltype == RT.HYPERLINK and rel.is_external and rel.target_ref == expected_url
    ]
    return matching_row and bool(relationships)


def generate_synthetic_document(candidate_path: Path, output: Path) -> tuple[Path | None, list[dict[str, str]]]:
    try:
        document = Document(candidate_path)
        context = _synthetic_context(document)
        generated = replace_keys_in_doc(
            document,
            context,
            [SYNTHETIC_ISSUE],
            "SYNTHETIC-1",
            jira_base_url=SYNTHETIC_JIRA_BASE,
        )
        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        temporary = output.parent / f".{output.name}.generating"
        generated.save(temporary)
        payload = temporary.read_bytes()
        temporary.unlink(missing_ok=True)
        atomic_write_bytes(output, payload)
        errors, _ = inspect_docx(output)
        if errors:
            output.unlink(missing_ok=True)
            return None, [{"code": "generation_security", "message": "Тестовый документ не прошёл контрольное открытие.", "group": "generation"}]
        reopened = Document(output)
        unresolved = []
        for paragraph in _body_paragraphs(reopened):
            unresolved.extend(PLACEHOLDER_RE.findall(paragraph.text))
        if unresolved:
            output.unlink(missing_ok=True)
            return None, [{"code": "generation_placeholders", "message": "В тестовом документе остались незаполненные служебные поля.", "group": "generation"}]
        if not _verify_synthetic_jira(reopened):
            output.unlink(missing_ok=True)
            return None, [{"code": "generation_jira", "message": "Тестовая Jira-строка или безопасная ссылка сформированы неверно.", "group": "generation"}]
        return output, []
    except Exception:
        return None, [{"code": "generation_failed", "message": "Не удалось создать тестовый документ.", "group": "generation"}]


def generate_test_document(candidate_uuid: str, candidate_path: Path) -> tuple[Path | None, list[dict[str, str]]]:
    return generate_synthetic_document(candidate_path, candidate_directory(candidate_uuid) / "test.docx")
